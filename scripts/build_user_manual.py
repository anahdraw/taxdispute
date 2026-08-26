from __future__ import annotations

from datetime import datetime
from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "user_manual_assets"
OUT = DOCS / "RSM_Tax_Dispute_Agentic_Advisor_User_Manual_ID.docx"

RSM_BLUE = "#009CDE"
RSM_GREEN = "#43A047"
RSM_GRAY = "#54585A"
INK = "#2F3340"
MUTED = "#667085"
LINE = "#D7DDE3"
SOFT = "#F4F7F9"
SOFT_BLUE = "#EAF7FD"
SOFT_GREEN = "#EDF8EF"
SOFT_ORANGE = "#FFF7ED"


def font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/Library/Fonts/Arial Bold.ttf" if bold else "/Library/Fonts/Arial.ttf",
        "/System/Library/Fonts/Helvetica.ttc",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except Exception:
            continue
    return ImageFont.load_default()


F = {
    "tiny": font(16),
    "small": font(20),
    "small_b": font(20, True),
    "body": font(24),
    "body_b": font(24, True),
    "h3": font(30, True),
    "h2": font(38, True),
    "h1": font(54, True),
    "metric": font(50, True),
}


def rounded(draw: ImageDraw.ImageDraw, box, fill, outline=LINE, radius=14, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def text(draw: ImageDraw.ImageDraw, xy, value, fill=INK, fnt=None, max_width=None, line_gap=6):
    fnt = fnt or F["body"]
    if not max_width:
        draw.text(xy, value, fill=fill, font=fnt)
        return
    words = value.split()
    lines = []
    line = ""
    for word in words:
        trial = f"{line} {word}".strip()
        if draw.textlength(trial, font=fnt) <= max_width:
            line = trial
        else:
            if line:
                lines.append(line)
            line = word
    if line:
        lines.append(line)
    x, y = xy
    for line in lines:
        draw.text((x, y), line, fill=fill, font=fnt)
        y += fnt.size + line_gap


def rsm_logo(draw: ImageDraw.ImageDraw, x: int, y: int, scale: float = 1.0):
    draw.rectangle((x, y, x + int(22 * scale), y + int(14 * scale)), fill="#8A8F93")
    draw.rectangle((x + int(34 * scale), y, x + int(110 * scale), y + int(14 * scale)), fill=RSM_GREEN)
    draw.rectangle((x + int(124 * scale), y, x + int(295 * scale), y + int(14 * scale)), fill=RSM_BLUE)
    draw.text((x, y + int(28 * scale)), "RSM", fill=RSM_GRAY, font=font(int(74 * scale), True))


def card(draw, x, y, w, h, title, body="", accent=RSM_BLUE, fill="white"):
    rounded(draw, (x, y, x + w, y + h), fill, LINE, 12)
    draw.rounded_rectangle((x, y, x + w, y + 8), radius=8, fill=accent)
    text(draw, (x + 22, y + 30), title, INK, F["h3"], w - 44)
    if body:
        text(draw, (x + 22, y + 80), body, MUTED, F["body"], w - 44)


def pill(draw, x, y, label, fill=SOFT_GREEN, color="#2E7D32"):
    width = int(draw.textlength(label, font=F["small_b"])) + 28
    draw.rounded_rectangle((x, y, x + width, y + 34), radius=17, fill=fill)
    draw.text((x + 14, y + 7), label, fill=color, font=F["small_b"])
    return width


def button(draw, x, y, label, fill=RSM_BLUE, color="white", w=None):
    w = w or int(draw.textlength(label, font=F["body_b"])) + 48
    draw.rounded_rectangle((x, y, x + w, y + 52), radius=10, fill=fill, outline="#B9E5FA")
    tw = draw.textlength(label, font=F["body_b"])
    draw.text((x + (w - tw) / 2, y + 13), label, fill=color, font=F["body_b"])
    return w


def base_canvas(title, subtitle=None):
    img = Image.new("RGB", (1280, 720), "white")
    d = ImageDraw.Draw(img)
    d.rectangle((0, 0, 1280, 720), fill="#FFFFFF")
    rsm_logo(d, 38, 30, 0.72)
    text(d, (300, 42), title, RSM_GRAY, F["h1"], 820)
    if subtitle:
        text(d, (302, 170), subtitle, RSM_GRAY, F["body"], 780)
    return img, d


def save(img: Image.Image, name: str) -> Path:
    ASSETS.mkdir(parents=True, exist_ok=True)
    path = ASSETS / name
    img.save(path, quality=95)
    return path


def fig_login():
    img, d = base_canvas("RSM Tax Dispute Agentic Advisor", "Masuk sebagai Admin atau User sesuai alur kerja.")
    rounded(d, (370, 245, 900, 610), "#FFFFFF", LINE, 14)
    text(d, (410, 275), "Sign in", INK, F["h2"])
    text(d, (410, 335), "Language", MUTED, F["small_b"])
    rounded(d, (410, 365, 860, 420), "white", LINE, 10)
    text(d, (432, 380), "English / Bahasa Indonesia", INK, F["body"])
    button(d, 410, 445, "Admin", RSM_BLUE, "white", 210)
    button(d, 640, 445, "User", "white", RSM_GRAY, 210)
    text(d, (410, 520), "Username: admin", INK, F["body_b"])
    text(d, (410, 562), "Password: Admin@RSM2026", INK, F["body_b"])
    rounded(d, (930, 315, 1215, 520), SOFT_GREEN, "#D7F0DB", 12)
    text(d, (960, 350), "Role tip", "#2E7D32", F["h3"])
    text(d, (960, 400), "Admin dapat mengelola database, peraturan, user, log, dan API check. User fokus pada analisis dan report.", INK, F["small"], 230)
    return save(img, "01_login.png")


def fig_dashboard():
    img, d = base_canvas("Dashboard", "Ringkasan database, coverage ekstraksi, dan komposisi putusan.")
    x0, y0 = 50, 250
    labels = [("Indexed\nDecisions", "305", RSM_BLUE), ("Extraction\nCoverage", "100%", RSM_GREEN), ("VAT/TP\nDocs", "297", RSM_BLUE), ("Local\nRegulations", "23", RSM_GRAY), ("LLM\nLabels", "305", RSM_GRAY)]
    for i, (lab, val, accent) in enumerate(labels):
        x = x0 + i * 240
        rounded(d, (x, y0, x + 210, y0 + 150), "white", LINE, 12)
        draw_y = y0
        d.rounded_rectangle((x, draw_y, x + 210, draw_y + 8), radius=8, fill=accent)
        for line_index, line in enumerate(lab.split("\n")):
            text(d, (x + 22, y0 + 28 + line_index * 24), line, MUTED, F["small_b"])
        text(d, (x + 22, y0 + 92), val, RSM_GRAY, F["h2"])
    rounded(d, (58, 445, 588, 660), "white", LINE, 12)
    text(d, (86, 468), "Outcome distribution", INK, F["h3"])
    colors = [RSM_BLUE, RSM_GREEN, RSM_GRAY, "#8A8F93", "#66C3E5"]
    cx, cy = 225, 560
    for i, col in enumerate(colors):
        d.pieslice((cx - 88, cy - 88, cx + 88, cy + 88), start=i * 72, end=(i + 1) * 72 - 4, fill=col)
    d.ellipse((cx - 48, cy - 48, cx + 48, cy + 48), fill="white")
    text(d, (370, 525), "Pola menang/kalah dan isu utama dipakai sebagai sinyal awal sebelum analisis kasus baru.", MUTED, font(17), 170)
    rounded(d, (620, 445, 1220, 660), "white", LINE, 12)
    text(d, (650, 468), "Top issues", INK, F["h3"])
    for j, (name, pct, col) in enumerate([("Transfer Pricing", 0.86, RSM_GREEN), ("PPN Kurang Bayar", 0.52, RSM_BLUE), ("Pajak Masukan", 0.34, RSM_GRAY)]):
        y = 520 + j * 42
        text(d, (650, y - 4), name, MUTED, F["small_b"])
        d.rounded_rectangle((900, y, 1160, y + 18), radius=9, fill="#E7EDF3")
        d.rounded_rectangle((900, y, 900 + int(260 * pct), y + 18), radius=9, fill=col)
    return save(img, "02_dashboard.png")


def fig_guided():
    img, d = base_canvas("Guided Flow", "Upload dokumen, ekstraksi LLM, analisis risiko, lalu buat Word/PDF.")
    card(d, 55, 240, 520, 395, "Upload + Extract", "Upload PDF putusan atau dokumen sengketa. Aplikasi menyimpan file, melakukan ekstraksi LLM, mengisi parameter kasus, dan menolak duplikat nomor putusan.", RSM_BLUE, SOFT_BLUE)
    rounded(d, (95, 475, 520, 530), "white", LINE, 10)
    text(d, (120, 490), "Pilih PDF/DOCX", INK, F["body_b"])
    button(d, 95, 555, "Upload + Extract", RSM_GREEN, "white", 260)
    card(d, 620, 240, 555, 395, "Analysis Result", "Skor transparan dihitung dari bukti, kelengkapan kasus, putusan pembanding, dasar aturan, dan kesiapan prosedural.", RSM_GREEN)
    for i, (name, val, accent) in enumerate([("Score", "75.5", RSM_BLUE), ("Confidence", "high", RSM_GREEN), ("Evidence", "86", RSM_GRAY)]):
        x = 655 + i * 170
        rounded(d, (x, 410, x + 145, 520), "white", LINE, 10)
        d.rectangle((x, 410, x + 145, 418), fill=accent)
        text(d, (x + 18, 435), name, MUTED, F["small_b"])
        text(d, (x + 18, 468), val, RSM_GRAY, F["h3"])
    rounded(d, (655, 525, 1125, 585), SOFT_GREEN, "#D7F0DB", 10)
    text(d, (680, 542), "Draft report dapat disimpan dan diunduh ulang.", INK, F["body_b"])
    return save(img, "03_guided_flow.png")


def fig_database():
    img, d = base_canvas("Decision Database", "Kelola putusan tersimpan dan buka halaman detail per putusan.")
    rounded(d, (55, 235, 1180, 650), "white", LINE, 14)
    text(d, (85, 265), "Dokumen Tersimpan", INK, F["h2"])
    rounded(d, (85, 330, 1145, 390), "#F8FAFC", LINE, 10)
    text(d, (115, 348), "Total records: 305     Showing 1-10", INK, F["small_b"])
    button(d, 910, 342, "Previous", SOFT_BLUE, "#006F9F", 120)
    button(d, 1040, 342, "Next", SOFT_BLUE, "#006F9F", 90)
    headers = ["Dokumen", "Status", "Confidence", "Nomor Putusan", "Aksi"]
    xs = [110, 410, 565, 720, 1030]
    for x, h in zip(xs, headers):
        text(d, (x, 425), h, MUTED, F["small_b"])
    rows = [
        ("PUT-012095...pdf", "PT Pertamina Power", "Extracted", "100%", "PUT-012095.16/2023", "Detail"),
        ("PUT-011723...pdf", "PT Asricipta Indah", "Extracted", "100%", "PUT-011723.12/2024", "Detail"),
        ("PUT-007235...xlsx", "Yamaha Motor Parts", "Extracted", "43%", "PUT-007235.15/2018", "Re-extract"),
    ]
    for r, (filename, taxpayer, status, confidence, number, action) in enumerate(rows):
        y = 472 + r * 64
        rounded(d, (92, y - 18, 1135, y + 48), "#FFFFFF", "#E5EAF0", 10)
        text(d, (xs[0], y - 4), filename, INK, font(18, True), 240)
        text(d, (xs[0], y + 22), taxpayer, MUTED, font(16, True), 240)
        pill(d, xs[1], y, status, SOFT_GREEN, "#2E7D32")
        pill(d, xs[2], y, confidence, SOFT_GREEN if confidence != "43%" else "#FFF1F2", "#2E7D32" if confidence != "43%" else "#B91C1C")
        text(d, (xs[3], y + 4), number, INK, font(18, True), 235)
        button(d, xs[4] - 5, y - 3, action, SOFT_BLUE, "#006F9F", 120)
    return save(img, "04_database.png")


def fig_detail():
    img, d = base_canvas("Halaman Detail Putusan", "Ringkas satu putusan dalam format satu halaman dengan tab metadata, perhitungan, dan paragraf penting.")
    rounded(d, (55, 235, 1180, 650), "white", LINE, 14)
    text(d, (85, 260), "PUT-012095.16/2023/PP/M.XIIIA", INK, F["h2"])
    text(d, (85, 315), "Tahun 2025 · PPN · DJP win / appeal rejected", MUTED, F["body_b"])
    pill(d, 85, 360, "Metadata penting", SOFT_BLUE, "#006F9F")
    pill(d, 270, 360, "Perhitungan", "#F8FAFC", MUTED)
    pill(d, 420, 360, "Paragraf penting", "#F8FAFC", MUTED)
    pill(d, 620, 360, "100% extraction confidence", SOFT_GREEN, "#2E7D32")
    rounded(d, (85, 420, 570, 615), "#FFFFFF", LINE, 10)
    text(d, (110, 445), "Pemohon Banding / WP", INK, F["body_b"])
    for i, (k, v) in enumerate([("Nama", "PT Pertamina Power Indonesia"), ("NPWP", "80.340.570.3-081.000"), ("Kuasa", "Bunga Wiladatika Savitri")]):
        y = 490 + i * 34
        text(d, (110, y), k, MUTED, F["small_b"])
        text(d, (220, y), v, INK, F["small"], 310)
    rounded(d, (600, 420, 1145, 615), "#FFFFFF", LINE, 10)
    text(d, (625, 445), "Komponen PPN", INK, F["body_b"])
    for i, (k, v) in enumerate([("DPP Pengadilan", "Rp 136.997.791"), ("Pajak Masukan", "Rp 31.213.391"), ("Kurang/Lebih Bayar", "Rp 109.099.175")]):
        y = 490 + i * 34
        text(d, (625, y), k, MUTED, F["small_b"])
        text(d, (860, y), v, INK, F["small_b"])
    return save(img, "05_decision_detail.png")


def fig_dispute_bot():
    img, d = base_canvas("Smart Dispute Bot", "Tanya putusan dan peraturan dalam satu tempat dengan RAG retrieval.")
    card(d, 55, 235, 430, 430, "Pertanyaan", "Contoh: apakah Pertamina pernah menang dalam kasus PPN? Pilih sumber: Putusan + Peraturan, Putusan saja, atau Peraturan saja.", RSM_BLUE)
    rounded(d, (90, 390, 450, 505), "white", LINE, 10)
    text(d, (110, 420), "Tulis pertanyaan pajak atau keyword kasus...", MUTED, F["small"], 300)
    button(d, 90, 540, "Tanya Dispute Analysis", RSM_GREEN, "white", 300)
    rounded(d, (525, 235, 1180, 665), "white", LINE, 14)
    text(d, (555, 265), "Jawaban berbasis konteks retrieval", INK, F["h3"])
    rounded(d, (555, 320, 1145, 455), "#F8FAFC", "#EDF0F4", 10)
    text(d, (585, 350), "Bot menjawab berdasarkan putusan dan aturan paling relevan, lalu menunjukkan sumber RAG, skor cosine, dan ringkasan outcome.", INK, F["body"], 520)
    text(d, (555, 500), "Visualisasi", INK, F["body_b"])
    d.pieslice((590, 535, 720, 665), 0, 150, fill=RSM_GREEN)
    d.pieslice((590, 535, 720, 665), 150, 280, fill=RSM_GRAY)
    d.pieslice((590, 535, 720, 665), 280, 360, fill=RSM_BLUE)
    d.ellipse((628, 573, 682, 627), fill="white")
    text(d, (760, 540), "RAG decision sources\n• similarity %\n• outcome\n• open reference", MUTED, F["small_b"], 300)
    return save(img, "06_dispute_bot.png")


def fig_regulations():
    img, d = base_canvas("Regulations", "Update aturan, import daftar peraturan, dan tanya Smart Regulation Bot.")
    rounded(d, (55, 235, 1180, 655), "white", LINE, 14)
    text(d, (85, 265), "Regulatory Basis", INK, F["h2"])
    for i, name in enumerate(["Bot aturan", "Update", "List aturan", "Manual input"]):
        pill(d, 85 + i * 170, 330, name, SOFT_BLUE if i == 0 else "#F8FAFC", "#006F9F" if i == 0 else MUTED)
    rounded(d, (85, 390, 610, 605), SOFT_BLUE, "#B9E5FA", 10)
    text(d, (115, 420), "Smart Regulation Bot", RSM_BLUE, F["h3"])
    text(d, (115, 475), "Menelaah banyak aturan tersimpan atau satu aturan yang dipilih. Jawaban tetap berbasis sumber.", INK, F["small"], 430)
    button(d, 115, 545, "Ask Regulation Bot", RSM_GREEN, "white", 250)
    rounded(d, (650, 390, 1145, 605), "#FFFFFF", LINE, 10)
    text(d, (680, 420), "Manajemen aturan", INK, F["h3"])
    for i, item in enumerate(["Update dari Ortax berdasarkan topik", "Upload Excel/CSV list aturan dan link", "Input manual jika PDF belum tersedia", "Delete, update, enrich source, buka viewer"]):
        text(d, (700, 470 + i * 34), f"• {item}", MUTED, font(17), 390)
    return save(img, "07_regulations.png")


def fig_reports():
    img, d = base_canvas("Reports", "Database report membuat hasil analisis tidak perlu dibuat ulang.")
    rounded(d, (55, 235, 1180, 655), "white", LINE, 14)
    text(d, (85, 265), "Report Database", INK, F["h2"])
    text(d, (85, 315), "Setelah analisis dibuat, aplikasi menyimpan report key agar dapat dibuka dan diunduh ulang.", MUTED, F["body"], 780)
    headers = ["Tanggal", "Taxpayer", "Nomor kasus", "Bahasa", "Aksi"]
    xs = [90, 245, 520, 800, 950]
    for x, h in zip(xs, headers):
        text(d, (x, 390), h, MUTED, F["small_b"])
    rows = [
        ("29/05/2026", "PT Pertamina Power", "PUT-012095", "ID", "Open / Download"),
        ("28/05/2026", "PT Asricipta Indah", "PUT-011723", "EN", "Open / Download"),
    ]
    for r, row in enumerate(rows):
        y = 440 + r * 70
        d.line((85, y - 18, 1145, y - 18), fill="#E5EAF0", width=1)
        for x, val in zip(xs, row):
            text(d, (x, y), val, INK, F["small_b" if x in (245, 950) else "small"], 220)
    rounded(d, (85, 575, 1145, 625), SOFT_GREEN, "#D7F0DB", 10)
    text(d, (110, 588), "Gunakan Update Analysis jika data berubah; kalau tidak, pakai report tersimpan untuk menghemat token LLM.", INK, F["small_b"])
    return save(img, "08_reports.png")


def fig_admin():
    img, d = base_canvas("Admin Center", "Log aktivitas, user management, dan API check untuk kesiapan aplikasi.")
    rounded(d, (55, 235, 1180, 655), "white", LINE, 14)
    text(d, (85, 265), "Admin", INK, F["h2"])
    for i, name in enumerate(["Log aktivitas", "User management", "Check API"]):
        pill(d, 85 + i * 200, 330, name, SOFT_BLUE if i == 0 else "#F8FAFC", "#006F9F" if i == 0 else MUTED)
    card(d, 85, 390, 340, 205, "Logs", "Mencatat login, upload, ekstraksi, chatbot, report, regulation update, dan delete.", RSM_BLUE)
    card(d, 465, 390, 340, 205, "Users", "Tambah, edit, nonaktifkan, atau hapus user demo admin/user.", RSM_GREEN)
    card(d, 845, 390, 300, 205, "API", "Cek OpenAI, Blob, database, health page, dan jumlah tabel utama.", RSM_GRAY)
    return save(img, "09_admin.png")


def make_figures():
    return [
        fig_login(),
        fig_dashboard(),
        fig_guided(),
        fig_database(),
        fig_detail(),
        fig_dispute_bot(),
        fig_regulations(),
        fig_reports(),
        fig_admin(),
    ]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_border(cell, color="D7DDE3"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def style_table(table, header=True):
    table.style = "Table Grid"
    for r, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            if header and r == 0:
                set_cell_shading(cell, "#F2F6F9")
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(84, 88, 90)


def set_font(run, size=10.5, bold=None, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def add_para(doc, text_value, size=10.5, bold=False, color=INK, after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.12
    if align:
        p.alignment = align
    r = p.add_run(text_value)
    set_font(r, size, bold, color)
    return p


def add_heading(doc, value, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(value)
    set_font(run, 16 if level == 1 else 13 if level == 2 else 11.5, True, RSM_BLUE if level <= 2 else RSM_GRAY)
    return p


def add_bullet(doc, value):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(value)
    set_font(run, 10.3, False, INK)
    return p


def add_number(doc, value):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    run = p.add_run(value)
    set_font(run, 10.3, False, INK)
    return p


def add_callout(doc, title, body, fill="EDF8EF", accent=RSM_GREEN):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.45])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, accent.replace("#", ""))
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_font(r, 10.5, True, RSM_GRAY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_font(r2, 9.5, False, INK)
    add_para(doc, "", after=2)


def add_screenshot(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.45))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_font(r, 8.5, False, MUTED)


def add_kv_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Topik"
    hdr[1].text = "Ringkasan"
    for key, val in rows:
        cells = table.add_row().cells
        cells[0].text = key
        cells[1].text = val
    set_table_width(table, [1.7, 4.75])
    style_table(table)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    set_font(run, 9.5, run.bold, RSM_GRAY if row == table.rows[0] else INK)
    add_para(doc, "", after=2)


def set_doc_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("RSM Tax Dispute Agentic Advisor | User Manual | Prototype guidance, subject to advisor review")
    set_font(r, 8.5, False, MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    r = p.add_run("RSM")
    set_font(r, 42, True, RSM_GRAY)
    p2 = doc.add_paragraph()
    for color, width in [(RSM_GRAY, 8), (RSM_GREEN, 26), (RSM_BLUE, 70)]:
        run = p2.add_run(" " * width)
        run.font.highlight_color = None
        set_font(run, 6, False, color)
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(10)
    tr = title.add_run("RSM Tax Dispute Agentic Advisor\nUser Manual")
    set_font(tr, 28, True, RSM_GRAY)
    subtitle = doc.add_paragraph()
    sr = subtitle.add_run("Panduan penggunaan aplikasi untuk upload putusan, ekstraksi data, analisis sengketa, RAG chatbot, pengelolaan peraturan, report, dan administrasi.")
    set_font(sr, 12, False, INK)
    add_callout(
        doc,
        "Dokumen untuk pengguna bisnis pajak",
        "Manual ini ditulis untuk tim pajak/advisor non-teknis. Beberapa istilah teknis seperti RAG dan LLM dijelaskan dalam bahasa operasional.",
        "EDF8EF",
        RSM_GREEN,
    )
    meta = [
        ("Aplikasi", "RSM Tax Dispute Agentic Advisor"),
        ("Lingkup", "Prototype Vercel/Next.js dan workflow analisis sengketa pajak"),
        ("Bahasa", "Indonesia"),
        ("Tanggal", datetime.now().strftime("%d %B %Y")),
        ("Akun demo Admin", "admin / Admin@RSM2026"),
        ("Akun demo User", "user / User@RSM2026"),
    ]
    add_kv_table(doc, meta)
    doc.add_page_break()


def add_toc(doc):
    add_heading(doc, "Daftar Isi", 1)
    for item in [
        "1. Gambaran umum aplikasi",
        "2. Login, role, dan navigasi",
        "3. Dashboard",
        "4. Guided Flow: upload, ekstraksi, analisis, dan report",
        "5. Decision Database dan halaman detail putusan",
        "6. Dispute Analysis / Smart Dispute Bot",
        "7. Regulations dan Smart Regulation Bot",
        "8. Reports database",
        "9. Admin Center: log, user management, dan API check",
        "10. Praktik penggunaan yang disarankan",
        "11. Troubleshooting singkat",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()


def build_doc(figures):
    doc = Document()
    set_doc_styles(doc)
    add_cover(doc)
    add_toc(doc)

    add_heading(doc, "1. Gambaran Umum Aplikasi", 1)
    add_para(
        doc,
        "RSM Tax Dispute Agentic Advisor adalah aplikasi prototype untuk membantu tim pajak melakukan ekstraksi dokumen sengketa, mencari putusan pembanding, menelaah regulasi, membuat analisis risiko, dan menyusun draft rekomendasi Word/PDF untuk review advisor.",
    )
    add_callout(
        doc,
        "Prinsip penting",
        "Aplikasi membantu mempercepat analisis awal. Hasil akhir tetap harus direview oleh advisor pajak karena skor, rekomendasi, dan chatbot bersifat indikatif.",
        "EAF7FD",
        RSM_BLUE,
    )
    add_kv_table(
        doc,
        [
            ("LLM extraction", "Membaca dokumen dan mengisi field seperti nomor putusan, pihak, objek sengketa, argumen, pertimbangan, outcome, dan komponen PPN."),
            ("RAG chatbot", "Mengambil konteks dari putusan/peraturan yang paling relevan terlebih dahulu, lalu menjawab dengan sumber yang ditampilkan."),
            ("Report database", "Menyimpan hasil analisis agar tidak perlu melakukan analisis ulang jika dokumen dan bahasa report sama."),
            ("Admin Center", "Mencatat log, mengelola user demo, dan memeriksa kesiapan OpenAI, Blob, database, serta health API."),
        ],
    )

    add_heading(doc, "2. Login, Role, dan Navigasi", 1)
    add_screenshot(doc, figures[0], "Gambar 1. Halaman login dan pemilihan role.")
    add_para(doc, "Aplikasi memiliki dua role demo: Admin dan User. Admin dapat mengakses seluruh menu termasuk Decision Database, Regulations, Reports, dan Admin Center. User difokuskan pada dashboard, guided flow, dispute analysis, dan report.")
    for step in [
        "Pilih bahasa antarmuka: English atau Bahasa Indonesia.",
        "Pilih role Admin atau User.",
        "Isi username dan password demo.",
        "Klik Sign in. Setelah berhasil, menu samping akan muncul sesuai hak akses role.",
    ]:
        add_bullet(doc, step)

    add_heading(doc, "3. Dashboard", 1)
    add_screenshot(doc, figures[1], "Gambar 2. Dashboard ringkasan database dan visualisasi.")
    add_para(doc, "Dashboard dipakai untuk melihat kesiapan data secara cepat: jumlah putusan terindeks, coverage ekstraksi, jumlah dokumen PPN/Transfer Pricing, aturan lokal, label LLM, distribusi outcome, dan isu dominan.")
    add_bullet(doc, "Gunakan dashboard sebelum analisis untuk memahami apakah database pembanding sudah cukup kaya untuk topik sengketa tertentu.")
    add_bullet(doc, "Angka dashboard bersifat dinamis mengikuti data yang tersimpan di database aplikasi.")

    add_heading(doc, "4. Guided Flow: Upload, Ekstraksi, Analisis, dan Report", 1)
    add_screenshot(doc, figures[2], "Gambar 3. Guided Flow untuk upload, ekstraksi, dan analisis.")
    add_para(doc, "Guided Flow adalah jalur kerja utama untuk menganalisis satu kasus wajib pajak. User dapat mengunggah dokumen, menjalankan ekstraksi LLM, meninjau parameter, membaca analisis, lalu membuat report.")
    add_heading(doc, "4.1 Langkah penggunaan", 2)
    for step in [
        "Upload dokumen PDF/DOCX yang ingin dianalisis.",
        "Klik Upload + Extract atau Extract PDF dengan LLM.",
        "Periksa hasil ekstraksi: nama WP, nomor putusan/nomor kasus, jenis pajak, masa pajak, nilai sengketa, posisi DJP, posisi WP, bukti, dan outcome.",
        "Jalankan analisis untuk memperoleh skor, confidence, evidence score, faktor pendukung, faktor risiko, celah bukti, dasar peraturan, dan putusan pembanding.",
        "Jika hasil sudah layak, klik Generate Word/PDF. Report akan tersimpan di database report.",
    ]:
        add_bullet(doc, step)
    add_heading(doc, "4.2 Cara membaca skor", 2)
    add_kv_table(
        doc,
        [
            ("Evidence strength", "Kekuatan bukti: invoice, SPT, pembayaran, rekonsiliasi, korespondensi, kontrak, atau dokumen pendukung."),
            ("Case specificity", "Seberapa lengkap field kasus: taxpayer, jenis pajak, issue, nilai sengketa, masa pajak, nomor SKP/KEP, dan posisi para pihak."),
            ("Comparable support", "Seberapa kuat putusan pembanding mendukung posisi WP atau menunjukkan risiko."),
            ("Regulatory basis", "Ketersediaan aturan relevan dan dasar hukum di dokumen."),
            ("Procedural readiness", "Kesiapan prosedural seperti tahap banding/keberatan, dokumen formal, dan dukungan pembuktian."),
        ],
    )

    add_heading(doc, "5. Decision Database dan Halaman Detail Putusan", 1)
    add_screenshot(doc, figures[3], "Gambar 4. Decision Database dan daftar dokumen tersimpan.")
    add_para(doc, "Decision Database menyimpan dokumen putusan sebagai basis pencarian pembanding. Dari tabel, user dapat melihat status ekstraksi, confidence, nomor putusan, taxpayer, ukuran file, tanggal upload, dan aksi.")
    add_bullet(doc, "Klik file atau nomor putusan untuk membuka halaman detail putusan.")
    add_bullet(doc, "Gunakan Re-extract jika confidence rendah atau data penting masih kosong.")
    add_bullet(doc, "Gunakan Delete hanya untuk dokumen yang salah upload atau duplikat.")
    add_screenshot(doc, figures[4], "Gambar 5. Halaman detail putusan dengan tab metadata, perhitungan, dan paragraf penting.")
    add_heading(doc, "5.1 Tab pada halaman detail", 2)
    add_kv_table(
        doc,
        [
            ("Metadata penting", "Outcome, klasifikasi menang/kalah, jenis pajak, masa pajak, nomor putusan, pihak, hakim, panitera, nilai sengketa, dan objek koreksi."),
            ("Perhitungan", "Komponen PPN seperti DPP, pajak keluaran, pajak masukan, kurang/lebih bayar, kompensasi, sanksi, dan nilai final."),
            ("Paragraf penting", "Ringkasan posisi WP, posisi DJP, pertimbangan majelis, amar, dan kutipan penting untuk review advisor."),
        ],
    )
    add_para(doc, "Jika PDF tersedia di Blob, tombol Open Original PDF/Reference membuka PDF viewer sehingga user dapat mencari keyword langsung di dokumen sumber.")

    add_heading(doc, "6. Dispute Analysis / Smart Dispute Bot", 1)
    add_screenshot(doc, figures[5], "Gambar 6. Dispute Analysis dengan RAG chatbot, sumber, dan visualisasi.")
    add_para(doc, "Dispute Analysis menggantikan menu pencarian kasus mirip. User dapat bertanya dalam bahasa natural atau memasukkan keyword kasus untuk mencari putusan/peraturan yang relevan.")
    add_heading(doc, "6.1 Cara kerja singkat RAG", 2)
    for item in [
        "Aplikasi menghitung relevansi awal dengan metode retrieval/cosine untuk memilih konteks yang paling mirip.",
        "LLM hanya menerima konteks terpilih sehingga token lebih hemat.",
        "Jawaban menampilkan sumber RAG, termasuk skor cosine, ringkasan putusan, dan link referensi.",
        "Jika pertanyaan bersifat statistik seperti menang/kalah, aplikasi menampilkan visualisasi outcome dari matched decisions.",
    ]:
        add_bullet(doc, item)
    add_callout(doc, "Catatan relevansi", "Skor cosine rendah dapat terjadi bila pertanyaan terlalu umum atau istilahnya tidak sama dengan metadata/isi putusan. Gunakan keyword yang lebih spesifik, misalnya nama WP, jenis pajak, objek sengketa, atau istilah teknis.", "FFF7ED", "#F97316")

    add_heading(doc, "7. Regulations dan Smart Regulation Bot", 1)
    add_screenshot(doc, figures[6], "Gambar 7. Menu Regulations dengan tabs Bot, Update, List, dan Manual.")
    add_para(doc, "Menu Regulations mengelola basis pengetahuan peraturan. User dapat update dari Ortax, upload daftar peraturan Excel/CSV, input manual ringkasan aturan, memperkaya sumber, membuka viewer, menghapus aturan, dan bertanya ke Smart Regulation Bot.")
    add_heading(doc, "7.1 Format upload Excel/CSV peraturan", 2)
    add_kv_table(
        doc,
        [
            ("title/name", "Nama peraturan atau judul kartu aturan."),
            ("citation/number", "Nomor peraturan, misalnya PMK, PER, SE, UU, atau PP."),
            ("topic/topik", "Topik seperti PPN, Transfer Pricing, General, atau topik lain yang didukung."),
            ("focus/summary", "Ringkasan singkat isi atau relevansi peraturan."),
            ("sourceUrl/link", "Link sumber peraturan, misalnya dari Ortax atau sumber resmi lain."),
            ("content/notes", "Catatan isi aturan. Semakin jelas catatan, semakin baik kualitas jawaban chatbot."),
            ("relevance", "Opsional, angka relevansi/priority untuk membantu pengurutan."),
        ],
    )

    add_heading(doc, "8. Reports Database", 1)
    add_screenshot(doc, figures[7], "Gambar 8. Report database untuk membuka dan mengunduh ulang report.")
    add_para(doc, "Reports Database menyimpan report yang sudah pernah dibuat. Ini berguna agar analisis tidak perlu dijalankan ulang kecuali dokumen, ekstraksi, atau instruksi report berubah.")
    add_bullet(doc, "Klik report untuk melihat detail analisis yang pernah dibuat.")
    add_bullet(doc, "Gunakan Download Word/PDF untuk mengunduh ulang hasilnya.")
    add_bullet(doc, "Gunakan Update Analysis jika ada data baru, peraturan baru, atau ekstraksi diperbaiki.")

    add_heading(doc, "9. Admin Center", 1)
    add_screenshot(doc, figures[8], "Gambar 9. Admin Center untuk log, user management, dan API check.")
    add_para(doc, "Admin Center hanya tersedia untuk role Admin. Menu ini membantu mengawasi penggunaan aplikasi dan memastikan konfigurasi teknis siap.")
    add_kv_table(
        doc,
        [
            ("Activity logs", "Mencatat login, upload, extraction, chatbot, report, regulation update, delete, dan aksi penting lain."),
            ("User management", "Membuat, mengedit, menonaktifkan, atau menghapus user demo."),
            ("API check", "Memeriksa status OpenAI, Vercel Blob, database, tabel utama, dan health page."),
            ("Health page", "Buka /api/health untuk tampilan status yang rapi atau /api/health?format=json untuk JSON mentah."),
        ],
    )

    add_heading(doc, "10. Praktik Penggunaan yang Disarankan", 1)
    for item in [
        "Mulai dari Dashboard untuk melihat kesiapan data dan topik dominan.",
        "Upload dokumen sumber di Decision Database jika putusan akan menjadi bagian dari database pembanding.",
        "Gunakan Guided Flow untuk analisis satu kasus WP dan pembuatan report.",
        "Gunakan Dispute Analysis untuk pertanyaan lintas putusan/peraturan atau pencarian kasus mirip.",
        "Update Regulations sebelum analisis jika isu sangat bergantung pada aturan terbaru atau topik transfer pricing/PPN tertentu.",
        "Review semua hasil LLM secara profesional sebelum mengirimkan rekomendasi ke wajib pajak.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "11. Troubleshooting Singkat", 1)
    add_kv_table(
        doc,
        [
            ("Upload gagal", "Periksa ukuran file, format PDF/DOCX, dan status Blob di Admin > API check."),
            ("Ekstraksi kosong", "Jalankan Re-extract, gunakan file yang teksnya dapat dibaca, atau cek apakah PDF hasil scan perlu OCR/vision processing."),
            ("Jawaban chatbot kurang relevan", "Persempit pertanyaan dengan nama WP, nomor putusan, jenis pajak, isu, objek koreksi, atau tahun pajak."),
            ("Report tidak berubah", "Pastikan melakukan Update Analysis atau hapus report lama jika ingin membuat ulang dari ekstraksi terbaru."),
            ("Tidak bisa melihat menu admin", "Login dengan role Admin. User biasa hanya melihat menu yang relevan untuk analisis."),
            ("Health page masih JSON", "Buka /api/health untuk tampilan rapi. Buka /api/health?format=json hanya jika membutuhkan payload JSON."),
        ],
    )

    add_heading(doc, "12. Checklist Sebelum Dipakai untuk Review Advisor", 1)
    for item in [
        "Dokumen utama sudah terupload dan status extraction minimal memadai.",
        "Nomor putusan/nomor kasus, nama WP, jenis pajak, nilai sengketa, dan posisi para pihak sudah terisi.",
        "Putusan pembanding dibaca ulang, bukan hanya dilihat outcome-nya.",
        "Dasar peraturan sudah diperiksa dan relevan dengan masa/tahun pajak.",
        "Celah bukti dan risiko sudah ditindaklanjuti dengan dokumen pendukung.",
        "Draft Word/PDF sudah direview manual sebelum dikirim ke pihak eksternal.",
    ]:
        add_bullet(doc, item)

    doc.save(OUT)
    return OUT


def main():
    DOCS.mkdir(parents=True, exist_ok=True)
    figures = make_figures()
    output = build_doc(figures)
    print(output)


if __name__ == "__main__":
    main()
