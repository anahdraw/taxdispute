import json
import math
import os
import re
import base64
import io
import sqlite3
import subprocess
import tempfile
import textwrap
import uuid
from collections import Counter, defaultdict
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple

try:
    from pypdf import PdfReader
except Exception:  # pragma: no cover - handled at runtime
    PdfReader = None


ROOT_DIR = Path(__file__).resolve().parent
DATA_DIR = ROOT_DIR / "data"
DEFAULT_DB_PATH = DATA_DIR / "tax_dispute_prototype.sqlite"
ENV_PATH = ROOT_DIR / ".env"

RSM_BLUE_HEX = "009CDE"
RSM_GREEN_HEX = "43A047"
RSM_GRAY_HEX = "54585A"
RSM_MID_GRAY_HEX = "8A8F93"
RSM_SOFT_GRAY_HEX = "F4F6F8"

OUTCOME_LABELS = {
    "WP_FULL_WIN": "WP dikabulkan seluruhnya",
    "WP_PARTIAL_WIN": "WP dikabulkan sebagian",
    "DJP_WIN": "DJP menang / banding ditolak",
    "FORMAL_REJECTED": "Tidak dapat diterima / gugur",
    "UNKNOWN": "Belum terklasifikasi",
}

OUTCOME_LABELS_EN = {
    "WP_FULL_WIN": "Taxpayer fully prevailed",
    "WP_PARTIAL_WIN": "Taxpayer partially prevailed",
    "DJP_WIN": "Tax authority prevailed / appeal rejected",
    "FORMAL_REJECTED": "Dismissed on formal grounds",
    "UNKNOWN": "Unclassified",
}

TAX_KEYWORDS = [
    ("PPN", ["pajak pertambahan nilai", " ppn ", "dpp ppn", "pajak masukan"]),
    ("PPh Badan", ["pajak penghasilan badan", "pph badan", "pajak penghasilan wajib pajak badan"]),
    ("PPh 21", ["pph pasal 21", "pajak penghasilan pasal 21", "pph 21"]),
    ("PPh 23", ["pph pasal 23", "pajak penghasilan pasal 23", "pph 23"]),
    ("PPh 26", ["pph pasal 26", "pajak penghasilan pasal 26", "pph 26"]),
    ("PBB", ["pajak bumi dan bangunan", " pbb "]),
]

ISSUE_KEYWORDS = [
    ("DPP_PPN", ["dpp ppn", "dasar pengenaan pajak", "penyerahan yang ppn", "peredaran usaha"]),
    ("PAJAK_MASUKAN", ["pajak masukan", "dikreditkan", "pengkreditan pajak masukan"]),
    ("FAKTUR_PAJAK", ["faktur pajak", "faktur pajak dobel", "faktur pengganti", "faktur masukan"]),
    ("PKPM_KONFIRMASI", ["pkpm", "konfirmasi lawan transaksi", "kpp penjual", "jawaban konfirmasi"]),
    ("SANKSI", ["sanksi administrasi", "pasal 13 ayat", "bunga", "kenaikan"]),
    ("FORMAL", ["tidak memenuhi ketentuan formal", "jangka waktu", "tidak dapat diterima"]),
    ("PPh", ["pajak penghasilan", "pph"]),
]

POSITIVE_KEYWORDS = [
    "faktur pajak dobel",
    "sudah dilaporkan",
    "spt masa ppn",
    "spt pembetulan",
    "bukti pembayaran",
    "dokumen pendukung",
    "memenuhi syarat formal",
    "memenuhi syarat material",
    "koreksi tidak dapat dipertahankan",
]

NEGATIVE_KEYWORDS = [
    "tidak dapat menunjukkan",
    "tidak dapat meyakini",
    "jawaban kpp penjual tidak ada",
    "tidak memberikan alasan",
    "bukti pendukung lainnya tidak ada",
    "tidak memenuhi ketentuan formal",
    "surat bantahan tidak dimasukkan",
    "koreksi tetap dipertahankan",
]

STOPWORDS = {
    "yang",
    "dan",
    "di",
    "ke",
    "dari",
    "dalam",
    "untuk",
    "dengan",
    "atas",
    "atau",
    "ini",
    "itu",
    "pada",
    "para",
    "bahwa",
    "adalah",
    "telah",
    "tidak",
    "oleh",
    "sebagai",
    "maka",
    "akan",
    "dapat",
    "karena",
    "dalam",
    "berdasarkan",
    "pemohon",
    "banding",
    "terbanding",
    "majelis",
}

RICH_EXTRACTION_FIELDS = [
    "procedure_type",
    "examination_level",
    "case_file_number",
    "decision_date",
    "tax_period",
    "wp_claim_amount",
    "djp_claim_amount",
    "representative_name",
    "appellee_name",
    "issue_subtype",
    "correction_object",
    "correction_reason",
    "wp_rebuttal_reason",
    "evidence_submitted",
    "accepted_arguments",
    "rejected_arguments",
    "sufficient_evidence_summary",
    "insufficient_evidence_summary",
    "verdict_text",
    "per_issue_outcome",
    "tax_before_amount",
    "tax_after_amount",
    "correction_reduction_amount",
    "sanctions_amount",
    "success_level",
]


@dataclass
class DocumentRecord:
    document_id: str
    filename: str
    file_path: str
    document_type: str
    putusan_number: str
    putusan_year: Optional[int]
    court_panel: str
    judge_names: str
    taxpayer_name: str
    taxpayer_npwp: str
    taxpayer_address: str
    legal_counsel_name: str
    legal_counsel_license: str
    djp_unit: str
    djp_decision_number: str
    skp_number: str
    tax_type: str
    issue_type: str
    issue_summary: str
    outcome: str
    amount_disputed: Optional[float]
    wp_position_summary: str
    djp_position_summary: str
    evidence_summary: str
    legal_references_summary: str
    court_reasoning_summary: str
    text: str
    extraction_status: str
    extracted_at: str


class DuplicateDocumentError(Exception):
    def __init__(self, message: str, duplicates: List[Dict[str, Any]]):
        super().__init__(message)
        self.duplicates = duplicates


class ExtractionError(Exception):
    def __init__(self, message: str, document_id: str = "", detail: str = ""):
        super().__init__(message)
        self.document_id = document_id
        self.detail = detail


def now_iso() -> str:
    return datetime.utcnow().isoformat(timespec="seconds") + "Z"


def ensure_data_dir() -> None:
    DATA_DIR.mkdir(exist_ok=True)


def load_local_env(env_path: Path = ENV_PATH) -> None:
    if not env_path.exists():
        return
    for raw_line in env_path.read_text(errors="ignore").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


def connect(db_path: Path = DEFAULT_DB_PATH) -> sqlite3.Connection:
    ensure_data_dir()
    conn = sqlite3.connect(str(db_path))
    conn.row_factory = sqlite3.Row
    return conn


def init_db(db_path: Path = DEFAULT_DB_PATH) -> None:
    with connect(db_path) as conn:
        conn.executescript(
            """
            CREATE TABLE IF NOT EXISTS documents (
                document_id TEXT PRIMARY KEY,
                filename TEXT NOT NULL,
                file_path TEXT NOT NULL UNIQUE,
                document_type TEXT,
                file_size INTEGER,
                page_count INTEGER,
                text TEXT,
                putusan_number TEXT,
                putusan_year INTEGER,
                court_panel TEXT,
                judge_names TEXT,
                taxpayer_name TEXT,
                taxpayer_npwp TEXT,
                taxpayer_address TEXT,
                legal_counsel_name TEXT,
                legal_counsel_license TEXT,
                djp_unit TEXT,
                djp_decision_number TEXT,
                skp_number TEXT,
                tax_type TEXT,
                issue_type TEXT,
                issue_summary TEXT,
                outcome TEXT,
                amount_disputed REAL,
                wp_position_summary TEXT,
                djp_position_summary TEXT,
                evidence_summary TEXT,
                legal_references_summary TEXT,
                court_reasoning_summary TEXT,
                extraction_status TEXT NOT NULL DEFAULT 'pending',
                extracted_at TEXT,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS chunks (
                chunk_id TEXT PRIMARY KEY,
                document_id TEXT NOT NULL,
                chunk_order INTEGER NOT NULL,
                section_type TEXT,
                text TEXT NOT NULL,
                token_count INTEGER,
                created_at TEXT NOT NULL,
                FOREIGN KEY(document_id) REFERENCES documents(document_id)
            );

            CREATE TABLE IF NOT EXISTS analysis_reports (
                report_id TEXT PRIMARY KEY,
                created_at TEXT NOT NULL,
                input_json TEXT NOT NULL,
                similar_cases_json TEXT NOT NULL,
                result_json TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS llm_labels (
                label_id TEXT PRIMARY KEY,
                document_id TEXT NOT NULL,
                model TEXT NOT NULL,
                label_json TEXT NOT NULL,
                tax_type TEXT,
                dispute_stage TEXT,
                issue_type TEXT,
                issue_subtype TEXT,
                outcome TEXT,
                issue_summary TEXT,
                wp_position_summary TEXT,
                djp_position_summary TEXT,
                court_reasoning_summary TEXT,
                evidence_summary TEXT,
                legal_references_summary TEXT,
                document_type TEXT,
                taxpayer_name TEXT,
                taxpayer_npwp TEXT,
                taxpayer_address TEXT,
                legal_counsel_name TEXT,
                legal_counsel_license TEXT,
                djp_unit TEXT,
                djp_decision_number TEXT,
                skp_number TEXT,
                court_panel TEXT,
                judge_names TEXT,
                confidence REAL,
                created_at TEXT NOT NULL,
                FOREIGN KEY(document_id) REFERENCES documents(document_id)
            );

            CREATE TABLE IF NOT EXISTS dashboard_metrics (
                metric_id TEXT PRIMARY KEY,
                metric_group TEXT NOT NULL,
                metric_key TEXT NOT NULL,
                metric_value REAL NOT NULL,
                metric_label TEXT,
                updated_at TEXT NOT NULL
            );

            CREATE TABLE IF NOT EXISTS document_extractions (
                extraction_id TEXT PRIMARY KEY,
                document_id TEXT NOT NULL UNIQUE,
                extraction_source TEXT,
                document_type TEXT,
                putusan_number TEXT,
                putusan_year INTEGER,
                court_panel TEXT,
                judge_names TEXT,
                procedure_type TEXT,
                examination_level TEXT,
                case_file_number TEXT,
                decision_date TEXT,
                tax_type TEXT,
                tax_period TEXT,
                skp_number TEXT,
                djp_decision_number TEXT,
                wp_claim_amount REAL,
                djp_claim_amount REAL,
                taxpayer_name TEXT,
                taxpayer_npwp TEXT,
                taxpayer_address TEXT,
                representative_name TEXT,
                legal_counsel_name TEXT,
                legal_counsel_license TEXT,
                appellee_name TEXT,
                djp_unit TEXT,
                issue_type TEXT,
                issue_subtype TEXT,
                amount_disputed REAL,
                correction_object TEXT,
                correction_reason TEXT,
                wp_rebuttal_reason TEXT,
                issue_summary TEXT,
                wp_position_summary TEXT,
                djp_position_summary TEXT,
                legal_references_summary TEXT,
                evidence_summary TEXT,
                evidence_submitted TEXT,
                court_reasoning_summary TEXT,
                accepted_arguments TEXT,
                rejected_arguments TEXT,
                sufficient_evidence_summary TEXT,
                insufficient_evidence_summary TEXT,
                verdict_text TEXT,
                per_issue_outcome TEXT,
                tax_before_amount REAL,
                tax_after_amount REAL,
                correction_reduction_amount REAL,
                sanctions_amount REAL,
                outcome TEXT,
                success_level TEXT,
                extraction_json TEXT NOT NULL,
                created_at TEXT NOT NULL,
                updated_at TEXT NOT NULL,
                FOREIGN KEY(document_id) REFERENCES documents(document_id)
            );

            CREATE INDEX IF NOT EXISTS idx_documents_tax_type ON documents(tax_type);
            CREATE INDEX IF NOT EXISTS idx_documents_issue_type ON documents(issue_type);
            CREATE INDEX IF NOT EXISTS idx_documents_outcome ON documents(outcome);
            CREATE INDEX IF NOT EXISTS idx_documents_year ON documents(putusan_year);
            CREATE INDEX IF NOT EXISTS idx_documents_filename ON documents(filename);
            CREATE INDEX IF NOT EXISTS idx_documents_putusan_number ON documents(putusan_number);
            CREATE INDEX IF NOT EXISTS idx_documents_taxpayer_name ON documents(taxpayer_name);
            CREATE INDEX IF NOT EXISTS idx_chunks_document ON chunks(document_id);
            CREATE INDEX IF NOT EXISTS idx_llm_labels_document ON llm_labels(document_id);
            CREATE INDEX IF NOT EXISTS idx_document_extractions_document ON document_extractions(document_id);
            """
        )
        ensure_column(conn, "documents", "label_source", "TEXT")
        ensure_column(conn, "documents", "llm_labeled_at", "TEXT")
        ensure_column(conn, "documents", "llm_confidence", "REAL")
        for column, definition in {
            "document_type": "TEXT",
            "judge_names": "TEXT",
            "taxpayer_name": "TEXT",
            "taxpayer_npwp": "TEXT",
            "taxpayer_address": "TEXT",
            "legal_counsel_name": "TEXT",
            "legal_counsel_license": "TEXT",
            "djp_unit": "TEXT",
            "djp_decision_number": "TEXT",
            "skp_number": "TEXT",
            "wp_position_summary": "TEXT",
            "djp_position_summary": "TEXT",
            "evidence_summary": "TEXT",
            "legal_references_summary": "TEXT",
            "procedure_type": "TEXT",
            "examination_level": "TEXT",
            "case_file_number": "TEXT",
            "decision_date": "TEXT",
            "tax_period": "TEXT",
            "wp_claim_amount": "REAL",
            "djp_claim_amount": "REAL",
            "representative_name": "TEXT",
            "appellee_name": "TEXT",
            "issue_subtype": "TEXT",
            "correction_object": "TEXT",
            "correction_reason": "TEXT",
            "wp_rebuttal_reason": "TEXT",
            "evidence_submitted": "TEXT",
            "accepted_arguments": "TEXT",
            "rejected_arguments": "TEXT",
            "sufficient_evidence_summary": "TEXT",
            "insufficient_evidence_summary": "TEXT",
            "verdict_text": "TEXT",
            "per_issue_outcome": "TEXT",
            "tax_before_amount": "REAL",
            "tax_after_amount": "REAL",
            "correction_reduction_amount": "REAL",
            "sanctions_amount": "REAL",
            "success_level": "TEXT",
        }.items():
            ensure_column(conn, "documents", column, definition)
        for column, definition in {
            "document_type": "TEXT",
            "taxpayer_name": "TEXT",
            "taxpayer_npwp": "TEXT",
            "taxpayer_address": "TEXT",
            "legal_counsel_name": "TEXT",
            "legal_counsel_license": "TEXT",
            "djp_unit": "TEXT",
            "djp_decision_number": "TEXT",
            "skp_number": "TEXT",
            "court_panel": "TEXT",
            "judge_names": "TEXT",
            "procedure_type": "TEXT",
            "examination_level": "TEXT",
            "case_file_number": "TEXT",
            "decision_date": "TEXT",
            "tax_period": "TEXT",
            "wp_claim_amount": "REAL",
            "djp_claim_amount": "REAL",
            "representative_name": "TEXT",
            "appellee_name": "TEXT",
            "issue_subtype": "TEXT",
            "correction_object": "TEXT",
            "correction_reason": "TEXT",
            "wp_rebuttal_reason": "TEXT",
            "evidence_submitted": "TEXT",
            "accepted_arguments": "TEXT",
            "rejected_arguments": "TEXT",
            "sufficient_evidence_summary": "TEXT",
            "insufficient_evidence_summary": "TEXT",
            "verdict_text": "TEXT",
            "per_issue_outcome": "TEXT",
            "tax_before_amount": "REAL",
            "tax_after_amount": "REAL",
            "correction_reduction_amount": "REAL",
            "sanctions_amount": "REAL",
            "success_level": "TEXT",
        }.items():
            ensure_column(conn, "llm_labels", column, definition)


def ensure_column(conn: sqlite3.Connection, table: str, column: str, definition: str) -> None:
    existing = {row["name"] for row in conn.execute(f"PRAGMA table_info({table})").fetchall()}
    if column not in existing:
        conn.execute(f"ALTER TABLE {table} ADD COLUMN {column} {definition}")


def table_exists(conn: sqlite3.Connection, table: str) -> bool:
    row = conn.execute("SELECT name FROM sqlite_master WHERE type = 'table' AND name = ?", (table,)).fetchone()
    return row is not None


def normalize_duplicate_key(value: Any, key_type: str = "") -> str:
    value = normalize_spaces(str(value or "")).lower()
    value = re.sub(r"\s+", " ", value)
    if key_type == "putusan_number":
        value = re.sub(r"[^a-z0-9]", "", value)
    return value


def is_meaningful_duplicate_key(value: Any, key_type: str) -> bool:
    value_norm = normalize_duplicate_key(value, key_type)
    if not value_norm:
        return False
    if key_type == "taxpayer_name" and value_norm in {"wajib pajak", "pemohon banding", "terbanding", "tergugat", "nomor", "nama", "unknown", "-"}:
        return False
    if key_type == "putusan_number" and value_norm in {"unknown", "-"}:
        return False
    return len(value_norm) >= 4


def duplicate_row_to_dict(row: sqlite3.Row, reasons: List[str]) -> Dict[str, Any]:
    return {
        "match_reason": ", ".join(reasons),
        "document_id": row["document_id"],
        "filename": row["filename"],
        "putusan_number": row["putusan_number"] or "",
        "taxpayer_name": row["taxpayer_name"] or "",
        "taxpayer_npwp": row["taxpayer_npwp"] or "",
        "tax_type": row["tax_type"] or "",
        "issue_type": row["issue_type"] or "",
        "outcome": row["outcome"] or "",
        "extracted_at": row["extracted_at"] or "",
        "file_path": row["file_path"] or "",
    }


def putusan_duplicate_match(input_key: str, existing_value: str) -> bool:
    existing_key = normalize_duplicate_key(existing_value, "putusan_number") if is_meaningful_duplicate_key(existing_value, "putusan_number") else ""
    if not input_key or not existing_key:
        return False
    if input_key == existing_key:
        return True
    shortest = min(len(input_key), len(existing_key))
    return shortest >= 10 and (input_key in existing_key or existing_key in input_key)


def find_duplicate_documents(
    filename: str = "",
    putusan_number: str = "",
    taxpayer_name: str = "",
    exclude_document_id: str = "",
    db_path: Path = DEFAULT_DB_PATH,
) -> List[Dict[str, Any]]:
    init_db(db_path)
    putusan_key = normalize_duplicate_key(putusan_number, "putusan_number") if is_meaningful_duplicate_key(putusan_number, "putusan_number") else ""

    if not putusan_key:
        return []

    with connect(db_path) as conn:
        rows = conn.execute("SELECT * FROM documents WHERE extraction_status = 'completed'").fetchall()

    duplicates = []
    seen = set()
    for row in rows:
        if exclude_document_id and row["document_id"] == exclude_document_id:
            continue
        reasons = []
        if putusan_duplicate_match(putusan_key, row["putusan_number"]):
            reasons.append("nomor putusan sama")
        if reasons and row["document_id"] not in seen:
            seen.add(row["document_id"])
            duplicates.append(duplicate_row_to_dict(row, reasons))
    return duplicates


def assert_no_duplicate_document(
    filename: str = "",
    putusan_number: str = "",
    taxpayer_name: str = "",
    exclude_document_id: str = "",
    db_path: Path = DEFAULT_DB_PATH,
) -> None:
    duplicates = find_duplicate_documents(putusan_number=putusan_number, exclude_document_id=exclude_document_id, db_path=db_path)
    if duplicates:
        raise DuplicateDocumentError("Ekstraksi ditolak karena dokumen duplikat terdeteksi.", duplicates)


def find_pdfs(pdf_dir: Path = ROOT_DIR) -> List[Path]:
    return sorted(pdf_dir.glob("*.pdf"), key=lambda p: p.name.lower())


def extract_pdf_text(pdf_path: Path) -> Tuple[str, int, str]:
    text = ""
    page_count = 0
    method = "pypdf"

    if PdfReader is not None:
        try:
            reader = PdfReader(str(pdf_path))
            page_count = len(reader.pages)
            parts = []
            for page in reader.pages:
                parts.append(page.extract_text() or "")
            text = "\n\n".join(parts)
        except Exception:
            text = ""

    if len(text.strip()) < 500:
        pdftotext = _which("pdftotext")
        if pdftotext:
            method = "pdftotext"
            with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
                out_path = tmp.name
            try:
                subprocess.run(
                    [pdftotext, "-layout", str(pdf_path), out_path],
                    check=False,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                )
                text = Path(out_path).read_text(errors="ignore")
            finally:
                try:
                    Path(out_path).unlink()
                except OSError:
                    pass

    return clean_text(text), page_count, method


def pdf_page_count(pdf_path: Path) -> int:
    if PdfReader is not None:
        try:
            return len(PdfReader(str(pdf_path)).pages)
        except Exception:
            pass

    pdfinfo = _which("pdfinfo")
    if not pdfinfo:
        return 0
    try:
        result = subprocess.run(
            [pdfinfo, str(pdf_path)],
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
        )
        match = re.search(r"^Pages:\s*(\d+)", result.stdout or "", re.MULTILINE)
        return int(match.group(1)) if match else 0
    except Exception:
        return 0


def select_pdf_pages_for_vision(page_count: int, max_pages: Optional[int] = None) -> List[int]:
    max_pages = max_pages or int(os.environ.get("TDP_VISION_MAX_PAGES", "14"))
    max_pages = max(1, max_pages)
    if page_count <= 0:
        return [1]
    if page_count <= max_pages:
        return list(range(1, page_count + 1))

    midpoint = max(1, page_count // 2)
    candidates = (
        list(range(1, min(page_count, 6) + 1))
        + list(range(max(1, midpoint - 1), min(page_count, midpoint + 1) + 1))
        + list(range(max(1, page_count - 4), page_count + 1))
    )
    pages = []
    for page in sorted(set(candidates)):
        if page not in pages:
            pages.append(page)
    return pages[:max_pages]


def render_pdf_pages_to_data_urls(pdf_path: Path, max_pages: Optional[int] = None) -> Tuple[List[str], List[int]]:
    pdftoppm = _which("pdftoppm")
    if not pdftoppm:
        raise RuntimeError("pdftoppm tidak ditemukan. Install Poppler agar PDF scan bisa dirender untuk ekstraksi vision.")
    try:
        from PIL import Image
    except Exception as exc:
        raise RuntimeError("Pillow tidak tersedia. Install paket pillow agar PDF scan bisa diproses.") from exc

    page_count = pdf_page_count(pdf_path)
    pages = select_pdf_pages_for_vision(page_count, max_pages=max_pages)
    dpi = int(os.environ.get("TDP_VISION_DPI", "130"))
    max_side = int(os.environ.get("TDP_VISION_MAX_SIDE", "1800"))
    quality = int(os.environ.get("TDP_VISION_JPEG_QUALITY", "72"))

    data_urls = []
    rendered_pages = []
    with tempfile.TemporaryDirectory() as tmpdir:
        tmp_path = Path(tmpdir)
        for page in pages:
            prefix = tmp_path / f"page_{page}"
            result = subprocess.run(
                [pdftoppm, "-f", str(page), "-l", str(page), "-r", str(dpi), "-png", str(pdf_path), str(prefix)],
                check=False,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
            )
            if result.returncode != 0:
                raise RuntimeError(f"Gagal render halaman {page}: {(result.stderr or '').strip()[:300]}")
            candidates = sorted(tmp_path.glob(f"page_{page}*.png"))
            if not candidates:
                continue
            with Image.open(candidates[0]) as image:
                image = image.convert("RGB")
                image.thumbnail((max_side, max_side))
                buffer = io.BytesIO()
                image.save(buffer, format="JPEG", quality=quality, optimize=True)
            encoded = base64.b64encode(buffer.getvalue()).decode("ascii")
            data_urls.append(f"data:image/jpeg;base64,{encoded}")
            rendered_pages.append(page)
    if not data_urls:
        raise RuntimeError("Tidak ada halaman PDF yang berhasil dirender untuk ekstraksi vision.")
    return data_urls, rendered_pages


def _which(name: str) -> Optional[str]:
    for part in os.environ.get("PATH", "").split(os.pathsep):
        candidate = Path(part) / name
        if candidate.exists() and os.access(candidate, os.X_OK):
            return str(candidate)
    return None


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"Halaman\s+\d+\s+dari\s+\d+", " ", text, flags=re.I)
    return text.strip()


def extract_metadata(pdf_path: Path, text: str) -> Dict[str, Any]:
    sample = f"{pdf_path.stem}\n{text[:50000]}"
    lower = f" {sample.lower()} "
    putusan_number = extract_putusan_number(sample, pdf_path.stem)
    year = extract_year(putusan_number) or extract_year(sample)
    panel = extract_panel(putusan_number or pdf_path.stem)
    tax_type = classify_by_keywords(lower, TAX_KEYWORDS, default="UNKNOWN")
    issue_scope = extract_section_snippet(
        text,
        ["pokok sengketa", "menurut pemohon banding", "menurut terbanding"],
        fallback_keywords=["koreksi", "sengketa", "pajak masukan", "dpp"],
        limit=3500,
    )
    issue_type = classify_issue_type(issue_scope, lower, tax_type)
    outcome = classify_outcome(sample)
    amount = extract_amount(sample)
    issue_summary = normalize_spaces(issue_scope[:900])
    court_reasoning = extract_section_snippet(
        text,
        ["pendapat majelis", "menurut pendapat majelis", "menimbang"],
        fallback_keywords=["majelis berpendapat", "mengadili", "koreksi"],
    )
    parties = extract_party_metadata(text)
    rich = extract_rich_dispute_metadata(text, sample, tax_type, issue_type, outcome)

    return {
        "document_type": infer_document_type(text, pdf_path.name),
        "putusan_number": putusan_number,
        "putusan_year": year,
        "court_panel": panel,
        "judge_names": parties.get("judge_names"),
        "taxpayer_name": parties.get("taxpayer_name"),
        "taxpayer_npwp": parties.get("taxpayer_npwp"),
        "taxpayer_address": parties.get("taxpayer_address"),
        "representative_name": parties.get("representative_name"),
        "legal_counsel_name": parties.get("legal_counsel_name"),
        "legal_counsel_license": parties.get("legal_counsel_license"),
        "appellee_name": parties.get("appellee_name"),
        "djp_unit": parties.get("djp_unit"),
        "djp_decision_number": parties.get("djp_decision_number"),
        "skp_number": parties.get("skp_number"),
        "tax_type": tax_type,
        "issue_type": issue_type,
        "issue_subtype": rich.get("issue_subtype"),
        "issue_summary": issue_summary,
        "outcome": outcome,
        "amount_disputed": amount,
        "wp_position_summary": "",
        "djp_position_summary": "",
        "evidence_summary": "",
        "legal_references_summary": "",
        "court_reasoning_summary": court_reasoning,
        **rich,
    }


def infer_document_type(text: str, filename: str = "") -> str:
    lower = f"{filename}\n{text[:12000]}".lower()
    if "mengadili" in lower and "pengadilan pajak" in lower and "putusan" in lower:
        return "putusan_pengadilan"
    if "surat banding" in lower or "permohonan banding" in lower:
        return "surat_banding"
    if "surat keberatan" in lower or "permohonan keberatan" in lower:
        return "surat_keberatan"
    if "surat pemberitahuan hasil pemeriksaan" in lower or "sphp" in lower:
        return "jawaban_sphp"
    if "uraian banding" in lower:
        return "uraian_banding"
    if "surat bantahan" in lower:
        return "bantahan"
    return "dokumen_pendukung"


def extract_party_metadata(text: str) -> Dict[str, str]:
    sample = text[:50000]
    parties = {
        "taxpayer_name": regex_after(sample, [r"Pemohon Banding\s*[:\-]?\s*([^\n;]{3,180})", r"Nama\s*[:\-]\s*([^\n;]{3,180})"]),
        "taxpayer_npwp": first_match(sample, r"\b\d{2}\.\d{3}\.\d{3}\.\d[-]\d{3}\.\d{3}\b"),
        "taxpayer_address": regex_after(sample, [r"beralamat di\s*([^\n;]{8,240})", r"Alamat\s*[:\-]\s*([^\n;]{8,240})"]),
        "representative_name": regex_after(sample, [r"Wakil\s*[:\-]?\s*([^\n;]{3,180})", r"diwakili oleh\s*([^\n;]{3,180})"]),
        "legal_counsel_name": regex_after(sample, [r"Kuasa Hukum\s*[:\-]?\s*([^\n;]{3,180})", r"kuasa\s+Pemohon Banding\s*[:\-]?\s*([^\n;]{3,180})"]),
        "legal_counsel_license": first_match(sample, r"(?:KEP|SI|KIP)[-\s]?\d+[^\s,;]{0,60}"),
        "appellee_name": regex_after(sample, [r"Terbanding\s*[:\-]?\s*([^\n;]{3,180})", r"Tergugat\s*[:\-]?\s*([^\n;]{3,180})"]),
        "djp_unit": regex_after(sample, [r"(Kantor Pelayanan Pajak[^\n;]{3,180})", r"(Kantor Wilayah DJP[^\n;]{3,180})"]),
        "djp_decision_number": first_match(sample, r"KEP[-\s]?\d+[^\s,;]{0,80}"),
        "skp_number": first_match(sample, r"\b\d{5}/\d{3}/\d{2}/\d{3}/\d{2}\b"),
        "judge_names": regex_after(sample, [r"Majelis\s+Hakim\s*[:\-]?\s*([^\n]{5,260})", r"Hakim\s+Ketua\s*[:\-]?\s*([^\n]{5,180})"]),
    }
    return {key: clean_extracted_party_value(value, key) for key, value in parties.items()}


def clean_extracted_party_value(value: str, key: str) -> str:
    value = normalize_spaces(value)
    generic = {"nomor", "nama", "alamat", "npwp", "terbanding", "pemohon banding", "wajib pajak"}
    if normalize_duplicate_key(value) in generic:
        return ""
    if key in {"taxpayer_name", "appellee_name"}:
        value = re.sub(r"^(Nomor|Nama|NPWP|Alamat)\s*[:\-]?\s*", "", value, flags=re.I)
        value = re.split(r"\s{2,}| NPWP | Alamat | Nomor ", value, maxsplit=1, flags=re.I)[0]
    return value[:240]


def extract_rich_dispute_metadata(text: str, sample: str, tax_type: str, issue_type: str, outcome: str) -> Dict[str, Any]:
    lower = sample.lower()
    wp_position = extract_section_snippet(
        text,
        ["menurut pemohon banding", "pemohon banding menyatakan", "alasan banding", "surat bantahan"],
        ["pemohon banding", "wajib pajak", "bantahan"],
        limit=1400,
    )
    djp_position = extract_section_snippet(
        text,
        ["menurut terbanding", "terbanding menyatakan", "uraian banding", "alasan koreksi"],
        ["terbanding", "koreksi", "pemeriksa"],
        limit=1400,
    )
    verdict_text = extract_section_snippet(
        text,
        ["mengadili", "memutuskan", "amar putusan"],
        ["mengabulkan", "menolak", "tidak dapat diterima"],
        limit=1600,
    )
    evidence = extract_section_snippet(
        text,
        ["bukti", "dokumen pendukung", "alat bukti"],
        ["faktur pajak", "spt masa", "invoice", "bukti pembayaran"],
        limit=1200,
    )
    legal_refs = extract_section_snippet(
        text,
        ["dasar hukum", "menurut ketentuan", "berdasarkan pasal"],
        ["undang-undang", "peraturan", "pasal"],
        limit=1300,
    )

    return {
        "procedure_type": extract_procedure_type(sample),
        "examination_level": extract_examination_level(sample),
        "case_file_number": regex_after(sample, [r"Nomor\s+Berkas\s*[:\-]?\s*([^\n;]{3,120})", r"No\.\s*Berkas\s*[:\-]?\s*([^\n;]{3,120})"]),
        "decision_date": extract_decision_date(sample),
        "tax_period": extract_tax_period(sample),
        "wp_claim_amount": extract_amount_near(text, ["menurut pemohon", "menurut wajib pajak", "nilai menurut wp"]),
        "djp_claim_amount": extract_amount_near(text, ["menurut terbanding", "menurut djp", "nilai menurut djp", "koreksi"]),
        "issue_subtype": infer_issue_subtype(text, issue_type),
        "correction_object": extract_correction_object(text, issue_type),
        "correction_reason": normalize_spaces(djp_position[:900]),
        "wp_rebuttal_reason": normalize_spaces(wp_position[:900]),
        "wp_position_summary": normalize_spaces(wp_position[:900]),
        "djp_position_summary": normalize_spaces(djp_position[:900]),
        "evidence_summary": normalize_spaces(evidence[:900]),
        "evidence_submitted": extract_evidence_list(text),
        "legal_references_summary": normalize_spaces(legal_refs[:900]),
        "accepted_arguments": extract_section_snippet(text, ["dapat diterima", "terbukti", "meyakini majelis"], ["terbukti", "meyakini"], limit=700),
        "rejected_arguments": extract_section_snippet(text, ["tidak dapat diterima", "tidak terbukti", "tidak meyakini"], ["tidak terbukti", "tidak meyakini"], limit=700),
        "sufficient_evidence_summary": extract_section_snippet(text, ["bukti cukup", "dapat membuktikan", "terbukti"], ["bukti", "terbukti"], limit=700),
        "insufficient_evidence_summary": extract_section_snippet(text, ["bukti tidak cukup", "tidak dapat membuktikan", "tidak dapat menunjukkan"], ["tidak dapat menunjukkan", "tidak cukup"], limit=700),
        "verdict_text": normalize_spaces(verdict_text[:1200]),
        "per_issue_outcome": summarize_per_issue_outcome(verdict_text, outcome),
        "tax_before_amount": extract_amount_near(text[-25000:], ["semula", "menurut terbanding", "jumlah pajak"]),
        "tax_after_amount": extract_amount_near(text[-25000:], ["menjadi", "menurut majelis", "ditetapkan"]),
        "correction_reduction_amount": extract_amount_near(text[-25000:], ["pengurangan", "koreksi dikurangi", "dikabulkan sebagian"]),
        "sanctions_amount": extract_amount_near(text, ["sanksi", "bunga", "kenaikan", "denda"]),
        "success_level": outcome_to_success_level(outcome),
    }


def extract_procedure_type(text: str) -> str:
    lower = text.lower()
    if "gugatan" in lower:
        return "Gugatan"
    if "banding" in lower:
        return "Banding"
    return regex_after(text, [r"Jenis\s+Acara\s*[:\-]?\s*([^\n;]{3,120})"]) or ""


def extract_examination_level(text: str) -> str:
    lower = text.lower()
    if "pengadilan pajak" in lower:
        return "Pengadilan Pajak"
    return regex_after(text, [r"Tingkat\s+Pemeriksaan\s*[:\-]?\s*([^\n;]{3,120})"]) or ""


def extract_decision_date(text: str) -> str:
    patterns = [
        r"tanggal\s+putusan\s*[:\-]?\s*([0-9]{1,2}\s+[A-Za-z]+\s+\d{4})",
        r"diputus[^\\n]{0,120}?pada\s+tanggal\s+([0-9]{1,2}\s+[A-Za-z]+\s+\d{4})",
        r"diucapkan[^\\n]{0,120}?pada\s+tanggal\s+([0-9]{1,2}\s+[A-Za-z]+\s+\d{4})",
    ]
    return regex_after(text, patterns)


def extract_tax_period(text: str) -> str:
    patterns = [
        r"Masa\s+Pajak\s*[:\-]?\s*([A-Za-z0-9\s./,-]{4,120})",
        r"Tahun\s+Pajak\s*[:\-]?\s*((?:19|20)\d{2})",
        r"Masa/Tahun\s+Pajak\s*[:\-]?\s*([A-Za-z0-9\s./,-]{4,120})",
    ]
    return regex_after(text[:50000], patterns)


def extract_amount_near(text: str, anchors: Sequence[str]) -> Optional[float]:
    lower = text.lower()
    for anchor in anchors:
        pos = lower.find(anchor.lower())
        if pos >= 0:
            snippet = text[pos : pos + 2500]
            amount = extract_amount(snippet)
            if amount:
                return amount
    return None


def infer_issue_subtype(text: str, issue_type: str) -> str:
    lower = text.lower()
    if "pajak masukan" in lower and "konfirmasi" in lower:
        return "Pajak Masukan - Konfirmasi lawan transaksi"
    if "faktur pajak" in lower and "pengganti" in lower:
        return "Faktur Pajak Pengganti"
    if "dasar pengenaan pajak" in lower or "dpp" in lower:
        return "DPP / Dasar Pengenaan Pajak"
    if "sanksi" in lower:
        return "Sanksi administrasi"
    return issue_type.replace("_", " ").title() if issue_type else ""


def extract_correction_object(text: str, issue_type: str) -> str:
    snippet = extract_section_snippet(
        text,
        ["koreksi", "pokok sengketa", "objek sengketa"],
        ["pajak masukan", "dpp", "faktur pajak", "sanksi"],
        limit=700,
    )
    if snippet:
        return snippet[:700]
    return issue_type.replace("_", " ")


def extract_evidence_list(text: str) -> str:
    lower = text.lower()
    evidence = []
    mapping = [
        ("Faktur Pajak", "faktur pajak"),
        ("SPT Masa PPN", "spt masa"),
        ("Invoice", "invoice"),
        ("Bukti Pembayaran", "bukti pembayaran"),
        ("Rekening Koran", "rekening koran"),
        ("Kontrak", "kontrak"),
        ("Purchase Order", "purchase order"),
        ("Delivery Order/Surat Jalan", "delivery order"),
        ("Konfirmasi Lawan Transaksi", "konfirmasi"),
        ("Bukti Potong/Pungut", "bukti potong"),
    ]
    for label, keyword in mapping:
        if keyword in lower:
            evidence.append(label)
    return "; ".join(dedupe(evidence))


def summarize_per_issue_outcome(verdict_text: str, outcome: str) -> str:
    verdict = normalize_spaces(verdict_text)
    if verdict:
        return verdict[:600]
    return OUTCOME_LABELS.get(outcome, outcome or "UNKNOWN")


def outcome_to_success_level(outcome: str) -> str:
    return {
        "WP_FULL_WIN": "WP full win",
        "WP_PARTIAL_WIN": "WP partial win",
        "DJP_WIN": "DJP win",
        "FORMAL_REJECTED": "Formal rejected",
        "DROPPED": "Dropped",
    }.get(outcome or "UNKNOWN", "Unknown")


def first_match(text: str, pattern: str) -> str:
    match = re.search(pattern, text, flags=re.I)
    return normalize_spaces(match.group(0)) if match else ""


def regex_after(text: str, patterns: Sequence[str]) -> str:
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.I)
        if match:
            return normalize_spaces(match.group(1))
    return ""


def extract_putusan_number(text: str, fallback: str) -> str:
    patterns = [
        r"PUT[-\s]?\d{3,6}[./]\d{1,3}[./]\d{4}[./]PP[./\s]M[.\s]?[A-Z0-9IVX]+[A-Z]?\s+Tahun\s+\d{4}",
        r"PUT[-\s]?\d{3,6}[./]\d{1,3}[./]\d{4}[./]PP[./\s]M[.\s]?[A-Z0-9IVX]+[A-Z]?",
        r"PUT[-\s]?\d{3,6}[./]\d{1,3}[./]\d{4}[./]PP",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.I)
        if match:
            return normalize_spaces(match.group(0))
    return fallback.replace("w", "").strip()


def extract_year(text: str) -> Optional[int]:
    if not text:
        return None
    matches = re.findall(r"Tahun\s+((?:19|20)\d{2})", text, flags=re.I)
    if matches:
        return int(matches[-1])
    matches = re.findall(r"\b((?:19|20)\d{2})\b", text)
    if matches:
        return int(matches[-1])
    return None


def extract_panel(text: str) -> str:
    match = re.search(r"PP[./\s]M[.\s]?([A-Z0-9IVX]+[A-Z]?)", text, flags=re.I)
    if match:
        return f"M.{match.group(1).replace(' ', '').upper()}"
    return "UNKNOWN"


def classify_by_keywords(lower_text: str, groups: Sequence[Tuple[str, Sequence[str]]], default: str) -> str:
    best_label = default
    best_score = 0
    for label, keywords in groups:
        score = sum(lower_text.count(keyword.lower()) for keyword in keywords)
        if score > best_score:
            best_score = score
            best_label = label
    return best_label


def classify_issue_type(issue_scope: str, lower_text: str, tax_type: str) -> str:
    scope = f" {issue_scope.lower()} "
    scores: Dict[str, int] = {}
    for label, keywords in ISSUE_KEYWORDS:
        scores[label] = sum(scope.count(keyword.lower()) * 3 for keyword in keywords)
        scores[label] += sum(lower_text.count(keyword.lower()) for keyword in keywords[:2])

    if tax_type == "PPN":
        ppn_labels = ["PAJAK_MASUKAN", "DPP_PPN", "FAKTUR_PAJAK", "PKPM_KONFIRMASI", "SANKSI", "FORMAL"]
        best = max(ppn_labels, key=lambda label: scores.get(label, 0))
        if scores.get(best, 0) > 0:
            return best
        return "UNKNOWN"

    best_label = max(scores, key=lambda label: scores.get(label, 0))
    if scores.get(best_label, 0) <= 0:
        return "UNKNOWN"
    return best_label


def classify_outcome(text: str) -> str:
    lower = text.lower()
    tail = lower[-35000:]
    if "mengabulkan seluruhnya" in tail:
        return "WP_FULL_WIN"
    if "mengabulkan sebagian" in tail:
        return "WP_PARTIAL_WIN"
    if "tidak dapat diterima" in tail:
        return "FORMAL_REJECTED"
    if re.search(r"\bmenolak\b", tail) or "menolak banding" in tail:
        return "DJP_WIN"
    if "gugur" in tail:
        return "FORMAL_REJECTED"
    if "mengabulkan sebagian" in lower:
        return "WP_PARTIAL_WIN"
    if "mengabulkan seluruhnya" in lower:
        return "WP_FULL_WIN"
    return "UNKNOWN"


def extract_amount(text: str) -> Optional[float]:
    amounts = []
    for match in re.finditer(r"Rp\.?\s*([0-9][0-9.\s]{3,}(?:,\d{1,2})?)", text):
        value = parse_rupiah(match.group(1))
        if value and value > 100000:
            amounts.append(value)
    if not amounts:
        return None
    amounts = sorted(amounts)
    index = max(0, int(len(amounts) * 0.85) - 1)
    return amounts[index]


def parse_rupiah(value: str) -> Optional[float]:
    cleaned = value.replace(" ", "").replace(".", "")
    cleaned = cleaned.replace(",", ".")
    try:
        return float(cleaned)
    except ValueError:
        return None


def extract_section_snippet(text: str, anchors: Sequence[str], fallback_keywords: Sequence[str], limit: int = 900) -> str:
    lower = text.lower()
    for anchor in anchors:
        pos = lower.find(anchor.lower())
        if pos >= 0:
            return normalize_spaces(text[pos : pos + limit])
    for keyword in fallback_keywords:
        pos = lower.find(keyword.lower())
        if pos >= 0:
            start = max(0, pos - 160)
            return normalize_spaces(text[start : pos + limit])
    return normalize_spaces(text[:limit])


def normalize_spaces(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip()


def chunk_text(text: str, max_chars: int = 1800, overlap: int = 180) -> List[str]:
    if not text:
        return []
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks = []
    current = ""
    for paragraph in paragraphs:
        if len(current) + len(paragraph) + 2 <= max_chars:
            current = f"{current}\n\n{paragraph}".strip()
        else:
            if current:
                chunks.append(current)
            prefix = current[-overlap:] if current and overlap else ""
            current = f"{prefix}\n\n{paragraph}".strip()
    if current:
        chunks.append(current)
    if not chunks:
        for i in range(0, len(text), max_chars - overlap):
            chunks.append(text[i : i + max_chars])
    return chunks[:180]


def tokenize(text: str) -> List[str]:
    words = re.findall(r"[a-zA-Z0-9_]{3,}", (text or "").lower())
    return [w for w in words if w not in STOPWORDS and not w.isdigit()]


def row_to_document(row: sqlite3.Row) -> DocumentRecord:
    return DocumentRecord(
        document_id=row["document_id"],
        filename=row["filename"],
        file_path=row["file_path"],
        document_type=row["document_type"] or "dokumen_pendukung",
        putusan_number=row["putusan_number"] or "",
        putusan_year=row["putusan_year"],
        court_panel=row["court_panel"] or "UNKNOWN",
        judge_names=row["judge_names"] or "",
        taxpayer_name=row["taxpayer_name"] or "",
        taxpayer_npwp=row["taxpayer_npwp"] or "",
        taxpayer_address=row["taxpayer_address"] or "",
        legal_counsel_name=row["legal_counsel_name"] or "",
        legal_counsel_license=row["legal_counsel_license"] or "",
        djp_unit=row["djp_unit"] or "",
        djp_decision_number=row["djp_decision_number"] or "",
        skp_number=row["skp_number"] or "",
        tax_type=row["tax_type"] or "UNKNOWN",
        issue_type=row["issue_type"] or "UNKNOWN",
        issue_summary=row["issue_summary"] or "",
        outcome=row["outcome"] or "UNKNOWN",
        amount_disputed=row["amount_disputed"],
        wp_position_summary=row["wp_position_summary"] or "",
        djp_position_summary=row["djp_position_summary"] or "",
        evidence_summary=row["evidence_summary"] or "",
        legal_references_summary=row["legal_references_summary"] or "",
        court_reasoning_summary=row["court_reasoning_summary"] or "",
        text=row["text"] or "",
        extraction_status=row["extraction_status"] or "unknown",
        extracted_at=row["extracted_at"] or "",
    )


def build_document_extraction_payload(meta: Dict[str, Any], text: str) -> Dict[str, Any]:
    payload = {
        "metadata_putusan": {
            "document_type": meta.get("document_type"),
            "putusan_number": meta.get("putusan_number"),
            "putusan_year": meta.get("putusan_year"),
            "court_panel": meta.get("court_panel"),
            "judge_names": meta.get("judge_names"),
            "procedure_type": meta.get("procedure_type"),
            "examination_level": meta.get("examination_level"),
            "case_file_number": meta.get("case_file_number"),
            "decision_date": meta.get("decision_date"),
        },
        "objek_sengketa": {
            "tax_type": meta.get("tax_type"),
            "tax_period": meta.get("tax_period"),
            "skp_number": meta.get("skp_number"),
            "djp_decision_number": meta.get("djp_decision_number"),
            "wp_claim_amount": meta.get("wp_claim_amount"),
            "djp_claim_amount": meta.get("djp_claim_amount"),
        },
        "pihak": {
            "taxpayer_name": meta.get("taxpayer_name"),
            "taxpayer_npwp": meta.get("taxpayer_npwp"),
            "taxpayer_address": meta.get("taxpayer_address"),
            "representative_name": meta.get("representative_name"),
            "legal_counsel_name": meta.get("legal_counsel_name"),
            "legal_counsel_license": meta.get("legal_counsel_license"),
            "appellee_name": meta.get("appellee_name"),
            "djp_unit": meta.get("djp_unit"),
        },
        "pokok_sengketa": {
            "issue_type": meta.get("issue_type"),
            "issue_subtype": meta.get("issue_subtype"),
            "amount_disputed": meta.get("amount_disputed"),
            "correction_object": meta.get("correction_object"),
            "correction_reason": meta.get("correction_reason"),
            "wp_rebuttal_reason": meta.get("wp_rebuttal_reason"),
            "issue_summary": meta.get("issue_summary"),
        },
        "argumen": {
            "wp_position_summary": meta.get("wp_position_summary"),
            "djp_position_summary": meta.get("djp_position_summary"),
            "legal_references_summary": meta.get("legal_references_summary"),
            "evidence_summary": meta.get("evidence_summary"),
            "evidence_submitted": meta.get("evidence_submitted"),
        },
        "pertimbangan": {
            "court_reasoning_summary": meta.get("court_reasoning_summary"),
            "accepted_arguments": meta.get("accepted_arguments"),
            "rejected_arguments": meta.get("rejected_arguments"),
            "sufficient_evidence_summary": meta.get("sufficient_evidence_summary"),
            "insufficient_evidence_summary": meta.get("insufficient_evidence_summary"),
        },
        "outcome": {
            "verdict_text": meta.get("verdict_text"),
            "per_issue_outcome": meta.get("per_issue_outcome"),
            "tax_before_amount": meta.get("tax_before_amount"),
            "tax_after_amount": meta.get("tax_after_amount"),
            "correction_reduction_amount": meta.get("correction_reduction_amount"),
            "sanctions_amount": meta.get("sanctions_amount"),
            "outcome": meta.get("outcome"),
            "success_level": meta.get("success_level"),
        },
    }
    payload["source_text_chars"] = len(text or "")
    return payload


def upsert_document_extraction(
    conn: sqlite3.Connection,
    document_id: str,
    meta: Dict[str, Any],
    extraction_source: str,
    timestamp: str,
    text: str = "",
) -> None:
    payload = build_document_extraction_payload(meta, text)
    existing = conn.execute("SELECT extraction_id, created_at FROM document_extractions WHERE document_id = ?", (document_id,)).fetchone()
    extraction_id = existing["extraction_id"] if existing else str(uuid.uuid4())
    created_at = existing["created_at"] if existing else timestamp
    conn.execute(
        """
        INSERT INTO document_extractions (
            extraction_id, document_id, extraction_source, document_type, putusan_number, putusan_year,
            court_panel, judge_names, procedure_type, examination_level, case_file_number, decision_date,
            tax_type, tax_period, skp_number, djp_decision_number, wp_claim_amount, djp_claim_amount,
            taxpayer_name, taxpayer_npwp, taxpayer_address, representative_name, legal_counsel_name,
            legal_counsel_license, appellee_name, djp_unit, issue_type, issue_subtype, amount_disputed,
            correction_object, correction_reason, wp_rebuttal_reason, issue_summary, wp_position_summary,
            djp_position_summary, legal_references_summary, evidence_summary, evidence_submitted,
            court_reasoning_summary, accepted_arguments, rejected_arguments, sufficient_evidence_summary,
            insufficient_evidence_summary, verdict_text, per_issue_outcome, tax_before_amount,
            tax_after_amount, correction_reduction_amount, sanctions_amount, outcome, success_level,
            extraction_json, created_at, updated_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        ON CONFLICT(document_id) DO UPDATE SET
            extraction_source=excluded.extraction_source,
            document_type=excluded.document_type,
            putusan_number=excluded.putusan_number,
            putusan_year=excluded.putusan_year,
            court_panel=excluded.court_panel,
            judge_names=excluded.judge_names,
            procedure_type=excluded.procedure_type,
            examination_level=excluded.examination_level,
            case_file_number=excluded.case_file_number,
            decision_date=excluded.decision_date,
            tax_type=excluded.tax_type,
            tax_period=excluded.tax_period,
            skp_number=excluded.skp_number,
            djp_decision_number=excluded.djp_decision_number,
            wp_claim_amount=excluded.wp_claim_amount,
            djp_claim_amount=excluded.djp_claim_amount,
            taxpayer_name=excluded.taxpayer_name,
            taxpayer_npwp=excluded.taxpayer_npwp,
            taxpayer_address=excluded.taxpayer_address,
            representative_name=excluded.representative_name,
            legal_counsel_name=excluded.legal_counsel_name,
            legal_counsel_license=excluded.legal_counsel_license,
            appellee_name=excluded.appellee_name,
            djp_unit=excluded.djp_unit,
            issue_type=excluded.issue_type,
            issue_subtype=excluded.issue_subtype,
            amount_disputed=excluded.amount_disputed,
            correction_object=excluded.correction_object,
            correction_reason=excluded.correction_reason,
            wp_rebuttal_reason=excluded.wp_rebuttal_reason,
            issue_summary=excluded.issue_summary,
            wp_position_summary=excluded.wp_position_summary,
            djp_position_summary=excluded.djp_position_summary,
            legal_references_summary=excluded.legal_references_summary,
            evidence_summary=excluded.evidence_summary,
            evidence_submitted=excluded.evidence_submitted,
            court_reasoning_summary=excluded.court_reasoning_summary,
            accepted_arguments=excluded.accepted_arguments,
            rejected_arguments=excluded.rejected_arguments,
            sufficient_evidence_summary=excluded.sufficient_evidence_summary,
            insufficient_evidence_summary=excluded.insufficient_evidence_summary,
            verdict_text=excluded.verdict_text,
            per_issue_outcome=excluded.per_issue_outcome,
            tax_before_amount=excluded.tax_before_amount,
            tax_after_amount=excluded.tax_after_amount,
            correction_reduction_amount=excluded.correction_reduction_amount,
            sanctions_amount=excluded.sanctions_amount,
            outcome=excluded.outcome,
            success_level=excluded.success_level,
            extraction_json=excluded.extraction_json,
            updated_at=excluded.updated_at
        """,
        (
            extraction_id,
            document_id,
            extraction_source,
            meta.get("document_type"),
            meta.get("putusan_number"),
            meta.get("putusan_year"),
            meta.get("court_panel"),
            meta.get("judge_names"),
            meta.get("procedure_type"),
            meta.get("examination_level"),
            meta.get("case_file_number"),
            meta.get("decision_date"),
            meta.get("tax_type"),
            meta.get("tax_period"),
            meta.get("skp_number"),
            meta.get("djp_decision_number"),
            meta.get("wp_claim_amount"),
            meta.get("djp_claim_amount"),
            meta.get("taxpayer_name"),
            meta.get("taxpayer_npwp"),
            meta.get("taxpayer_address"),
            meta.get("representative_name"),
            meta.get("legal_counsel_name"),
            meta.get("legal_counsel_license"),
            meta.get("appellee_name"),
            meta.get("djp_unit"),
            meta.get("issue_type"),
            meta.get("issue_subtype"),
            meta.get("amount_disputed"),
            meta.get("correction_object"),
            meta.get("correction_reason"),
            meta.get("wp_rebuttal_reason"),
            meta.get("issue_summary"),
            meta.get("wp_position_summary"),
            meta.get("djp_position_summary"),
            meta.get("legal_references_summary"),
            meta.get("evidence_summary"),
            meta.get("evidence_submitted"),
            meta.get("court_reasoning_summary"),
            meta.get("accepted_arguments"),
            meta.get("rejected_arguments"),
            meta.get("sufficient_evidence_summary"),
            meta.get("insufficient_evidence_summary"),
            meta.get("verdict_text"),
            meta.get("per_issue_outcome"),
            meta.get("tax_before_amount"),
            meta.get("tax_after_amount"),
            meta.get("correction_reduction_amount"),
            meta.get("sanctions_amount"),
            meta.get("outcome"),
            meta.get("success_level"),
            json.dumps(payload, ensure_ascii=False),
            created_at,
            timestamp,
        ),
    )


def update_document_rich_fields(conn: sqlite3.Connection, document_id: str, meta: Dict[str, Any], timestamp: str) -> None:
    assignments = ", ".join(f"{field} = COALESCE(NULLIF(?, ''), {field})" for field in RICH_EXTRACTION_FIELDS)
    values = [meta.get(field) for field in RICH_EXTRACTION_FIELDS]
    conn.execute(
        f"UPDATE documents SET {assignments}, updated_at = ? WHERE document_id = ?",
        (*values, timestamp, document_id),
    )


def delete_document_records(document_id: str, db_path: Path = DEFAULT_DB_PATH) -> None:
    init_db(db_path)
    with connect(db_path) as conn:
        conn.execute("DELETE FROM chunks WHERE document_id = ?", (document_id,))
        conn.execute("DELETE FROM document_extractions WHERE document_id = ?", (document_id,))
        conn.execute("DELETE FROM llm_labels WHERE document_id = ?", (document_id,))
        conn.execute("DELETE FROM documents WHERE document_id = ?", (document_id,))


def upsert_document(
    pdf_path: Path,
    overwrite: bool = False,
    db_path: Path = DEFAULT_DB_PATH,
    document_type: Optional[str] = None,
    use_llm_extraction: bool = False,
) -> DocumentRecord:
    init_db(db_path)
    abs_path = str(pdf_path.resolve())
    with connect(db_path) as conn:
        existing = conn.execute("SELECT * FROM documents WHERE file_path = ?", (abs_path,)).fetchone()
        exclude_document_id = existing["document_id"] if existing and overwrite else ""

    if existing and not overwrite:
        reason = "nomor putusan sama" if existing["putusan_number"] else "dokumen sudah ada"
        raise DuplicateDocumentError(
            "Ekstraksi ditolak karena nomor putusan sudah pernah diekstraksi.",
            [duplicate_row_to_dict(existing, [reason])],
        )

    text, page_count, _method = extract_pdf_text(pdf_path)
    meta = extract_metadata(pdf_path, text)
    if document_type:
        meta["document_type"] = document_type
    assert_no_duplicate_document(
        putusan_number=meta.get("putusan_number"),
        exclude_document_id=exclude_document_id,
        db_path=db_path,
    )
    document_id = existing["document_id"] if existing else str(uuid.uuid4())
    is_new_document = existing is None
    timestamp = now_iso()
    file_size = pdf_path.stat().st_size
    status = "completed" if text else "failed"

    with connect(db_path) as conn:
        conn.execute(
            """
            INSERT INTO documents (
                document_id, filename, file_path, document_type, file_size, page_count, text,
                putusan_number, putusan_year, court_panel, judge_names, taxpayer_name,
                taxpayer_npwp, taxpayer_address, legal_counsel_name, legal_counsel_license,
                djp_unit, djp_decision_number, skp_number, tax_type, issue_type,
                issue_summary, outcome, amount_disputed, wp_position_summary,
                djp_position_summary, evidence_summary, legal_references_summary,
                court_reasoning_summary,
                extraction_status, extracted_at, created_at, updated_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            ON CONFLICT(file_path) DO UPDATE SET
                filename=excluded.filename,
                document_type=excluded.document_type,
                file_size=excluded.file_size,
                page_count=excluded.page_count,
                text=excluded.text,
                putusan_number=excluded.putusan_number,
                putusan_year=excluded.putusan_year,
                court_panel=excluded.court_panel,
                judge_names=excluded.judge_names,
                taxpayer_name=excluded.taxpayer_name,
                taxpayer_npwp=excluded.taxpayer_npwp,
                taxpayer_address=excluded.taxpayer_address,
                legal_counsel_name=excluded.legal_counsel_name,
                legal_counsel_license=excluded.legal_counsel_license,
                djp_unit=excluded.djp_unit,
                djp_decision_number=excluded.djp_decision_number,
                skp_number=excluded.skp_number,
                tax_type=excluded.tax_type,
                issue_type=excluded.issue_type,
                issue_summary=excluded.issue_summary,
                outcome=excluded.outcome,
                amount_disputed=excluded.amount_disputed,
                wp_position_summary=excluded.wp_position_summary,
                djp_position_summary=excluded.djp_position_summary,
                evidence_summary=excluded.evidence_summary,
                legal_references_summary=excluded.legal_references_summary,
                court_reasoning_summary=excluded.court_reasoning_summary,
                extraction_status=excluded.extraction_status,
                extracted_at=excluded.extracted_at,
                updated_at=excluded.updated_at
            """,
            (
                document_id,
                pdf_path.name,
                abs_path,
                meta["document_type"],
                file_size,
                page_count,
                text,
                meta["putusan_number"],
                meta["putusan_year"],
                meta["court_panel"],
                meta["judge_names"],
                meta["taxpayer_name"],
                meta["taxpayer_npwp"],
                meta["taxpayer_address"],
                meta["legal_counsel_name"],
                meta["legal_counsel_license"],
                meta["djp_unit"],
                meta["djp_decision_number"],
                meta["skp_number"],
                meta["tax_type"],
                meta["issue_type"],
                meta["issue_summary"],
                meta["outcome"],
                meta["amount_disputed"],
                meta["wp_position_summary"],
                meta["djp_position_summary"],
                meta["evidence_summary"],
                meta["legal_references_summary"],
                meta["court_reasoning_summary"],
                status,
                timestamp,
                timestamp,
                timestamp,
            ),
        )
        update_document_rich_fields(conn, document_id, meta, timestamp)
        upsert_document_extraction(conn, document_id, meta, "heuristic", timestamp, text)
        conn.execute("DELETE FROM chunks WHERE document_id = ?", (document_id,))
        for idx, chunk in enumerate(chunk_text(text)):
            conn.execute(
                """
                INSERT INTO chunks (chunk_id, document_id, chunk_order, section_type, text, token_count, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?)
                """,
                (
                    str(uuid.uuid4()),
                    document_id,
                    idx,
                    infer_section_type(chunk),
                    chunk,
                    len(tokenize(chunk)),
                    timestamp,
                ),
            )
        row = conn.execute("SELECT * FROM documents WHERE document_id = ?", (document_id,)).fetchone()

    if use_llm_extraction:
        llm_result = label_document_with_llm(document_id, update_document=True, db_path=db_path)
        if llm_result.get("duplicates"):
            if is_new_document:
                delete_document_records(document_id, db_path=db_path)
            raise DuplicateDocumentError(llm_result.get("error") or "Ekstraksi LLM ditolak karena nomor putusan ganda.", llm_result["duplicates"])
        if llm_result.get("ok"):
            refreshed = get_document(document_id, db_path=db_path)
            if refreshed:
                return refreshed
        if len((text or "").strip()) < 500:
            if is_new_document:
                delete_document_records(document_id, db_path=db_path)
            raise ExtractionError(
                "Ekstraksi gagal karena PDF tidak memiliki text layer dan OCR/vision LLM belum berhasil.",
                document_id=document_id,
                detail=llm_result.get("error") or "",
            )
    return row_to_document(row)


def infer_section_type(text: str) -> str:
    lower = text.lower()
    if "mengadili" in lower or "memutuskan" in lower:
        return "verdict"
    if "pendapat majelis" in lower or "menurut pendapat majelis" in lower:
        return "court_opinion"
    if "menurut pemohon banding" in lower:
        return "wp_position"
    if "menurut terbanding" in lower:
        return "djp_position"
    if "pokok sengketa" in lower:
        return "dispute_issues"
    return "body"


def ingest_directory(
    pdf_dir: Path = ROOT_DIR,
    max_files: Optional[int] = None,
    overwrite: bool = False,
    db_path: Path = DEFAULT_DB_PATH,
) -> List[DocumentRecord]:
    pdfs = find_pdfs(pdf_dir)
    if max_files:
        pdfs = pdfs[:max_files]
    records = []
    for pdf in pdfs:
        records.append(upsert_document(pdf, overwrite=overwrite, db_path=db_path))
    return records


def get_stats(db_path: Path = DEFAULT_DB_PATH) -> Dict[str, Any]:
    init_db(db_path)
    with connect(db_path) as conn:
        total = conn.execute("SELECT COUNT(*) FROM documents").fetchone()[0]
        chunks = conn.execute("SELECT COUNT(*) FROM chunks").fetchone()[0]
        extractions = conn.execute("SELECT COUNT(*) FROM document_extractions").fetchone()[0] if table_exists(conn, "document_extractions") else 0
        regulations = conn.execute("SELECT COUNT(*) FROM tax_regulations").fetchone()[0] if table_exists(conn, "tax_regulations") else 0
        regulation_chunks = conn.execute("SELECT COUNT(*) FROM tax_regulation_chunks").fetchone()[0] if table_exists(conn, "tax_regulation_chunks") else 0
        by_outcome = {
            row["outcome"] or "UNKNOWN": row["count"]
            for row in conn.execute("SELECT outcome, COUNT(*) AS count FROM documents GROUP BY outcome")
        }
        by_tax = {
            row["tax_type"] or "UNKNOWN": row["count"]
            for row in conn.execute("SELECT tax_type, COUNT(*) AS count FROM documents GROUP BY tax_type")
        }
    return {
        "documents": total,
        "chunks": chunks,
        "extractions": extractions,
        "regulations": regulations,
        "regulation_chunks": regulation_chunks,
        "by_outcome": by_outcome,
        "by_tax": by_tax,
    }


def refresh_dashboard_metrics(db_path: Path = DEFAULT_DB_PATH) -> None:
    init_db(db_path)
    timestamp = now_iso()
    with connect(db_path) as conn:
        conn.execute("DELETE FROM dashboard_metrics")
        rows_to_insert = []

        def add(group: str, key: str, value: float, label: str = "") -> None:
            rows_to_insert.append((str(uuid.uuid4()), group, key, float(value), label or key, timestamp))

        add("summary", "documents", conn.execute("SELECT COUNT(*) FROM documents").fetchone()[0], "Dokumen")
        add("summary", "chunks", conn.execute("SELECT COUNT(*) FROM chunks").fetchone()[0], "Chunks")
        add("summary", "llm_labels", conn.execute("SELECT COUNT(*) FROM llm_labels").fetchone()[0], "Label LLM")
        if table_exists(conn, "document_extractions"):
            add("summary", "document_extractions", conn.execute("SELECT COUNT(*) FROM document_extractions").fetchone()[0], "Detail Ekstraksi")
        if table_exists(conn, "tax_regulations"):
            add("summary", "tax_regulations", conn.execute("SELECT COUNT(*) FROM tax_regulations").fetchone()[0], "Peraturan Pajak")
        if table_exists(conn, "tax_regulation_chunks"):
            add("summary", "tax_regulation_chunks", conn.execute("SELECT COUNT(*) FROM tax_regulation_chunks").fetchone()[0], "Chunks Peraturan")

        for group, column in [("outcome", "outcome"), ("tax_type", "tax_type"), ("issue_type", "issue_type"), ("document_type", "document_type")]:
            for row in conn.execute(f"SELECT COALESCE({column}, 'UNKNOWN') AS key, COUNT(*) AS value FROM documents GROUP BY {column}"):
                add(group, row["key"], row["value"], row["key"])
        if table_exists(conn, "tax_regulations"):
            for group, column in [("regulation_type", "regulation_type"), ("regulation_topic", "topic")]:
                for row in conn.execute(f"SELECT COALESCE({column}, 'UNKNOWN') AS key, COUNT(*) AS value FROM tax_regulations GROUP BY {column}"):
                    add(group, row["key"], row["value"], row["key"])

        conn.executemany(
            """
            INSERT INTO dashboard_metrics (metric_id, metric_group, metric_key, metric_value, metric_label, updated_at)
            VALUES (?, ?, ?, ?, ?, ?)
            """,
            rows_to_insert,
        )


def get_dashboard_metrics(db_path: Path = DEFAULT_DB_PATH) -> Dict[str, List[Dict[str, Any]]]:
    init_db(db_path)
    refresh_dashboard_metrics(db_path)
    with connect(db_path) as conn:
        rows = conn.execute(
            "SELECT * FROM dashboard_metrics ORDER BY metric_group, metric_value DESC, metric_key"
        ).fetchall()
    grouped: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
    for row in rows:
        grouped[row["metric_group"]].append(dict(row))
    return dict(grouped)


def list_documents(limit: int = 100, db_path: Path = DEFAULT_DB_PATH) -> List[DocumentRecord]:
    init_db(db_path)
    with connect(db_path) as conn:
        rows = conn.execute(
            "SELECT * FROM documents ORDER BY extracted_at DESC, filename ASC LIMIT ?",
            (limit,),
        ).fetchall()
    return [row_to_document(row) for row in rows]


def search_cases(
    query: str,
    tax_type: str = "ANY",
    issue_type: str = "ANY",
    outcome: str = "ANY",
    limit: int = 10,
    db_path: Path = DEFAULT_DB_PATH,
) -> List[Dict[str, Any]]:
    init_db(db_path)
    with connect(db_path) as conn:
        rows = conn.execute("SELECT * FROM documents WHERE extraction_status = 'completed'").fetchall()
    docs = [row_to_document(row) for row in rows]
    query_tokens = Counter(tokenize(query))
    results = []

    for doc in docs:
        if tax_type != "ANY" and doc.tax_type != tax_type:
            continue
        if issue_type != "ANY" and doc.issue_type != issue_type:
            continue
        if outcome != "ANY" and doc.outcome != outcome:
            continue
        score, reasons = score_document(query, query_tokens, doc)
        if score > 0 or not query_tokens:
            results.append({"document": doc, "score": score, "reasons": reasons})

    results.sort(key=lambda item: item["score"], reverse=True)
    return results[:limit]


def score_document(query: str, query_tokens: Counter, doc: DocumentRecord) -> Tuple[float, List[str]]:
    text = (
        f"{doc.putusan_number} {doc.taxpayer_name} {doc.djp_unit} {doc.tax_type} {doc.issue_type} "
        f"{doc.issue_summary} {doc.wp_position_summary} {doc.djp_position_summary} "
        f"{doc.evidence_summary} {doc.legal_references_summary} {doc.court_reasoning_summary} {doc.text[:20000]}"
    )
    doc_tokens = Counter(tokenize(text))
    reasons = []
    score = 0.0

    if query_tokens:
        overlap = set(query_tokens) & set(doc_tokens)
        weighted_overlap = sum(query_tokens[t] * (1 + math.log1p(doc_tokens[t])) for t in overlap)
        norm = math.sqrt(sum(v * v for v in query_tokens.values())) * math.sqrt(sum(v * v for v in doc_tokens.values()))
        cosine = weighted_overlap / norm if norm else 0
        score += cosine * 100
        if overlap:
            reasons.append("Keyword sama: " + ", ".join(sorted(list(overlap))[:8]))

    lower_query = (query or "").lower()
    if doc.tax_type != "UNKNOWN" and doc.tax_type.lower() in lower_query:
        score += 12
        reasons.append(f"Jenis pajak cocok: {doc.tax_type}")
    issue_words = doc.issue_type.lower().replace("_", " ").split()
    if any(word in lower_query for word in issue_words if len(word) > 3):
        score += 10
        reasons.append(f"Isu cocok: {doc.issue_type}")
    if doc.outcome in {"WP_FULL_WIN", "WP_PARTIAL_WIN"}:
        score += 2
    return score, reasons


def find_similar_cases(intake: Dict[str, Any], limit: int = 8, db_path: Path = DEFAULT_DB_PATH) -> List[Dict[str, Any]]:
    query = " ".join(
        str(intake.get(key, ""))
        for key in [
            "tax_type",
            "issue_type",
            "djp_reason",
            "wp_reason",
            "available_evidence",
            "case_notes",
        ]
    )
    results = search_cases(
        query=query,
        tax_type=intake.get("tax_type_filter", "ANY"),
        issue_type="ANY",
        outcome="ANY",
        limit=limit,
        db_path=db_path,
    )
    return results


def analyze_case(
    intake: Dict[str, Any],
    similar_cases: List[Dict[str, Any]],
    use_llm: bool = False,
    db_path: Path = DEFAULT_DB_PATH,
) -> Dict[str, Any]:
    language = normalize_report_language(intake.get("report_language"))
    distribution = outcome_distribution(similar_cases)
    evidence_score, evidence_gaps = score_evidence(intake)
    positive, negative = keyword_signals(intake, similar_cases)
    historical_score = historical_success_score(distribution)
    similarity_score = average_similarity(similar_cases)
    formal_risk = has_formal_risk(intake)
    relevant_regulations = find_relevant_regulations(intake, db_path=db_path)
    top_case_analysis = build_top_case_analyses(intake, similar_cases, limit=2, language=language)

    score = (
        historical_score * 0.35
        + similarity_score * 0.20
        + evidence_score * 0.25
        + (0 if formal_risk else 70) * 0.10
        + max(0, 65 + len(positive) * 4 - len(negative) * 5) * 0.10
    )
    score = round(max(0, min(100, score)), 1)
    indication, confidence = score_to_label(score, len(similar_cases), distribution, language=language)

    result = {
        "report_language": language,
        "score": score,
        "indication": indication,
        "confidence": confidence,
        "historical_success_rate": distribution,
        "similar_cases_count": len(similar_cases),
        "key_positive_factors": positive[:5],
        "key_negative_factors": negative[:5],
        "supporting_factor_analysis": build_factor_analysis(positive[:5], "supporting", language=language),
        "risk_factor_analysis": build_factor_analysis(negative[:5], "risk", language=language),
        "evidence_gap_analysis": build_evidence_gap_analysis(evidence_gaps, intake, language=language),
        "regulation_analysis": build_regulation_analysis(relevant_regulations, intake, language=language),
        "relevant_regulations": relevant_regulations,
        "evidence_score": evidence_score,
        "evidence_gaps": evidence_gaps,
        "formal_risk": formal_risk,
        "top_case_analysis": top_case_analysis,
        "review": build_review(intake, similar_cases, distribution, positive, negative, evidence_gaps, score, relevant_regulations, language=language),
        "recommendation_draft": build_recommendation(intake, similar_cases, distribution, positive, negative, evidence_gaps, score, indication, relevant_regulations, language=language),
        "disclaimer": (
            "This assessment is indicative and based on comparable decisions and user input; it is not a guarantee of dispute outcome."
            if language == "en"
            else "Indikasi ini bersifat awal berdasarkan putusan pembanding dan data input, bukan kepastian hasil sengketa."
        ),
        "llm_used": False,
        "llm_note": "",
    }

    if use_llm:
        prompt = build_llm_prompt(intake, similar_cases, result)
        llm_text = call_openai_llm(prompt)
        if llm_text:
            result["llm_used"] = True
            result["recommendation_draft"] = llm_text
            result["llm_note"] = (
                "The recommendation draft was deepened with LLM and merged into the recommendation tab."
                if language == "en"
                else "Draft rekomendasi diperdalam dengan LLM dan sudah digabung ke tab rekomendasi."
            )
        else:
            result["llm_note"] = (
                "LLM was not active or failed; the app used the local fallback analysis."
                if language == "en"
                else "LLM tidak aktif atau gagal dipanggil; aplikasi memakai fallback analisis lokal."
            )

    return result


FACTOR_LABELS_EN = {
    "spt masa ppn": "VAT return",
    "bukti pembayaran": "Payment evidence",
    "dokumen pendukung": "Supporting documents",
    "faktur pajak": "VAT invoice",
    "rekonsiliasi": "Reconciliation",
    "konfirmasi lawan transaksi": "Counterparty confirmation",
    "skp/stp": "Tax assessment letter",
    "kep keberatan": "Objection decision",
    "sudah dilaporkan": "Reported position",
    "memenuhi syarat formal": "Formal requirements met",
    "memenuhi syarat material": "Substantive requirements met",
    "koreksi tidak dapat dipertahankan": "Correction may not be sustainable",
    "tidak dapat menunjukkan": "Missing evidence",
    "tidak dapat meyakini": "Insufficient substantiation",
    "jawaban kpp penjual tidak ada": "Seller tax office confirmation unavailable",
    "tidak memberikan alasan": "Reasoning not sufficiently explained",
    "bukti pendukung lainnya tidak ada": "Other supporting evidence unavailable",
    "tidak memenuhi ketentuan formal": "Formal requirement risk",
    "surat bantahan tidak dimasukkan": "Rebuttal letter not submitted",
    "koreksi tetap dipertahankan": "Correction upheld",
    "putusan pembanding teratas memuat outcome yang mendukung wp": "Top comparable decisions support the taxpayer",
    "sebagian putusan pembanding teratas ditolak": "Some top comparable decisions were rejected",
}


def translated_signal_title(value: str, language: str = "id") -> str:
    text = str(value or "").replace("_", " ").strip()
    if normalize_report_language(language) != "en":
        return text.capitalize()
    normalized = " ".join(text.lower().split())
    if normalized in FACTOR_LABELS_EN:
        return FACTOR_LABELS_EN[normalized]
    if "spt" in normalized and "ppn" in normalized:
        return "VAT return"
    if "bukti pembayaran" in normalized or "payment" in normalized:
        return "Payment evidence"
    if "faktur" in normalized:
        return "VAT invoice"
    if "konfirmasi" in normalized or "pkpm" in normalized:
        return "Counterparty confirmation"
    if "dokumen" in normalized:
        return "Supporting documents"
    if "ditolak" in normalized:
        return "Rejected comparable decisions"
    if "putusan pembanding" in normalized:
        return "Comparable decision pattern"
    return text.capitalize()


def translate_match_reason(reason: str, language: str = "id") -> str:
    text = str(reason or "")
    if normalize_report_language(language) != "en":
        return text
    lower = text.lower()
    if lower.startswith("keyword sama:"):
        return "Shared keywords:" + text.split(":", 1)[1]
    if lower.startswith("jenis pajak cocok:"):
        return "Matching tax type:" + text.split(":", 1)[1]
    if lower.startswith("isu cocok:"):
        return "Matching issue:" + text.split(":", 1)[1]
    return text


def translate_legacy_english_text(text: str, language: str = "id") -> str:
    rendered = str(text or "")
    if normalize_report_language(language) != "en":
        return rendered
    replacements = [
        ("putusan pembanding teratas memuat outcome yang mendukung WP", "top comparable decisions support the taxpayer"),
        ("putusan pembanding teratas memuat outcome yang mendukung wp", "top comparable decisions support the taxpayer"),
        ("sebagian putusan pembanding teratas ditolak", "some top comparable decisions were rejected"),
        ("spt masa ppn", "VAT return"),
        ("bukti pembayaran", "payment evidence"),
        ("dokumen pendukung", "supporting documents"),
        ("faktur pajak", "VAT invoice"),
        ("konfirmasi lawan transaksi", "counterparty confirmation"),
        ("rekonsiliasi", "reconciliation"),
        ("sudah dilaporkan", "reported position"),
        ("memenuhi syarat formal", "formal requirements met"),
        ("memenuhi syarat material", "substantive requirements met"),
        ("koreksi tidak dapat dipertahankan", "correction may not be sustainable"),
        ("tidak dapat menunjukkan", "missing evidence"),
        ("tidak dapat meyakini", "insufficient substantiation"),
        ("jawaban kpp penjual tidak ada", "seller tax office confirmation unavailable"),
        ("tidak memberikan alasan", "reasoning not sufficiently explained"),
        ("bukti pendukung lainnya tidak ada", "other supporting evidence unavailable"),
        ("tidak memenuhi ketentuan formal", "formal requirement risk"),
        ("surat bantahan tidak dimasukkan", "rebuttal letter not submitted"),
        ("koreksi tetap dipertahankan", "correction upheld"),
    ]
    for source, target in replacements:
        rendered = re.sub(re.escape(source), target, rendered, flags=re.IGNORECASE)
    return rendered


def build_factor_analysis(factors: List[str], kind: str, language: str = "id") -> List[Dict[str, str]]:
    language = normalize_report_language(language)
    if not factors:
        if language == "en":
            if kind == "supporting":
                return [
                    {
                        "title": "No strong supporting factor has been identified from the input",
                        "analysis": "The input should be enriched with evidence, chronology, legal basis, and the taxpayer's factual position so the system can identify the main strengths of the case.",
                    }
                ]
            return [
                {
                    "title": "No specific risk factor has been detected",
                    "analysis": "This does not mean the case has no risk. Key risks usually become visible after the correction documents, transaction evidence, and the tax authority's legal basis are compared in detail.",
                }
            ]
        if kind == "supporting":
            return [
                {
                    "title": "Belum ada faktor pendukung yang kuat dari input",
                    "analysis": "Data yang dimasukkan masih perlu diperkaya dengan bukti, kronologi, dasar hukum, dan posisi faktual WP agar sistem dapat mengidentifikasi kekuatan utama kasus.",
                }
            ]
        return [
            {
                "title": "Belum ada faktor risiko spesifik yang terbaca",
                "analysis": "Ini bukan berarti risiko tidak ada. Risiko utama biasanya baru terlihat setelah dokumen koreksi, bukti transaksi, dan dasar hukum DJP dibandingkan secara rinci.",
            }
        ]

    result = []
    for factor in factors:
        title = translated_signal_title(factor, language=language)
        if kind == "supporting":
            analysis = explain_positive_factor(factor, language=language)
        else:
            analysis = explain_negative_factor(factor, language=language)
        result.append({"title": title, "analysis": analysis})
    return result


def explain_positive_factor(factor: str, language: str = "id") -> str:
    language = normalize_report_language(language)
    lower = factor.lower()
    if language == "en":
        if "spt" in lower:
            return "The VAT return helps show that the disputed transaction or tax item was included in the taxpayer's reporting, which is useful to rebut an under-reporting narrative."
        if "bukti pembayaran" in lower:
            return "Payment evidence strengthens the substantive proof that the transaction occurred and was not merely supported by formal documents."
        if "faktur" in lower:
            return "The tax invoice is a key VAT document, especially if it can be reconciled with VAT returns, invoices, and payment flows."
        if "putusan pembanding" in lower:
            return "A comparable decision supporting the taxpayer suggests a judicial reasoning pattern that may be useful, provided the facts and evidence are genuinely comparable."
        return "This factor supports the taxpayer because it relates to evidence, reporting consistency, or a relevant historical decision pattern."
    if "spt" in lower:
        return "SPT membantu menunjukkan bahwa transaksi atau pajak yang disengketakan sudah masuk dalam pelaporan WP. Ini berguna untuk membantah narasi bahwa ada objek pajak yang belum dilaporkan."
    if "bukti pembayaran" in lower:
        return "Bukti pembayaran memperkuat pembuktian material bahwa transaksi benar terjadi dan bukan sekadar dokumen formal."
    if "faktur" in lower:
        return "Faktur pajak menjadi bukti utama untuk isu PPN, terutama jika dapat direkonsiliasi dengan SPT, invoice, dan arus pembayaran."
    if "putusan pembanding" in lower:
        return "Putusan pembanding dengan outcome mendukung WP menunjukkan ada pola pertimbangan majelis yang dapat dipakai sebagai rujukan argumentasi, sepanjang fakta dan buktinya sebanding."
    return "Faktor ini mendukung posisi WP karena berkaitan dengan pembuktian, konsistensi pelaporan, atau pola putusan historis yang relevan."


def explain_negative_factor(factor: str, language: str = "id") -> str:
    language = normalize_report_language(language)
    lower = factor.lower()
    if language == "en":
        if "tidak dapat menunjukkan" in lower or "bukti" in lower:
            return "Missing evidence can become the main attack point and may lead the panel to view the taxpayer's position as insufficiently persuasive."
        if "konfirmasi" in lower or "pkpm" in lower:
            return "Unfavorable counterparty confirmation is a key VAT/input-tax risk because it affects transaction and invoice validity."
        if "formal" in lower or "jangka waktu" in lower:
            return "Formal risk must be addressed first because it may prevent substantive review of the dispute."
        if "ditolak" in lower:
            return "Rejected comparable decisions indicate factual or evidentiary patterns that should be avoided in the taxpayer's argument."
        return "This factor should be monitored because it may weaken the link between facts, evidence, legal basis, and the taxpayer's request."
    if "tidak dapat menunjukkan" in lower or "bukti" in lower:
        return "Ketiadaan bukti akan menjadi titik serang utama DJP dan dapat membuat majelis menilai posisi WP tidak cukup meyakinkan."
    if "konfirmasi" in lower or "pkpm" in lower:
        return "Konfirmasi lawan transaksi yang tidak mendukung sering menjadi risiko penting dalam sengketa PPN/Pajak Masukan karena menyentuh validitas transaksi dan faktur."
    if "formal" in lower or "jangka waktu" in lower:
        return "Risiko formal harus diselesaikan terlebih dahulu karena dapat membuat sengketa tidak diperiksa substansinya."
    if "ditolak" in lower:
        return "Adanya putusan pembanding yang ditolak menunjukkan terdapat pola fakta atau bukti yang perlu dihindari dalam penyusunan argumen WP."
    return "Faktor ini perlu diperhatikan karena dapat melemahkan hubungan antara fakta, bukti, dasar hukum, dan permohonan WP."


def build_evidence_gap_analysis(gaps: List[str], intake: Dict[str, Any], language: str = "id") -> List[Dict[str, str]]:
    language = normalize_report_language(language)
    if not gaps:
        return [
            {
                "title": "No key evidence gap detected" if language == "en" else "Tidak ada gap bukti utama yang terdeteksi",
                "analysis": (
                    "Still perform physical document checks and amount reconciliation before submission."
                    if language == "en"
                    else "Tetap lakukan pengecekan fisik dokumen dan rekonsiliasi angka sebelum dokumen diajukan."
                ),
            }
        ]
    return [
        {
            "title": translated_signal_title(gap, language=language),
            "analysis": (
                f"This evidence should be completed or replaced with an explained substitute. For the issue {intake.get('issue_type', '-')}, it helps connect the taxpayer position with transaction facts and corrected amounts."
                if language == "en"
                else f"Bukti ini perlu dilengkapi atau dijelaskan penggantinya. Untuk isu {intake.get('issue_type', '-')}, dokumen tersebut membantu menghubungkan posisi WP dengan fakta transaksi dan angka koreksi."
            ),
        }
        for gap in gaps
    ]


def find_relevant_regulations(intake: Dict[str, Any], db_path: Path = DEFAULT_DB_PATH) -> List[Dict[str, Any]]:
    try:
        from tax_regulation_connector import find_relevant_regulations_for_intake

        results = find_relevant_regulations_for_intake(intake, limit=6, db_path=db_path)
        return [slim_regulation_result(item) for item in results]
    except Exception:
        return []


def slim_regulation_result(item: Dict[str, Any]) -> Dict[str, Any]:
    matched_chunks = item.get("matched_chunks") or []
    return {
        "regulation_id": item.get("regulation_id", ""),
        "source": item.get("source", ""),
        "source_id": item.get("source_id", ""),
        "url": item.get("url", ""),
        "title": item.get("title", ""),
        "regulation_type": item.get("regulation_type", ""),
        "number": item.get("number", ""),
        "year": item.get("year"),
        "topic": item.get("topic", ""),
        "category": item.get("category", ""),
        "published_date": item.get("published_date", ""),
        "summary": item.get("summary", ""),
        "score": item.get("score", 0),
        "reasons": item.get("reasons", [])[:3],
        "matched_chunks": [
            {
                "section_label": chunk.get("section_label", ""),
                "text": chunk.get("text", "")[:800],
                "score": chunk.get("score", 0),
            }
            for chunk in matched_chunks[:2]
        ],
    }


def build_regulation_analysis(regulations: List[Dict[str, Any]], intake: Dict[str, Any], language: str = "id") -> List[Dict[str, str]]:
    language = normalize_report_language(language)
    if not regulations:
        return [
                {
                    "title": "No matching local VAT regulation yet" if language == "en" else "Belum ada peraturan PPN lokal yang cocok",
                    "analysis": (
                        "Download or refresh regulations in the Regulations menu so the system can link the case to relevant laws/regulations."
                        if language == "en"
                        else "Unduh atau refresh peraturan di menu Peraturan agar sistem dapat menautkan kasus ke UU/PMK/PER terkait."
                    ),
                }
            ]
    analysis = []
    issue = intake.get("issue_type") or "isu PPN"
    for reg in regulations[:5]:
        snippet = ""
        if reg.get("matched_chunks"):
            snippet = reg["matched_chunks"][0].get("text", "")
        focus = snippet[:260] + ("..." if len(snippet) > 260 else "")
        analysis.append(
            {
                "title": f"{reg.get('title') or reg.get('number')} ({reg.get('score', 0)})",
                "analysis": (
                    (
                        f"Relevant as an initial legal check for {issue}. Summary: {reg.get('summary') or '-'}. "
                        f"Closest indexed excerpt: {focus or 'no specific excerpt yet'}"
                    )
                    if language == "en"
                    else (
                        f"Relevan sebagai dasar cek untuk {issue}. Ringkasan: {reg.get('summary') or '-'}. "
                        f"Bagian yang paling dekat dari indeks lokal: {focus or 'belum ada cuplikan spesifik'}"
                    )
                ),
            }
        )
    return analysis


def format_regulation_review_text(regulations: List[Dict[str, Any]], language: str = "id") -> str:
    language = normalize_report_language(language)
    if not regulations:
        return (
            "VAT regulation support is not yet included because the local regulation database is empty or no match was found."
            if language == "en"
            else "Dasar peraturan PPN belum masuk ke analisis karena database peraturan lokal masih kosong atau belum menemukan kecocokan."
        )
    labels = []
    for reg in regulations[:4]:
        label = reg.get("title") or reg.get("number") or reg.get("source_id")
        if reg.get("number") and reg.get("number") not in label:
            label = f"{label} ({reg['number']})"
        labels.append(label)
    if language == "en":
        return (
            "Initial regulation references to verify: "
            + "; ".join(labels)
            + ". These rules are used as context, not final citations; the relevant articles must still be verified against the tax period and transaction type."
        )
    return (
        "Dasar peraturan awal yang perlu dicek: "
        + "; ".join(labels)
        + ". Peraturan ini dipakai sebagai konteks, bukan sebagai kutipan final, sehingga pasal yang dipakai tetap harus diverifikasi terhadap masa pajak dan jenis transaksi."
    )


def format_regulation_recommendation_lines(regulations: List[Dict[str, Any]], language: str = "id") -> List[str]:
    language = normalize_report_language(language)
    if not regulations:
        return (
            ["No VAT regulation has been indexed yet. Run the VAT seed download from Ortax before preparing the final memo."]
            if language == "en"
            else ["Belum ada peraturan PPN yang terindeks. Jalankan unduh seed PPN dari Ortax sebelum membuat memo final."]
        )
    lines = []
    for idx, reg in enumerate(regulations[:6], start=1):
        title = reg.get("title") or reg.get("number") or reg.get("source_id")
        summary = reg.get("summary") or "-"
        chunk = ""
        if reg.get("matched_chunks"):
            first = reg["matched_chunks"][0]
            chunk = (
                f" Closest section: {first.get('section_label') or '-'}."
                if language == "en"
                else f" Bagian terdekat: {first.get('section_label') or '-'}."
            )
        if language == "en":
            lines.append(f"{idx}. {title} - relevance {reg.get('score', 0)}. Initial focus: {summary[:280]}.{chunk}")
        else:
            lines.append(f"{idx}. {title} - relevansi {reg.get('score', 0)}. Fokus awal: {summary[:280]}.{chunk}")
    return lines


def get_document(document_id: str, db_path: Path = DEFAULT_DB_PATH) -> Optional[DocumentRecord]:
    init_db(db_path)
    with connect(db_path) as conn:
        row = conn.execute("SELECT * FROM documents WHERE document_id = ?", (document_id,)).fetchone()
    return row_to_document(row) if row else None


def latest_llm_label(document_id: str, db_path: Path = DEFAULT_DB_PATH) -> Optional[Dict[str, Any]]:
    init_db(db_path)
    with connect(db_path) as conn:
        row = conn.execute(
            "SELECT * FROM llm_labels WHERE document_id = ? ORDER BY created_at DESC LIMIT 1",
            (document_id,),
        ).fetchone()
    if not row:
        return None
    data = dict(row)
    data["label"] = json.loads(row["label_json"])
    return data


def get_document_extraction(document_id: str, db_path: Path = DEFAULT_DB_PATH) -> Optional[Dict[str, Any]]:
    init_db(db_path)
    with connect(db_path) as conn:
        row = conn.execute(
            "SELECT * FROM document_extractions WHERE document_id = ?",
            (document_id,),
        ).fetchone()
    if not row:
        return None
    data = dict(row)
    try:
        data["extraction"] = json.loads(row["extraction_json"])
    except Exception:
        data["extraction"] = {}
    return data


def backfill_missing_document_extractions(db_path: Path = DEFAULT_DB_PATH) -> int:
    init_db(db_path)
    timestamp = now_iso()
    with connect(db_path) as conn:
        rows = conn.execute(
            """
            SELECT d.*
            FROM documents d
            LEFT JOIN document_extractions e ON e.document_id = d.document_id
            WHERE e.document_id IS NULL AND d.extraction_status = 'completed'
            """
        ).fetchall()
        for row in rows:
            meta = dict(row)
            if not meta.get("success_level"):
                meta["success_level"] = outcome_to_success_level(meta.get("outcome"))
            upsert_document_extraction(conn, row["document_id"], meta, "backfill", timestamp, row["text"] or "")
    return len(rows)


def list_llm_labels(limit: int = 100, db_path: Path = DEFAULT_DB_PATH) -> List[Dict[str, Any]]:
    init_db(db_path)
    with connect(db_path) as conn:
        rows = conn.execute(
            """
            SELECT l.*, d.filename, d.putusan_number
            FROM llm_labels l
            JOIN documents d ON d.document_id = l.document_id
            ORDER BY l.created_at DESC
            LIMIT ?
            """,
            (limit,),
        ).fetchall()
    labels = []
    for row in rows:
        item = dict(row)
        item["label"] = json.loads(row["label_json"])
        labels.append(item)
    return labels


def label_document_with_llm(
    document_id: str,
    update_document: bool = True,
    db_path: Path = DEFAULT_DB_PATH,
) -> Dict[str, Any]:
    doc = get_document(document_id, db_path=db_path)
    if not doc:
        return {"ok": False, "error": "Dokumen tidak ditemukan.", "label": None}

    prompt = build_labeling_prompt(doc)
    source = "text"
    rendered_pages: List[int] = []
    if len((doc.text or "").strip()) < 500:
        try:
            image_data_urls, rendered_pages = render_pdf_pages_to_data_urls(Path(doc.file_path))
        except Exception as exc:
            return {"ok": False, "error": f"PDF tidak punya text layer dan render OCR/vision gagal: {exc}", "label": None}
        prompt += (
            "\n\nCATATAN OCR/VISION:\n"
            "Teks PDF sangat minim atau kosong. Gunakan gambar halaman PDF yang dilampirkan untuk membaca visual dokumen. "
            "Ekstrak field dari gambar tersebut, terutama halaman awal, bagian tengah, dan halaman akhir. "
            f"Halaman yang dikirim: {', '.join(str(page) for page in rendered_pages)}. "
            "Jika data tidak terlihat di halaman yang dikirim, isi null atau UNKNOWN dan beri catatan pada review_notes."
        )
        llm = call_openai_vision(prompt, image_data_urls, max_output_tokens=4200)
        source = "vision"
    else:
        llm = call_openai_text(prompt, max_output_tokens=4200)
    if not llm.get("text"):
        return {"ok": False, "error": llm.get("error") or "LLM tidak mengembalikan teks.", "label": None}

    try:
        label = parse_json_from_text(llm["text"] or "")
    except ValueError as exc:
        return {"ok": False, "error": f"Output LLM bukan JSON valid: {exc}. Raw: {(llm['text'] or '')[:500]}", "label": None}

    normalized = normalize_llm_label(label, doc)
    if update_document:
        duplicates = find_duplicate_documents(
            putusan_number=normalized.get("putusan_number"),
            exclude_document_id=doc.document_id,
            db_path=db_path,
        )
        if duplicates:
            return {
                "ok": False,
                "error": "Label LLM ditolak karena nomor putusan sama dengan dokumen yang sudah diekstraksi.",
                "label": normalized,
                "duplicates": duplicates,
                "model": llm.get("model"),
                "source": source,
            }
    save_llm_label(doc.document_id, normalized, llm.get("model") or "unknown", update_document, db_path)
    return {"ok": True, "error": None, "label": normalized, "model": llm.get("model"), "source": source, "pages": rendered_pages}


def batch_label_documents_with_llm(
    limit: int = 5,
    only_unlabeled: bool = True,
    update_document: bool = True,
    db_path: Path = DEFAULT_DB_PATH,
) -> List[Dict[str, Any]]:
    init_db(db_path)
    with connect(db_path) as conn:
        if only_unlabeled:
            rows = conn.execute(
                """
                SELECT d.document_id
                FROM documents d
                LEFT JOIN llm_labels l ON l.document_id = d.document_id
                WHERE (d.extraction_status = 'completed' OR length(trim(COALESCE(d.text, ''))) < 500)
                  AND l.document_id IS NULL
                ORDER BY d.filename ASC
                LIMIT ?
                """,
                (limit,),
            ).fetchall()
        else:
            rows = conn.execute(
                """
                SELECT document_id
                FROM documents
                WHERE extraction_status = 'completed' OR length(trim(COALESCE(text, ''))) < 500
                ORDER BY filename ASC
                LIMIT ?
                """,
                (limit,),
            ).fetchall()

    results = []
    for row in rows:
        results.append(label_document_with_llm(row["document_id"], update_document=update_document, db_path=db_path))
    return results


def build_labeling_prompt(doc: DocumentRecord) -> str:
    snippets = {
        "header": doc.text[:9000],
        "pokok_sengketa": extract_section_snippet(
            doc.text,
            ["pokok sengketa", "sengketa pajak", "menurut pemohon banding"],
            ["koreksi", "dasar pengenaan pajak", "pajak masukan"],
            limit=5000,
        ),
        "wp_position": extract_section_snippet(
            doc.text,
            ["menurut pemohon banding", "pemohon banding menyatakan"],
            ["pemohon banding"],
            limit=5000,
        ),
        "djp_position": extract_section_snippet(
            doc.text,
            ["menurut terbanding", "terbanding menyatakan"],
            ["terbanding"],
            limit=5000,
        ),
        "court_opinion": extract_section_snippet(
            doc.text,
            ["pendapat majelis", "menurut pendapat majelis", "majelis berpendapat"],
            ["menimbang"],
            limit=5000,
        ),
        "verdict": doc.text[-9000:],
    }
    return (
        "Anda adalah analis putusan Pengadilan Pajak Indonesia. Ekstrak dokumen berikut "
        "menjadi JSON terstruktur untuk database sengketa pajak. Gunakan hanya informasi dalam teks. Jangan mengarang. "
        "Jika tidak ada data, isi null atau UNKNOWN. Jangan isi nama WP dengan label umum seperti 'Nomor', 'Nama', atau 'Pemohon Banding'. "
        "Ringkasan harus substantif: jelaskan objek koreksi, alasan DJP, bantahan WP, bukti, pertimbangan majelis, amar, dan outcome per isu jika ada. "
        "Return hanya JSON valid tanpa markdown.\n\n"
        "Kelompok field minimum yang harus Anda usahakan lengkap:\n"
        "- metadata_putusan: nomor putusan, tahun, majelis, jenis acara, tingkat pemeriksaan, nomor berkas, tanggal putusan.\n"
        "- objek_sengketa: jenis pajak, masa pajak, nomor SKP/STP/SKPLB, nomor keputusan keberatan, nilai menurut WP dan DJP.\n"
        "- pihak: nama WP, NPWP, alamat, wakil, kuasa hukum, izin kuasa, Terbanding, unit DJP.\n"
        "- pokok_sengketa: issue type/subtype, nilai sengketa, objek koreksi, alasan koreksi, alasan bantahan WP.\n"
        "- argumen: posisi WP, posisi DJP, dasar hukum dikutip, bukti diajukan, bantahan.\n"
        "- pertimbangan: pendapat Majelis, argumen diterima/ditolak, bukti cukup/tidak cukup.\n"
        "- outcome: amar putusan, hasil per isu, nilai pajak sebelum/sesudah, pengurangan koreksi, sanksi, success level.\n\n"
        "Skema JSON wajib:\n"
        "{\n"
        '  "document_type": "putusan_pengadilan|surat_banding|surat_keberatan|jawaban_sphp|uraian_banding|bantahan|dokumen_pendukung|UNKNOWN",\n'
        '  "putusan_number": string,\n'
        '  "putusan_year": number|null,\n'
        '  "procedure_type": string|null,\n'
        '  "examination_level": string|null,\n'
        '  "case_file_number": string|null,\n'
        '  "decision_date": string|null,\n'
        '  "taxpayer_name": string|null,\n'
        '  "taxpayer_npwp": string|null,\n'
        '  "taxpayer_address": string|null,\n'
        '  "representative_name": string|null,\n'
        '  "legal_counsel_name": string|null,\n'
        '  "legal_counsel_license": string|null,\n'
        '  "appellee_name": string|null,\n'
        '  "djp_unit": string|null,\n'
        '  "djp_decision_number": string|null,\n'
        '  "skp_number": string|null,\n'
        '  "court_panel": string|null,\n'
        '  "judge_names": string|null,\n'
        '  "tax_type": "PPN|PPh Badan|PPh 21|PPh 23|PPh 26|PBB|Lainnya|UNKNOWN",\n'
        '  "tax_period": string|null,\n'
        '  "dispute_stage": "Banding|Gugatan|UNKNOWN",\n'
        '  "issue_type": "DPP_PPN|PAJAK_MASUKAN|FAKTUR_PAJAK|PKPM_KONFIRMASI|SANKSI|FORMAL|PPh|LAINNYA|UNKNOWN",\n'
        '  "issue_subtype": string|null,\n'
        '  "correction_object": string|null,\n'
        '  "correction_reason": string|null,\n'
        '  "wp_rebuttal_reason": string|null,\n'
        '  "outcome": "WP_FULL_WIN|WP_PARTIAL_WIN|DJP_WIN|FORMAL_REJECTED|DROPPED|UNKNOWN",\n'
        '  "issue_summary": string,\n'
        '  "wp_position_summary": string,\n'
        '  "djp_position_summary": string,\n'
        '  "court_reasoning_summary": string,\n'
        '  "evidence_summary": string,\n'
        '  "evidence_submitted": string,\n'
        '  "legal_references_summary": string,\n'
        '  "amount_disputed": number|null,\n'
        '  "wp_claim_amount": number|null,\n'
        '  "djp_claim_amount": number|null,\n'
        '  "accepted_arguments": string|null,\n'
        '  "rejected_arguments": string|null,\n'
        '  "sufficient_evidence_summary": string|null,\n'
        '  "insufficient_evidence_summary": string|null,\n'
        '  "verdict_text": string|null,\n'
        '  "per_issue_outcome": string|null,\n'
        '  "tax_before_amount": number|null,\n'
        '  "tax_after_amount": number|null,\n'
        '  "correction_reduction_amount": number|null,\n'
        '  "sanctions_amount": number|null,\n'
        '  "success_level": string|null,\n'
        '  "confidence": number,\n'
        '  "needs_human_review": boolean,\n'
        '  "review_notes": string\n'
        "}\n\n"
        f"FILE: {doc.filename}\n"
        f"HEURISTIC_LABEL: document_type={doc.document_type}, tax_type={doc.tax_type}, issue_type={doc.issue_type}, outcome={doc.outcome}, taxpayer={doc.taxpayer_name}\n\n"
        f"SNIPPETS:\n{json.dumps(snippets, ensure_ascii=False, indent=2)}"
    )


def parse_json_from_text(text: str) -> Dict[str, Any]:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?", "", cleaned).strip()
        cleaned = re.sub(r"```$", "", cleaned).strip()
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        start = cleaned.find("{")
        end = cleaned.rfind("}")
        if start >= 0 and end > start:
            return json.loads(cleaned[start : end + 1])
    raise ValueError(cleaned[:300])


def normalize_llm_label(label: Dict[str, Any], doc: DocumentRecord) -> Dict[str, Any]:
    allowed_outcomes = set(OUTCOME_LABELS) | {"DROPPED"}
    allowed_issues = {"DPP_PPN", "PAJAK_MASUKAN", "FAKTUR_PAJAK", "PKPM_KONFIRMASI", "SANKSI", "FORMAL", "PPh", "LAINNYA", "UNKNOWN"}
    normalized = dict(label)
    normalized["document_type"] = normalize_spaces(str(label.get("document_type") or doc.document_type or "UNKNOWN"))
    normalized["putusan_number"] = normalize_spaces(str(label.get("putusan_number") or doc.putusan_number))
    normalized["tax_type"] = normalize_spaces(str(label.get("tax_type") or doc.tax_type or "UNKNOWN"))
    issue_type = normalize_spaces(str(label.get("issue_type") or doc.issue_type or "UNKNOWN")).upper().replace(" ", "_")
    if issue_type == "PPH":
        normalized["issue_type"] = "PPh"
    else:
        normalized["issue_type"] = issue_type if issue_type in allowed_issues else "LAINNYA"
    outcome = normalize_spaces(str(label.get("outcome") or doc.outcome or "UNKNOWN")).upper()
    normalized["outcome"] = outcome if outcome in allowed_outcomes else "UNKNOWN"
    for key in [
        "dispute_stage",
        "issue_subtype",
        "issue_summary",
        "wp_position_summary",
        "djp_position_summary",
        "court_reasoning_summary",
        "evidence_summary",
        "legal_references_summary",
        "review_notes",
        "taxpayer_name",
        "taxpayer_npwp",
        "taxpayer_address",
        "legal_counsel_name",
        "legal_counsel_license",
        "djp_unit",
        "djp_decision_number",
        "skp_number",
        "court_panel",
        "judge_names",
        "procedure_type",
        "examination_level",
        "case_file_number",
        "decision_date",
        "tax_period",
        "representative_name",
        "appellee_name",
        "correction_object",
        "correction_reason",
        "wp_rebuttal_reason",
        "evidence_submitted",
        "accepted_arguments",
        "rejected_arguments",
        "sufficient_evidence_summary",
        "insufficient_evidence_summary",
        "verdict_text",
        "per_issue_outcome",
        "success_level",
    ]:
        value = label.get(key)
        normalized[key] = normalize_spaces(str(value)) if value is not None else None
    for numeric_key, fallback in {
        "amount_disputed": doc.amount_disputed,
        "wp_claim_amount": None,
        "djp_claim_amount": None,
        "tax_before_amount": None,
        "tax_after_amount": None,
        "correction_reduction_amount": None,
        "sanctions_amount": None,
    }.items():
        try:
            normalized[numeric_key] = float(label[numeric_key]) if label.get(numeric_key) is not None else fallback
        except (TypeError, ValueError):
            normalized[numeric_key] = fallback
    try:
        normalized["putusan_year"] = int(label["putusan_year"]) if label.get("putusan_year") is not None else doc.putusan_year
    except (TypeError, ValueError):
        normalized["putusan_year"] = doc.putusan_year
    try:
        normalized["confidence"] = max(0.0, min(1.0, float(label.get("confidence", 0.5))))
    except (TypeError, ValueError):
        normalized["confidence"] = 0.5
    normalized["needs_human_review"] = bool(label.get("needs_human_review", normalized["confidence"] < 0.7))
    return normalized


def save_llm_label(
    document_id: str,
    label: Dict[str, Any],
    model: str,
    update_document: bool,
    db_path: Path,
) -> None:
    timestamp = now_iso()
    with connect(db_path) as conn:
        conn.execute(
            """
            INSERT INTO llm_labels (
                label_id, document_id, model, label_json, tax_type, dispute_stage,
                issue_type, issue_subtype, outcome, issue_summary, wp_position_summary,
                djp_position_summary, court_reasoning_summary, evidence_summary,
                legal_references_summary, document_type, taxpayer_name, taxpayer_npwp,
                taxpayer_address, legal_counsel_name, legal_counsel_license, djp_unit,
                djp_decision_number, skp_number, court_panel, judge_names, confidence, created_at
            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """,
            (
                str(uuid.uuid4()),
                document_id,
                model,
                json.dumps(label, ensure_ascii=False),
                label.get("tax_type"),
                label.get("dispute_stage"),
                label.get("issue_type"),
                label.get("issue_subtype"),
                label.get("outcome"),
                label.get("issue_summary"),
                label.get("wp_position_summary"),
                label.get("djp_position_summary"),
                label.get("court_reasoning_summary"),
                label.get("evidence_summary"),
                label.get("legal_references_summary"),
                label.get("document_type"),
                label.get("taxpayer_name"),
                label.get("taxpayer_npwp"),
                label.get("taxpayer_address"),
                label.get("legal_counsel_name"),
                label.get("legal_counsel_license"),
                label.get("djp_unit"),
                label.get("djp_decision_number"),
                label.get("skp_number"),
                label.get("court_panel"),
                label.get("judge_names"),
                label.get("confidence"),
                timestamp,
            ),
        )
        if update_document:
            conn.execute(
                """
                UPDATE documents
                SET putusan_number = COALESCE(NULLIF(?, ''), putusan_number),
                    document_type = COALESCE(NULLIF(?, ''), document_type),
                    taxpayer_name = COALESCE(NULLIF(?, ''), taxpayer_name),
                    taxpayer_npwp = COALESCE(NULLIF(?, ''), taxpayer_npwp),
                    taxpayer_address = COALESCE(NULLIF(?, ''), taxpayer_address),
                    legal_counsel_name = COALESCE(NULLIF(?, ''), legal_counsel_name),
                    legal_counsel_license = COALESCE(NULLIF(?, ''), legal_counsel_license),
                    djp_unit = COALESCE(NULLIF(?, ''), djp_unit),
                    djp_decision_number = COALESCE(NULLIF(?, ''), djp_decision_number),
                    skp_number = COALESCE(NULLIF(?, ''), skp_number),
                    court_panel = COALESCE(NULLIF(?, ''), court_panel),
                    judge_names = COALESCE(NULLIF(?, ''), judge_names),
                    tax_type = COALESCE(NULLIF(?, ''), tax_type),
                    issue_type = COALESCE(NULLIF(?, ''), issue_type),
                    issue_summary = COALESCE(NULLIF(?, ''), issue_summary),
                    outcome = COALESCE(NULLIF(?, ''), outcome),
                    amount_disputed = COALESCE(?, amount_disputed),
                    wp_position_summary = COALESCE(NULLIF(?, ''), wp_position_summary),
                    djp_position_summary = COALESCE(NULLIF(?, ''), djp_position_summary),
                    evidence_summary = COALESCE(NULLIF(?, ''), evidence_summary),
                    legal_references_summary = COALESCE(NULLIF(?, ''), legal_references_summary),
                    court_reasoning_summary = COALESCE(NULLIF(?, ''), court_reasoning_summary),
                    label_source = 'llm',
                    llm_labeled_at = ?,
                    llm_confidence = ?,
                    updated_at = ?
                WHERE document_id = ?
                """,
                (
                    label.get("putusan_number"),
                    label.get("document_type"),
                    label.get("taxpayer_name"),
                    label.get("taxpayer_npwp"),
                    label.get("taxpayer_address"),
                    label.get("legal_counsel_name"),
                    label.get("legal_counsel_license"),
                    label.get("djp_unit"),
                    label.get("djp_decision_number"),
                    label.get("skp_number"),
                    label.get("court_panel"),
                    label.get("judge_names"),
                    label.get("tax_type"),
                    label.get("issue_type"),
                    label.get("issue_summary"),
                    label.get("outcome"),
                    label.get("amount_disputed"),
                    label.get("wp_position_summary"),
                    label.get("djp_position_summary"),
                    label.get("evidence_summary"),
                    label.get("legal_references_summary"),
                    label.get("court_reasoning_summary"),
                    timestamp,
                    label.get("confidence"),
                    timestamp,
                    document_id,
                ),
            )
            conn.execute(
                """
                UPDATE documents
                SET extraction_status = 'completed',
                    extracted_at = COALESCE(NULLIF(extracted_at, ''), ?),
                    updated_at = ?
                WHERE document_id = ?
                """,
                (timestamp, timestamp, document_id),
            )
            update_document_rich_fields(conn, document_id, label, timestamp)
            upsert_document_extraction(conn, document_id, label, "llm", timestamp)


def build_intake_from_document(document_id: str, db_path: Path = DEFAULT_DB_PATH) -> Dict[str, Any]:
    doc = get_document(document_id, db_path=db_path)
    if not doc:
        return {}
    label_row = latest_llm_label(document_id, db_path=db_path)
    label = label_row["label"] if label_row else {}
    issue_type = label.get("issue_type") or doc.issue_type
    display_issue = {
        "DPP_PPN": "DPP PPN",
        "PAJAK_MASUKAN": "Pajak Masukan",
        "FAKTUR_PAJAK": "Faktur Pajak",
        "PKPM_KONFIRMASI": "PKPM / Konfirmasi",
        "SANKSI": "Sanksi",
        "FORMAL": "Formal",
        "PPh": "PPh",
    }.get(issue_type, issue_type or "DPP PPN")
    tax_type = label.get("tax_type") or doc.tax_type
    if tax_type == "UNKNOWN":
        tax_type = "PPN"

    return {
        "taxpayer_name": label.get("taxpayer_name") or doc.taxpayer_name or "Wajib Pajak",
        "taxpayer_npwp": label.get("taxpayer_npwp") or doc.taxpayer_npwp,
        "taxpayer_address": label.get("taxpayer_address") or doc.taxpayer_address,
        "legal_counsel_name": label.get("legal_counsel_name") or doc.legal_counsel_name,
        "djp_unit": label.get("djp_unit") or doc.djp_unit,
        "djp_decision_number": label.get("djp_decision_number") or doc.djp_decision_number,
        "skp_number": label.get("skp_number") or doc.skp_number,
        "document_type": label.get("document_type") or doc.document_type,
        "stage": label.get("dispute_stage") or "Banding",
        "tax_type": tax_type,
        "tax_type_filter": tax_type,
        "issue_type": display_issue,
        "correction_amount": format_rupiah(label.get("amount_disputed") or doc.amount_disputed),
        "djp_reason": label.get("djp_position_summary") or doc.djp_position_summary or "DJP melakukan koreksi sesuai uraian dalam putusan/dokumen.",
        "wp_reason": label.get("wp_position_summary") or doc.wp_position_summary or "WP membantah koreksi dan menyiapkan bukti pendukung.",
        "available_evidence": infer_evidence_list(label.get("evidence_summary") or doc.evidence_summary or doc.text[:20000]),
        "case_notes": label.get("issue_summary") or doc.issue_summary,
    }


def format_rupiah(value: Any) -> str:
    if value is None:
        return "Rp 0"
    try:
        return "Rp {:,.0f}".format(float(value)).replace(",", ".")
    except (TypeError, ValueError):
        return str(value)


def infer_evidence_list(text: str) -> List[str]:
    lower = (text or "").lower()
    mapping = [
        ("Faktur Pajak", ["faktur pajak"]),
        ("SPT Masa PPN", ["spt masa ppn"]),
        ("SPT Pembetulan", ["spt pembetulan"]),
        ("Invoice", ["invoice"]),
        ("Kontrak", ["kontrak", "perjanjian"]),
        ("Purchase Order", ["purchase order", "po "]),
        ("Delivery Order / Surat Jalan", ["delivery order", "surat jalan"]),
        ("Bukti Pembayaran", ["bukti pembayaran", "transfer", "ssp"]),
        ("Rekening Koran", ["rekening koran"]),
        ("Rekonsiliasi", ["rekonsiliasi"]),
        ("Konfirmasi Lawan Transaksi", ["konfirmasi", "pkpm", "kpp penjual"]),
        ("SKP/STP", ["skp", "stp"]),
        ("KEP Keberatan", ["kep-", "keputusan keberatan"]),
        ("Surat Kuasa", ["surat kuasa"]),
        ("Bukti Pengiriman", ["bukti pengiriman", "pos"]),
    ]
    result = []
    for label, keywords in mapping:
        if any(keyword in lower for keyword in keywords):
            result.append(label)
    return result or ["Faktur Pajak", "SPT Masa PPN", "Bukti Pembayaran", "Rekonsiliasi"]


def outcome_distribution(similar_cases: List[Dict[str, Any]]) -> Dict[str, float]:
    weighted = defaultdict(float)
    total = 0.0
    for item in similar_cases:
        outcome = item["document"].outcome or "UNKNOWN"
        weight = max(1.0, float(item.get("score") or 0.0))
        weighted[outcome] += weight
        total += weight
    if total <= 0:
        return {label: 0.0 for label in OUTCOME_LABELS}
    return {label: round(weighted.get(label, 0.0) / total, 3) for label in OUTCOME_LABELS}


def normalize_report_language(value: Any) -> str:
    lang = str(value or "id").strip().lower()
    if lang in {"en", "english", "inggris"}:
        return "en"
    return "id"


def outcome_label(outcome: str, language: str = "id") -> str:
    labels = OUTCOME_LABELS_EN if normalize_report_language(language) == "en" else OUTCOME_LABELS
    return labels.get(outcome or "UNKNOWN", outcome or labels["UNKNOWN"])


def historical_success_score(distribution: Dict[str, float]) -> float:
    return round(
        100
        * (
            distribution.get("WP_FULL_WIN", 0.0)
            + distribution.get("WP_PARTIAL_WIN", 0.0) * 0.72
            + distribution.get("DJP_WIN", 0.0) * 0.18
            + distribution.get("FORMAL_REJECTED", 0.0) * 0.08
        ),
        1,
    )


def average_similarity(similar_cases: List[Dict[str, Any]]) -> float:
    if not similar_cases:
        return 0.0
    scores = [min(100.0, float(item.get("score") or 0.0)) for item in similar_cases]
    return round(sum(scores) / len(scores), 1)


def score_evidence(intake: Dict[str, Any]) -> Tuple[float, List[str]]:
    issue = str(intake.get("issue_type", "")).lower()
    selected = {str(item).lower() for item in intake.get("available_evidence", [])}
    expected = ["faktur pajak", "spt masa ppn", "rekonsiliasi", "bukti pembayaran"]
    if "pajak masukan" in issue:
        expected = ["faktur pajak", "spt masa ppn", "bukti pembayaran", "konfirmasi lawan transaksi"]
    if "sanksi" in issue:
        expected = ["skp/stp", "perhitungan sanksi", "dasar hukum", "bukti pembayaran"]
    if "formal" in issue:
        expected = ["tanggal terima keputusan", "bukti pengiriman", "bukti pembayaran", "surat kuasa"]

    hits = []
    gaps = []
    for item in expected:
        found = any(item in evidence or evidence in item for evidence in selected)
        if found:
            hits.append(item)
        else:
            gaps.append(item)
    score = round(100 * len(hits) / len(expected), 1) if expected else 60.0
    return score, gaps


def keyword_signals(intake: Dict[str, Any], similar_cases: List[Dict[str, Any]]) -> Tuple[List[str], List[str]]:
    text_parts = [
        str(intake.get("djp_reason", "")),
        str(intake.get("wp_reason", "")),
        str(intake.get("case_notes", "")),
        " ".join(str(x) for x in intake.get("available_evidence", [])),
    ]
    for item in similar_cases[:5]:
        doc = item["document"]
        text_parts.append(doc.issue_summary)
        text_parts.append(doc.court_reasoning_summary)
    lower = " ".join(text_parts).lower()
    positive = [keyword for keyword in POSITIVE_KEYWORDS if keyword in lower]
    negative = [keyword for keyword in NEGATIVE_KEYWORDS if keyword in lower]

    if any(item["document"].outcome in {"WP_FULL_WIN", "WP_PARTIAL_WIN"} for item in similar_cases[:5]):
        positive.append("putusan pembanding teratas memuat outcome yang mendukung WP")
    if any(item["document"].outcome == "DJP_WIN" for item in similar_cases[:5]):
        negative.append("sebagian putusan pembanding teratas ditolak")
    return dedupe(positive), dedupe(negative)


def dedupe(items: Iterable[str]) -> List[str]:
    seen = set()
    result = []
    for item in items:
        if item not in seen:
            seen.add(item)
            result.append(item)
    return result


def has_formal_risk(intake: Dict[str, Any]) -> bool:
    text = " ".join(
        [
            str(intake.get("stage", "")),
            str(intake.get("djp_reason", "")),
            str(intake.get("wp_reason", "")),
            str(intake.get("case_notes", "")),
        ]
    ).lower()
    return any(key in text for key in ["lewat jangka waktu", "tidak dapat diterima", "formal", "terlambat"])


def score_to_label(score: float, similar_count: int, distribution: Dict[str, float], language: str = "id") -> Tuple[str, str]:
    language = normalize_report_language(language)
    if similar_count == 0:
        return ("Comparable data is not sufficient yet", "low") if language == "en" else ("Data pembanding belum cukup", "low")
    if language == "en":
        if score >= 75:
            indication = "Strong chance of being granted or partially granted"
        elif score >= 55:
            indication = "Potentially partially granted"
        elif score >= 35:
            indication = "Medium-high risk; evidence needs strengthening"
        else:
            indication = "High risk of rejection or formal dismissal"
    else:
        if score >= 75:
            indication = "Berpeluang kuat dikabulkan atau dikabulkan sebagian"
        elif score >= 55:
            indication = "Berpeluang dikabulkan sebagian"
        elif score >= 35:
            indication = "Risiko sedang-tinggi, perlu penguatan bukti"
        else:
            indication = "Risiko tinggi ditolak atau tidak dapat diterima"

    top_share = max(distribution.values()) if distribution else 0
    confidence = "high" if similar_count >= 8 and top_share >= 0.55 else "medium" if similar_count >= 4 else "low"
    return indication, confidence


def compact_text(text: str, limit: int = 420) -> str:
    cleaned = normalize_spaces(text or "")
    if len(cleaned) <= limit:
        return cleaned
    return cleaned[:limit].rsplit(" ", 1)[0] + "..."


def build_top_case_analyses(
    intake: Dict[str, Any],
    similar_cases: List[Dict[str, Any]],
    limit: int = 2,
    language: str = "id",
) -> List[Dict[str, Any]]:
    language = normalize_report_language(language)
    analyses = []
    input_tax = normalize_spaces(str(intake.get("tax_type_filter") or intake.get("tax_type") or "")).lower()
    input_issue = normalize_spaces(str(intake.get("issue_type") or "")).lower()
    input_stage = normalize_spaces(str(intake.get("stage") or "")).lower()

    for idx, item in enumerate(similar_cases[:limit], start=1):
        doc = item["document"]
        outcome_text = outcome_label(doc.outcome, language)
        match_points = []
        if input_tax and doc.tax_type and input_tax == doc.tax_type.lower():
            match_points.append(f"same tax type ({doc.tax_type})" if language == "en" else f"jenis pajak sama ({doc.tax_type})")
        if input_issue and doc.issue_type and input_issue.replace(" ", "_") in doc.issue_type.lower():
            match_points.append(f"similar issue ({doc.issue_type})" if language == "en" else f"isu sejenis ({doc.issue_type})")
        if input_stage and input_stage in {"banding", "gugatan"}:
            match_points.append(
                "relevant to Tax Court litigation stage"
                if language == "en"
                else "relevan untuk tahap litigasi di Pengadilan Pajak"
            )
        if item.get("reasons"):
            match_points.extend(translate_match_reason(reason, language=language) for reason in item.get("reasons", [])[:2])
        if not match_points:
            match_points.append(
                "shares token and dispute-context similarity with the input"
                if language == "en"
                else "memiliki kemiripan token dan konteks sengketa dengan input"
            )

        reasoning = doc.court_reasoning_summary or doc.issue_summary or doc.wp_position_summary or doc.djp_position_summary
        if not reasoning:
            reasoning = (
                "The reasoning summary is not yet complete in the local database; open the PDF to validate details."
                if language == "en"
                else "Ringkasan pertimbangan belum lengkap di database lokal; perlu membuka PDF untuk validasi detail."
            )

        if language == "en" and doc.outcome in {"WP_FULL_WIN", "WP_PARTIAL_WIN"}:
            implication = (
                "This decision may support the taxpayer's position, especially where the evidentiary pattern "
                "and judicial reasoning align with the current case."
            )
        elif language == "en" and doc.outcome == "DJP_WIN":
            implication = (
                "This decision should be treated as a risk comparator. Use it to identify weaknesses in evidence "
                "or arguments that must be distinguished from the current case."
            )
        elif language == "en" and doc.outcome == "FORMAL_REJECTED":
            implication = (
                "This decision is mainly useful for testing formal risk, such as deadlines, authority of counsel, "
                "or administrative completeness."
            )
        elif language == "en":
            implication = (
                "This decision is useful as preliminary context, but its outcome and reasoning should be verified "
                "before relying on it as a key comparator."
            )
        elif doc.outcome in {"WP_FULL_WIN", "WP_PARTIAL_WIN"}:
            implication = (
                "Putusan ini dapat dipakai sebagai pembanding pendukung WP, terutama untuk menonjolkan pola bukti "
                "dan pertimbangan majelis yang menerima sebagian/seluruh posisi Wajib Pajak."
            )
        elif doc.outcome == "DJP_WIN":
            implication = (
                "Putusan ini penting sebagai pembanding risiko. Gunakan untuk mengidentifikasi kelemahan bukti atau "
                "argumentasi yang perlu dibedakan dari kasus WP saat ini."
            )
        elif doc.outcome == "FORMAL_REJECTED":
            implication = (
                "Putusan ini terutama relevan untuk menguji risiko formal seperti tenggat waktu, kewenangan kuasa, "
                "atau kelengkapan administrasi permohonan."
            )
        else:
            implication = (
                "Putusan ini relevan sebagai konteks awal, tetapi outcome atau pertimbangannya perlu diverifikasi lagi "
                "sebelum dipakai sebagai pembanding utama."
            )

        analyses.append(
            {
                "rank": idx,
                "putusan_number": doc.putusan_number,
                "filename": doc.filename,
                "tax_type": doc.tax_type,
                "issue_type": doc.issue_type,
                "outcome": doc.outcome,
                "outcome_label": outcome_text,
                "score": round(float(item.get("score") or 0), 1),
                "match_points": match_points[:4],
                "reasoning_summary": compact_text(reasoning, 620),
                "why_relevant": (
                    (
                        f"This decision is priority {idx} because {', '.join(match_points[:3])}. "
                        f"Its indicative relevance score is {round(float(item.get('score') or 0), 1)}."
                    )
                    if language == "en"
                    else (
                        f"Putusan ini masuk prioritas {idx} karena {', '.join(match_points[:3])}. "
                        f"Skor relevansi indikatifnya {round(float(item.get('score') or 0), 1)}."
                    )
                ),
                "case_implication": implication,
                "use_strategy": (
                    (
                        "Read the dispute issue and judicial reasoning sections, then cite only the portions with the closest facts. "
                        "Do not rely on the outcome alone without demonstrating evidentiary similarity."
                    )
                    if language == "en"
                    else (
                        "Baca bagian pokok sengketa dan pertimbangan majelis, lalu kutip hanya bagian yang faktanya "
                        "paling mirip. Hindari menjadikan outcome sebagai argumen tunggal tanpa menunjukkan kesamaan bukti."
                    )
                ),
            }
        )
    return analyses


def build_review(
    intake: Dict[str, Any],
    similar_cases: List[Dict[str, Any]],
    distribution: Dict[str, float],
    positive: List[str],
    negative: List[str],
    evidence_gaps: List[str],
    score: float,
    relevant_regulations: Optional[List[Dict[str, Any]]] = None,
    language: str = "id",
) -> str:
    language = normalize_report_language(language)
    top_cases = ", ".join(item["document"].putusan_number for item in similar_cases[:3]) or "belum ada"
    full = int(distribution.get("WP_FULL_WIN", 0) * 100)
    partial = int(distribution.get("WP_PARTIAL_WIN", 0) * 100)
    loss = int(distribution.get("DJP_WIN", 0) * 100)
    gaps = ", ".join(evidence_gaps) if evidence_gaps else "belum teridentifikasi signifikan"
    positives = "; ".join(positive[:4]) if positive else "belum ada faktor positif kuat dari input"
    negatives = "; ".join(negative[:4]) if negative else "belum ada faktor negatif kuat dari input"
    regulation_text = format_regulation_review_text(relevant_regulations or [])

    if language == "en":
        top_cases_en = ", ".join(item["document"].putusan_number for item in similar_cases[:3]) or "not yet available"
        gaps_en = ", ".join(translated_signal_title(item, language="en") for item in evidence_gaps) if evidence_gaps else "no significant gap identified yet"
        positives_en = "; ".join(translated_signal_title(item, language="en") for item in positive[:4]) if positive else "no strong positive factor has been identified from the input"
        negatives_en = "; ".join(translated_signal_title(item, language="en") for item in negative[:4]) if negative else "no strong negative factor has been identified from the input"
        regulation_en = format_regulation_review_text(relevant_regulations or [], language="en")
        return (
            f"The indicative score for this case is {score}/100. The score is driven mainly by comparable decisions, "
            f"evidence completeness, and issue similarity. The main comparable decisions to review are {top_cases_en}.\n\n"
            f"Based on weighted comparables, full taxpayer wins account for around {full}%, partial taxpayer wins around {partial}%, "
            f"and tax authority wins around {loss}%. These figures are not a prediction; they are an initial map for shaping the evidence strategy.\n\n"
            f"Current factors supporting the taxpayer: {positives_en}. Key risks to watch: {negatives_en}. "
            f"Evidence gaps to address first: {gaps_en}.\n\n"
            f"{regulation_en}\n\n"
            "Initial strategy: organize the chronology, reconcile the corrected amount, prepare key evidence, and use comparable decisions only where the judicial reasoning is factually close to the taxpayer's case."
        )

    return (
        f"Skor indikatif kasus ini adalah {score}/100. Angka ini terutama dipengaruhi oleh pola putusan pembanding, "
        f"kelengkapan bukti, dan kemiripan isu. Putusan pembanding utama yang perlu dibaca adalah {top_cases}.\n\n"
        f"Dari pembanding yang ditemukan, pola historis berbobot menunjukkan WP menang penuh sekitar {full}%, "
        f"menang sebagian sekitar {partial}%, dan DJP menang sekitar {loss}%. Angka ini bukan prediksi pasti, "
        "melainkan peta awal untuk menentukan strategi pembuktian.\n\n"
        f"Hal yang saat ini membantu posisi WP: {positives}. Hal yang perlu diwaspadai: {negatives}. "
        f"Celah bukti yang perlu dibereskan lebih dulu: {gaps}.\n\n"
        f"{regulation_text}\n\n"
        "Arah strategi awal: rapikan kronologi, pastikan angka koreksi dapat direkonsiliasi, siapkan bukti utama, "
        "lalu gunakan putusan pembanding hanya untuk bagian pertimbangan yang benar-benar mirip dengan fakta kasus WP."
    )


def build_recommendation(
    intake: Dict[str, Any],
    similar_cases: List[Dict[str, Any]],
    distribution: Dict[str, float],
    positive: List[str],
    negative: List[str],
    evidence_gaps: List[str],
    score: float,
    indication: str,
    relevant_regulations: Optional[List[Dict[str, Any]]] = None,
    language: str = "id",
) -> str:
    language = normalize_report_language(language)
    case_name = intake.get("taxpayer_name") or ("Taxpayer" if language == "en" else "Wajib Pajak")
    top_case_analyses = build_top_case_analyses(intake, similar_cases, limit=2, language=language)
    similar_lines = []
    for item in top_case_analyses:
        if language == "en":
            similar_lines.append(
                f"{item['rank']}. {item['putusan_number']} - {item['outcome_label']}; issue {item['issue_type']}; relevance {item['score']:.1f}."
            )
            similar_lines.append(f"   Relevance: {item['why_relevant']}")
            similar_lines.append(f"   Core reasoning: {item['reasoning_summary']}")
            similar_lines.append(f"   Implication for the taxpayer: {item['case_implication']}")
            similar_lines.append(f"   How to use it in the argument: {item['use_strategy']}")
        else:
            similar_lines.append(
                f"{item['rank']}. {item['putusan_number']} - {item['outcome_label']}; isu {item['issue_type']}; relevansi {item['score']:.1f}."
            )
            similar_lines.append(f"   Relevansi: {item['why_relevant']}")
            similar_lines.append(f"   Inti pertimbangan: {item['reasoning_summary']}")
            similar_lines.append(f"   Implikasi untuk WP: {item['case_implication']}")
            similar_lines.append(f"   Cara pakai dalam argumentasi: {item['use_strategy']}")
    if not similar_lines:
        similar_lines = [
            "No sufficiently relevant comparable decision was found in the local database."
            if language == "en"
            else "Belum ada putusan pembanding yang cukup relevan di database lokal."
        ]

    if language == "en":
        gaps = [translated_signal_title(item, language="en") for item in evidence_gaps] or ["Re-check the completeness of supporting documents."]
        positive_lines = [translated_signal_title(item, language="en") for item in positive] or ["The taxpayer narrative should be sharpened so the link between transaction facts, evidence, and legal basis is clear."]
        negative_lines = [translated_signal_title(item, language="en") for item in negative] or ["Risks cannot yet be fully mapped because the input remains limited."]
        regulation_lines = format_regulation_recommendation_lines(relevant_regulations or [], language="en")
        return "\n".join(
            [
                f"INITIAL RECOMMENDATION FOR {case_name}",
                "",
                "1. Case position summary",
                f"The case is at the {intake.get('stage', '-')} stage, tax type {intake.get('tax_type', '-')}, "
                f"with the main issue {intake.get('issue_type', '-')}. Input correction amount: {intake.get('correction_amount', '-')}.",
                f"Tax authority position: {intake.get('djp_reason', '-')}",
                f"Taxpayer position: {intake.get('wp_reason', '-')}",
                "",
                "2. Initial indication",
                f"{indication}. Indicative score: {score}/100. Confidence is based on the number and consistency of comparable decisions.",
                "",
                "3. VAT regulations to verify",
                *regulation_lines,
                "Use this list as an initial map. For the final argument, verify the relevant articles against the tax year, tax period, transaction type, and tax authority correction documents.",
                "",
                "4. Analysis of the key comparable decisions",
                *similar_lines,
                "At prototype stage, focus first on one or two most relevant decisions. Other comparables should remain as backup references so the analysis stays focused.",
                "",
                "5. Factors supporting the taxpayer and why they matter",
                *[f"- {item}" for item in positive_lines[:5]],
                "These supporting factors must be translated into an evidence narrative: which documents exist, which amounts reconcile, and how they answer the tax authority's correction.",
                "",
                "6. Risks and weaknesses to mitigate",
                *[f"- {item}" for item in negative_lines[:5]],
                "The main risk usually arises when evidence is inconsistent, incomplete, or does not directly answer the reason for correction. Each taxpayer rebuttal should therefore be tied to specific supporting documents.",
                "",
                "7. Evidence to complete and its evidentiary purpose",
                *[f"- {item}" for item in gaps],
                "",
                "8. Recommended next steps",
                "- Complete key evidence and reconcile amounts before submitting dispute documents.",
                "- Structure the argument by issue: tax authority position, taxpayer rebuttal, evidence, VAT legal basis, then comparable decisions.",
                "- Use comparable decisions only where the issue, evidence, and legal basis are genuinely similar.",
                "- Include a disclaimer that this is indicative and remains subject to review by tax/legal counsel.",
                "",
                "9. Main narrative draft",
                "The taxpayer should build the narrative that the correction is not appropriate because the transaction facts, reporting position, and supporting evidence indicate a stronger taxpayer position. This narrative should be supported by amount reconciliation and comparison with decisions that have similar reasoning patterns.",
            ]
        )

    gaps = evidence_gaps or ["Validasi ulang kelengkapan dokumen pendukung."]
    positive_lines = positive or ["Narasi WP perlu dipertajam agar hubungan transaksi, bukti, dan dasar hukum terlihat jelas."]
    negative_lines = negative or ["Risiko belum dapat dipetakan penuh karena data input masih terbatas."]
    regulation_lines = format_regulation_recommendation_lines(relevant_regulations or [], language="id")

    return "\n".join(
        [
            f"REKOMENDASI AWAL UNTUK {case_name}",
            "",
            "1. Ringkasan posisi perkara",
            f"Kasus berada pada tahap {intake.get('stage', '-')}, jenis pajak {intake.get('tax_type', '-')}, "
            f"dengan isu utama {intake.get('issue_type', '-')}. Nilai koreksi yang diinput: {intake.get('correction_amount', '-')}.",
            f"Posisi DJP: {intake.get('djp_reason', '-')}",
            f"Posisi WP: {intake.get('wp_reason', '-')}",
            "",
            "2. Indikasi awal",
            f"{indication}. Skor indikatif: {score}/100. Confidence: berbasis jumlah dan konsistensi putusan pembanding.",
            "",
            "3. Dasar peraturan PPN yang perlu diuji",
            *regulation_lines,
            "Gunakan daftar ini sebagai peta awal. Untuk argumentasi final, cocokkan kembali pasal yang relevan dengan tahun pajak, masa pajak, jenis transaksi, dan dokumen koreksi DJP.",
            "",
            "4. Analisis putusan pembanding utama",
            *similar_lines,
            "Untuk tahap prototype, gunakan satu sampai dua putusan paling relevan terlebih dahulu. Pembanding lain cukup menjadi daftar cadangan agar analisis tidak melebar.",
            "",
            "5. Faktor yang mendukung WP dan alasannya",
            *[f"- {item}" for item in positive_lines[:5]],
            "Faktor pendukung tersebut harus diterjemahkan ke dalam narasi bukti: dokumen apa yang ada, angka mana yang direkonsiliasi, dan bagaimana dokumen itu menjawab koreksi DJP.",
            "",
            "6. Risiko dan kelemahan yang perlu dimitigasi",
            *[f"- {item}" for item in negative_lines[:5]],
            "Risiko terbesar biasanya muncul saat bukti tidak konsisten, tidak lengkap, atau tidak langsung menjawab alasan koreksi. Karena itu, setiap bantahan WP perlu diikat ke dokumen pendukung yang spesifik.",
            "",
            "7. Bukti yang perlu dilengkapi dan tujuan pembuktiannya",
            *[f"- {item}" for item in gaps],
            "",
            "8. Rekomendasi langkah berikutnya",
            "- Lengkapi bukti utama dan rekonsiliasi angka sebelum menyampaikan dokumen sengketa.",
            "- Susun argumentasi per pokok sengketa: posisi DJP, bantahan WP, bukti, dasar hukum PPN, lalu putusan pembanding.",
            "- Gunakan putusan pembanding hanya jika isu, bukti, dan dasar hukumnya benar-benar mirip.",
            "- Cantumkan disclaimer bahwa analisis ini indikatif dan tetap perlu review konsultan pajak/kuasa hukum.",
            "",
            "9. Draft narasi utama",
            "WP sebaiknya membangun narasi bahwa koreksi DJP tidak tepat karena fakta transaksi, pelaporan, dan bukti pendukung menunjukkan posisi WP lebih sesuai. Narasi ini perlu diperkuat dengan rekonsiliasi angka dan pembandingan terhadap putusan yang memiliki pola pertimbangan serupa.",
        ]
    )


def build_llm_prompt(intake: Dict[str, Any], similar_cases: List[Dict[str, Any]], local_result: Dict[str, Any]) -> str:
    language = normalize_report_language(intake.get("report_language") or local_result.get("report_language"))
    language_instruction = (
        "Write the report in professional English for a taxpayer/client audience."
        if language == "en"
        else "Tulis dalam bahasa Indonesia profesional, praktis, dan mudah dipahami WP."
    )
    structure_instruction = (
        "Required structure: Executive Summary, Case Position, Relevant VAT Regulations, Comparable Decision Analysis, Supporting Factors, Risks and Mitigation, Evidence to Complete, Argument Strategy, Recommended Next Steps, Disclaimer."
        if language == "en"
        else "Struktur wajib: Ringkasan Eksekutif, Posisi Perkara, Dasar Peraturan PPN yang Relevan, Analisis Putusan Pembanding, Faktor Pendukung, Risiko dan Mitigasi, Bukti yang Harus Dilengkapi, Strategi Argumentasi, Rekomendasi Langkah Berikutnya, Disclaimer."
    )
    return_instruction = "Return plain Markdown, not JSON." if language == "en" else "Return dalam Markdown biasa, bukan JSON."
    comparable_instruction = (
        "In the Comparable Decision Analysis section, select at most 1-2 most relevant decisions and explain in depth why each is relevant, the core reasoning, its impact on the taxpayer position, and how it should be used in the argument."
        if language == "en"
        else "Pada bagian Analisis Putusan Pembanding, pilih maksimal 1-2 putusan paling relevan dan jelaskan secara mendalam mengapa relevan, inti pertimbangannya, dampaknya untuk posisi WP, serta cara memakainya dalam argumen."
    )
    table_instruction = (
        "If you create Markdown tables, keep them complete and simple."
        if language == "en"
        else "Jika membuat tabel Markdown, pastikan format tabel lengkap dan sederhana."
    )
    cases = []
    for item in similar_cases[:4]:
        doc = item["document"]
        cases.append(
            {
                "putusan_number": doc.putusan_number,
                "tax_type": doc.tax_type,
                "issue_type": doc.issue_type,
                "outcome": doc.outcome,
                "similarity_score": round(float(item.get("score") or 0), 1),
                "issue_summary": doc.issue_summary[:700],
                "court_reasoning_summary": doc.court_reasoning_summary[:700],
            }
        )
    regulations = local_result.get("relevant_regulations", [])[:6]
    return (
        "Anda adalah analis sengketa pajak Indonesia. Buat satu draft rekomendasi yang mendalam "
        "untuk Wajib Pajak berdasarkan data input, putusan pembanding, dan peraturan pajak terkait berikut. "
        "Jangan mengarang nomor putusan, pasal, nama pihak, atau fakta baru. Jika menyebut peraturan, pakai hanya yang tersedia di daftar konteks dan tetap beri arahan verifikasi pasal. Gunakan framing indikatif, bukan kepastian. "
        f"{language_instruction} {return_instruction}\n\n"
        f"INPUT KASUS:\n{json.dumps(intake, ensure_ascii=False, indent=2)}\n\n"
        f"PUTUSAN PEMBANDING:\n{json.dumps(cases, ensure_ascii=False, indent=2)}\n\n"
        f"PERATURAN PAJAK TERKAIT:\n{json.dumps(regulations, ensure_ascii=False, indent=2)}\n\n"
        f"HASIL SCORING LOKAL:\n{json.dumps(local_result, ensure_ascii=False, indent=2)[:6000]}\n\n"
        f"{structure_instruction} "
        f"{comparable_instruction} "
        f"{table_instruction}"
    )


def call_openai_llm(prompt: str) -> Optional[str]:
    result = call_openai_text(prompt, max_output_tokens=1800)
    text = result.get("text")
    if text and result.get("model"):
        return text
    return text


def call_openai_text(prompt: str, max_output_tokens: int = 1800) -> Dict[str, Optional[str]]:
    load_local_env()
    if not os.environ.get("OPENAI_API_KEY"):
        return {"text": None, "model": None, "error": "OPENAI_API_KEY belum diisi."}
    try:
        from openai import OpenAI

        client = OpenAI()
        errors = []
        for model in model_candidates():
            params = {
                "model": model,
                "input": prompt,
                "max_output_tokens": max_output_tokens,
            }
            if model.startswith("gpt-5"):
                params["reasoning"] = {"effort": os.environ.get("TDP_REASONING_EFFORT", "low")}
                params["text"] = {"verbosity": os.environ.get("TDP_TEXT_VERBOSITY", "low")}
            else:
                params["temperature"] = 0.2
            try:
                response = client.responses.create(**params)
                return {"text": getattr(response, "output_text", None), "model": model, "error": None}
            except Exception as exc:
                errors.append(f"{model}: {exc}")
        return {"text": None, "model": None, "error": " | ".join(errors)}
    except Exception as exc:
        return {"text": None, "model": None, "error": f"LLM call failed: {exc}"}


def call_openai_vision(prompt: str, image_data_urls: Sequence[str], max_output_tokens: int = 1800) -> Dict[str, Optional[str]]:
    load_local_env()
    if not os.environ.get("OPENAI_API_KEY"):
        return {"text": None, "model": None, "error": "OPENAI_API_KEY belum diisi."}
    if not image_data_urls:
        return {"text": None, "model": None, "error": "Tidak ada gambar halaman PDF untuk dikirim ke LLM."}
    try:
        from openai import OpenAI

        client = OpenAI()
        content = [{"type": "input_text", "text": prompt}]
        content.extend({"type": "input_image", "image_url": image_url} for image_url in image_data_urls)
        errors = []
        for model in model_candidates():
            params = {
                "model": model,
                "input": [{"role": "user", "content": content}],
                "max_output_tokens": max_output_tokens,
            }
            if model.startswith("gpt-5"):
                params["reasoning"] = {"effort": os.environ.get("TDP_REASONING_EFFORT", "low")}
                params["text"] = {"verbosity": os.environ.get("TDP_TEXT_VERBOSITY", "low")}
            else:
                params["temperature"] = 0.1
            try:
                response = client.responses.create(**params)
                return {"text": getattr(response, "output_text", None), "model": model, "error": None}
            except Exception as exc:
                errors.append(f"{model}: {exc}")
        return {"text": None, "model": None, "error": " | ".join(errors)}
    except Exception as exc:
        return {"text": None, "model": None, "error": f"LLM vision call failed: {exc}"}


def model_candidates() -> List[str]:
    primary = os.environ.get("TDP_LLM_MODEL", "gpt-5.5")
    fallback_raw = os.environ.get("TDP_LLM_FALLBACK_MODELS", "gpt-5.4-mini,gpt-5-mini,gpt-4o-mini")
    candidates = [primary] + [item.strip() for item in fallback_raw.split(",") if item.strip()]
    return dedupe(candidates)


def save_report(
    intake: Dict[str, Any],
    similar_cases: List[Dict[str, Any]],
    result: Dict[str, Any],
    db_path: Path = DEFAULT_DB_PATH,
) -> str:
    init_db(db_path)
    report_id = str(uuid.uuid4())
    slim_cases = []
    for item in similar_cases:
        doc = item["document"]
        slim_cases.append(
            {
                "putusan_number": doc.putusan_number,
                "filename": doc.filename,
                "tax_type": doc.tax_type,
                "issue_type": doc.issue_type,
                "outcome": doc.outcome,
                "score": item.get("score"),
                "reasons": item.get("reasons", []),
            }
        )
    with connect(db_path) as conn:
        conn.execute(
            """
            INSERT INTO analysis_reports (report_id, created_at, input_json, similar_cases_json, result_json)
            VALUES (?, ?, ?, ?, ?)
            """,
            (
                report_id,
                now_iso(),
                json.dumps(intake, ensure_ascii=False),
                json.dumps(slim_cases, ensure_ascii=False),
                json.dumps(result, ensure_ascii=False),
            ),
        )
    return report_id


def list_reports(limit: int = 20, db_path: Path = DEFAULT_DB_PATH) -> List[Dict[str, Any]]:
    init_db(db_path)
    with connect(db_path) as conn:
        rows = conn.execute(
            "SELECT * FROM analysis_reports ORDER BY created_at DESC LIMIT ?",
            (limit,),
        ).fetchall()
    return [
        {
            "report_id": row["report_id"],
            "created_at": row["created_at"],
            "input": json.loads(row["input_json"]),
            "similar_cases": json.loads(row["similar_cases_json"]),
            "result": json.loads(row["result_json"]),
        }
        for row in rows
    ]


def get_report(report_id: str, db_path: Path = DEFAULT_DB_PATH) -> Optional[Dict[str, Any]]:
    init_db(db_path)
    with connect(db_path) as conn:
        row = conn.execute("SELECT * FROM analysis_reports WHERE report_id = ?", (report_id,)).fetchone()
    if not row:
        return None
    return {
        "report_id": row["report_id"],
        "created_at": row["created_at"],
        "input": json.loads(row["input_json"]),
        "similar_cases": json.loads(row["similar_cases_json"]),
        "result": json.loads(row["result_json"]),
    }


def export_report_markdown(report: Dict[str, Any]) -> str:
    result = report["result"]
    intake = report["input"]
    language = normalize_report_language(intake.get("report_language") or result.get("report_language"))
    similar_cases = report.get("similar_cases", [])
    top_case_analyses = result.get("top_case_analysis") or []
    supporting = result.get("supporting_factor_analysis", [])
    risks = result.get("risk_factor_analysis", [])
    gaps = result.get("evidence_gap_analysis", [])
    regulations = result.get("regulation_analysis", [])
    h = {
        "indication": "Indication" if language == "en" else "Indikasi",
        "input_position": "Input Position" if language == "en" else "Input Posisi",
        "tax_authority": "Tax authority position" if language == "en" else "Alasan DJP",
        "taxpayer": "Taxpayer position" if language == "en" else "Alasan WP",
        "evidence": "Available evidence" if language == "en" else "Bukti tersedia",
        "risk_review": "Risk Review" if language == "en" else "Review Risiko",
        "supporting": "Supporting Factors" if language == "en" else "Faktor Pendukung",
        "risks": "Risk Factors" if language == "en" else "Faktor Risiko",
        "gaps": "Evidence Gaps" if language == "en" else "Celah Bukti",
        "regulations": "Related Regulations" if language == "en" else "Peraturan Terkait",
        "top_cases": "Most Relevant Decisions" if language == "en" else "Putusan Paling Terkait",
        "other_cases": "Comparable Decision List" if language == "en" else "Daftar Putusan Pembanding",
        "recommendation": "Recommendation Draft" if language == "en" else "Draft Rekomendasi",
        "disclaimer": "Disclaimer" if language == "en" else "Disclaimer",
        "llm_note": "LLM Note" if language == "en" else "Catatan LLM",
        "outcome": "Outcome" if language == "en" else "Outcome",
        "relevance": "Relevance" if language == "en" else "Relevansi",
        "reasoning": "Core reasoning" if language == "en" else "Inti pertimbangan",
        "implication": "Taxpayer implication" if language == "en" else "Implikasi untuk WP",
        "strategy": "How to use" if language == "en" else "Cara pakai",
    }

    def analysis_line(item: Dict[str, Any], translate_title: bool = True) -> str:
        raw_title = item.get("title", "-")
        title = translated_signal_title(raw_title, language=language) if translate_title else str(raw_title)
        analysis = translate_legacy_english_text(item.get("analysis", "-"), language=language)
        return f"- {title}: {analysis}"

    lines = [
        "# RSM Tax Dispute Analysis Report",
        "",
        "Prepared with RSM Tax Dispute Simple Advisor prototype.",
        "",
        f"- Report ID: {report['report_id']}",
        f"- Created at: {report['created_at']}",
        f"- WP: {intake.get('taxpayer_name', '-')}",
        f"- NPWP: {intake.get('taxpayer_npwp', '-') or '-'}",
        f"- Kuasa/Konsultan: {intake.get('legal_counsel_name', '-') or '-'}",
        f"- Unit DJP: {intake.get('djp_unit', '-') or '-'}",
        f"- Jenis pajak: {intake.get('tax_type', '-')}",
        f"- Isu: {intake.get('issue_type', '-')}",
        f"- Tahap: {intake.get('stage', '-')}",
        f"- Nilai koreksi: {intake.get('correction_amount', '-')}",
        "",
        f"## {h['indication']}",
        f"- Score: {result.get('score')}",
        f"- {h['indication']}: {result.get('indication')}",
        f"- Confidence: {result.get('confidence')}",
        f"- Evidence score: {result.get('evidence_score')}",
        "",
        f"## {h['input_position']}",
        f"**{h['tax_authority']}:** {intake.get('djp_reason', '-')}",
        "",
        f"**{h['taxpayer']}:** {intake.get('wp_reason', '-')}",
        "",
        f"**{h['evidence']}:** {', '.join(intake.get('available_evidence') or []) or '-'}",
        "",
        f"## {h['risk_review']}",
        translate_legacy_english_text(result.get("review", ""), language=language),
        "",
        f"## {h['supporting']}",
        *[analysis_line(item) for item in supporting],
        "",
        f"## {h['risks']}",
        *[analysis_line(item) for item in risks],
        "",
        f"## {h['gaps']}",
        *[analysis_line(item) for item in gaps],
        "",
        f"## {h['regulations']}",
        *[analysis_line(item, translate_title=False) for item in regulations],
        *[
            f"- {reg.get('title') or reg.get('number')} ({reg.get('url', '-')})"
            for reg in result.get("relevant_regulations", [])[:8]
        ],
        "",
        f"## {h['top_cases']}",
        *[
            "\n".join(
                [
                    f"### {item.get('rank', idx)}. {item.get('putusan_number', '-')}",
                    f"- {h['outcome']}: {item.get('outcome_label', item.get('outcome', '-'))}",
                    f"- {h['relevance']}: {item.get('why_relevant', '-')}",
                    f"- {h['reasoning']}: {item.get('reasoning_summary', '-')}",
                    f"- {h['implication']}: {item.get('case_implication', '-')}",
                    f"- {h['strategy']}: {item.get('use_strategy', '-')}",
                ]
            )
            for idx, item in enumerate(top_case_analyses[:2], start=1)
        ],
        "",
        f"## {h['other_cases']}",
        *[
            f"- {case.get('putusan_number', '-')} | {case.get('tax_type', '-')} | {case.get('issue_type', '-')} | {case.get('outcome', '-')} | skor {case.get('score', '-')}"
            for case in similar_cases[:5]
        ],
        "",
        f"## {h['recommendation']}",
        translate_legacy_english_text(result.get("recommendation_draft", ""), language=language),
        "",
        f"## {h['disclaimer']}",
        result.get("disclaimer", ""),
    ]
    if result.get("llm_used") and result.get("llm_note"):
        lines.extend(["", f"## {h['llm_note']}", result["llm_note"]])
    return "\n".join(lines)


def _docx_rgb(hex_color: str):
    from docx.shared import RGBColor

    value = hex_color.strip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def _shade_docx_cell(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _remove_docx_cell_borders(cell) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tc_pr.append(borders)


def _set_docx_margins(document) -> None:
    from docx.shared import Inches

    for section in document.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.72)
        section.right_margin = Inches(0.72)


def _add_rsm_docx_identity(document, report: Dict[str, Any]) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    _set_docx_margins(document)

    footer = document.sections[0].footer.paragraphs[0]
    footer.text = "RSM Tax Dispute Simple Advisor | Prototype analysis, subject to professional review"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = _docx_rgb(RSM_MID_GRAY_HEX)

    strip = document.add_table(rows=1, cols=3)
    strip.alignment = WD_TABLE_ALIGNMENT.LEFT
    for cell, fill, width in zip(
        strip.rows[0].cells,
        (RSM_MID_GRAY_HEX, RSM_GREEN_HEX, RSM_BLUE_HEX),
        (Inches(0.22), Inches(0.68), Inches(2.2)),
    ):
        cell.width = width
        cell.text = ""
        _shade_docx_cell(cell, fill)
        _remove_docx_cell_borders(cell)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    logo = document.add_paragraph()
    logo.paragraph_format.space_before = Pt(4)
    logo.paragraph_format.space_after = Pt(0)
    logo_run = logo.add_run("RSM")
    logo_run.bold = True
    logo_run.font.name = "Arial"
    logo_run.font.size = Pt(32)
    logo_run.font.color.rgb = _docx_rgb(RSM_GRAY_HEX)

    tagline = document.add_paragraph()
    tagline.paragraph_format.space_before = Pt(0)
    tagline.paragraph_format.space_after = Pt(10)
    tag_run = tagline.add_run("Tax Dispute Simple Advisor | Analysis Report")
    tag_run.font.name = "Arial"
    tag_run.font.size = Pt(9)
    tag_run.font.color.rgb = _docx_rgb(RSM_BLUE_HEX)

    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    meta_run = meta.add_run(f"Report ID: {report.get('report_id', '-')} | Generated: {report.get('created_at', '-')}")
    meta_run.font.name = "Arial"
    meta_run.font.size = Pt(8)
    meta_run.font.color.rgb = _docx_rgb(RSM_MID_GRAY_HEX)


def _clean_inline_markdown(text: str) -> str:
    cleaned = re.sub(r"`([^`]+)`", r"\1", text or "")
    cleaned = cleaned.replace("\\|", "|")
    cleaned = re.sub(r"\*\*([^*]+)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"__([^_]+)__", r"\1", cleaned)
    cleaned = re.sub(r"\*([^*]+)\*", r"\1", cleaned)
    cleaned = re.sub(r"_([^_]+)_", r"\1", cleaned)
    return cleaned.strip()


def _add_markdown_runs(paragraph, text: str) -> None:
    cursor = 0
    pattern = re.compile(r"(\*\*([^*]+)\*\*|__([^_]+)__|`([^`]+)`)")
    for match in pattern.finditer(text or ""):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        content = match.group(2) or match.group(3) or match.group(4) or ""
        run = paragraph.add_run(content)
        if match.group(2) or match.group(3):
            run.bold = True
        if match.group(4):
            run.font.name = "Courier New"
        cursor = match.end()
    if cursor < len(text or ""):
        paragraph.add_run((text or "")[cursor:])


def _is_markdown_table_separator(line: str) -> bool:
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell or "") for cell in cells)


def _parse_markdown_table_row(line: str) -> List[str]:
    return [_clean_inline_markdown(cell) for cell in line.strip().strip("|").split("|")]


def _read_markdown_table(lines: List[str], start_idx: int) -> Tuple[Optional[List[List[str]]], int]:
    if start_idx + 1 >= len(lines):
        return None, start_idx
    header = lines[start_idx].strip()
    separator = lines[start_idx + 1].strip()
    if "|" not in header or not _is_markdown_table_separator(separator):
        return None, start_idx
    rows = [_parse_markdown_table_row(header)]
    idx = start_idx + 2
    while idx < len(lines) and "|" in lines[idx].strip() and lines[idx].strip():
        rows.append(_parse_markdown_table_row(lines[idx]))
        idx += 1
    width = len(rows[0])
    rows = [row + [""] * (width - len(row)) if len(row) < width else row[:width] for row in rows]
    return rows, idx


def _add_docx_markdown_table(document, rows: List[List[str]]) -> None:
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.shared import Pt

    if not rows:
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.add_run(value)
            run.font.name = "Arial"
            run.font.size = Pt(9)
            if row_idx == 0:
                _shade_docx_cell(cell, RSM_BLUE_HEX)
                run.bold = True
                run.font.color.rgb = _docx_rgb("FFFFFF")
            else:
                run.font.color.rgb = _docx_rgb(RSM_GRAY_HEX)
                if row_idx % 2 == 0:
                    _shade_docx_cell(cell, RSM_SOFT_GRAY_HEX)
    document.add_paragraph()


def export_report_docx(report: Dict[str, Any]) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as exc:
        raise RuntimeError("Paket python-docx belum tersedia untuk membuat Word.") from exc

    markdown = export_report_markdown(report)
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for style_name, color in [("Heading 1", RSM_BLUE_HEX), ("Heading 2", RSM_GRAY_HEX), ("Heading 3", RSM_GRAY_HEX)]:
        if style_name in styles:
            styles[style_name].font.name = "Arial"
            styles[style_name].font.color.rgb = _docx_rgb(color)

    _add_rsm_docx_identity(document, report)

    lines = markdown.splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        if not stripped:
            document.add_paragraph()
            idx += 1
            continue
        table_rows, next_idx = _read_markdown_table(lines, idx)
        if table_rows:
            _add_docx_markdown_table(document, table_rows)
            idx = next_idx
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            level = min(max(len(heading.group(1)) - 1, 0), 3)
            document.add_heading(_clean_inline_markdown(heading.group(2)), level=level)
        elif stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_markdown_runs(paragraph, stripped[2:])
        elif re.match(r"^\d+\.\s+", stripped):
            paragraph = document.add_paragraph(style="List Number")
            _add_markdown_runs(paragraph, re.sub(r"^\d+\.\s+", "", stripped))
        elif stripped.startswith("**") and ":**" in stripped:
            title, body = stripped.split(":**", 1)
            paragraph = document.add_paragraph()
            paragraph.add_run(title.strip("*") + ":").bold = True
            _add_markdown_runs(paragraph, body.strip())
        else:
            paragraph = document.add_paragraph()
            _add_markdown_runs(paragraph, stripped)
        idx += 1

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def find_soffice() -> Optional[str]:
    for name in ("soffice", "libreoffice"):
        found = _which(name)
        if found:
            return found
    candidates = [
        Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),
        Path("/Applications/LibreOffice.app/Contents/MacOS/libreoffice"),
    ]
    for candidate in candidates:
        if candidate.exists() and os.access(candidate, os.X_OK):
            return str(candidate)
    return None


def export_report_pdf(report: Dict[str, Any]) -> bytes:
    docx_bytes = export_report_docx(report)
    soffice = find_soffice()
    if soffice:
        with tempfile.TemporaryDirectory() as tmpdir:
            tmp_path = Path(tmpdir)
            docx_path = tmp_path / "report.docx"
            docx_path.write_bytes(docx_bytes)
            result = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(tmp_path), str(docx_path)],
                check=False,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                text=True,
                timeout=90,
            )
            pdf_path = tmp_path / "report.pdf"
            if result.returncode == 0 and pdf_path.exists():
                return pdf_path.read_bytes()
    return export_report_basic_pdf(report)


def export_report_basic_pdf(report: Dict[str, Any]) -> bytes:
    text = export_report_markdown(report)
    wrapped_lines: List[str] = []
    raw_lines: List[str] = []
    markdown_lines = text.splitlines()
    idx = 0
    while idx < len(markdown_lines):
        table_rows, next_idx = _read_markdown_table(markdown_lines, idx)
        if table_rows:
            raw_lines.extend("   |   ".join(row) for row in table_rows)
            idx = next_idx
            continue
        raw_lines.append(markdown_lines[idx])
        idx += 1

    for raw_line in raw_lines:
        prefix = ""
        line = raw_line.strip()
        if line.startswith("#"):
            line = line.lstrip("#").strip().upper()
        if line.startswith("- "):
            prefix = "- "
            line = line[2:]
        line = _clean_inline_markdown(line)
        chunks = textwrap.wrap(line, width=92) or [""]
        for idx, chunk in enumerate(chunks):
            wrapped_lines.append((prefix if idx == 0 else "  ") + chunk)

    lines_per_page = 48
    pages = [wrapped_lines[idx : idx + lines_per_page] for idx in range(0, len(wrapped_lines), lines_per_page)] or [[""]]
    objects: List[bytes] = []

    def pdf_rgb(hex_color: str) -> str:
        value = hex_color.strip("#")
        channels = [int(value[idx : idx + 2], 16) / 255 for idx in (0, 2, 4)]
        return " ".join(f"{channel:.3f}" for channel in channels)

    def add_object(body: str) -> int:
        objects.append(body.encode("latin-1", errors="replace"))
        return len(objects)

    catalog_id = add_object("<< /Type /Catalog /Pages 2 0 R >>")
    pages_id = add_object("PAGES_PLACEHOLDER")
    font_id = add_object("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    page_ids = []
    for page_lines in pages:
        commands = [
            f"{pdf_rgb(RSM_MID_GRAY_HEX)} rg",
            "50 814 12 10 re f",
            f"{pdf_rgb(RSM_GREEN_HEX)} rg",
            "70 814 40 10 re f",
            f"{pdf_rgb(RSM_BLUE_HEX)} rg",
            "118 814 128 10 re f",
            f"{pdf_rgb(RSM_GRAY_HEX)} rg",
            "BT",
            "/F1 20 Tf",
            "50 788 Td",
            "(RSM) Tj",
            "ET",
            f"{pdf_rgb(RSM_BLUE_HEX)} rg",
            "BT",
            "/F1 9 Tf",
            "114 792 Td",
            "(Tax Dispute Simple Advisor | Analysis Report) Tj",
            "ET",
            "0 0 0 rg",
            "BT",
            "/F1 10 Tf",
            "50 760 Td",
            "14 TL",
        ]
        for line in page_lines:
            commands.append(f"({_pdf_escape(line)}) Tj")
            commands.append("T*")
        commands.append("ET")
        stream = "\n".join(commands)
        stream_id = add_object(f"<< /Length {len(stream.encode('latin-1', errors='replace'))} >>\nstream\n{stream}\nendstream")
        page_id = add_object(
            f"<< /Type /Page /Parent {pages_id} 0 R /MediaBox [0 0 612 842] "
            f"/Resources << /Font << /F1 {font_id} 0 R >> >> /Contents {stream_id} 0 R >>"
        )
        page_ids.append(page_id)
    kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
    objects[pages_id - 1] = f"<< /Type /Pages /Kids [{kids}] /Count {len(page_ids)} >>".encode("latin-1")

    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for idx, body in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{idx} 0 obj\n".encode("ascii"))
        output.extend(body)
        output.extend(b"\nendobj\n")
    xref_start = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    output.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root {catalog_id} 0 R >>\nstartxref\n{xref_start}\n%%EOF\n".encode("ascii")
    )
    return bytes(output)


def _pdf_escape(text: str) -> str:
    return text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
