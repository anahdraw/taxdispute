from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "LLM_ARCHITECTURE.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "222222"
MUTED = "666666"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
LIGHT_GREEN = "EAF7EA"
LIGHT_YELLOW = "FFF4CC"
BORDER = "B7C3D0"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, size=10, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_fixed(table, widths_inches):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.allow_autofit = False
    total_twips = sum(Inches(width).twips for width in widths_inches)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width_emu = Inches(widths_inches[idx])
            width_twips = Inches(widths_inches[idx]).twips
            cell.width = width_emu
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width_twips))
            tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = tr_pr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor.from_string(BLUE if level <= 2 else DARK_BLUE)
    return paragraph


def add_body(document, text, bold_prefix=None):
    paragraph = document.add_paragraph()
    paragraph.style = document.styles["Normal"]
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    return paragraph


def add_numbered(document, text):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    return paragraph


def add_callout(document, label, text, fill=LIGHT_YELLOW):
    table = document.add_table(rows=1, cols=1)
    set_table_fixed(table, [6.5])
    set_table_borders(table, color="D5C07A", size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(label + ": ")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    body = paragraph.add_run(text)
    body.font.name = "Calibri"
    body.font.size = Pt(10.5)
    document.add_paragraph()


def add_vertical_flow(document):
    add_heading(document, "Diagram 1 - Flow Utama LLM", level=1)
    steps = [
        ("1", "Input", "Upload PDF putusan/surat banding/dokumen pendukung, isi form kasus WP, atau refresh peraturan PPN dari Ortax."),
        ("2", "Document Intake", "File disimpan, text layer dibaca dengan pypdf/pdftotext, dan tipe dokumen dicatat."),
        ("3A", "LLM Text Extraction", "Jika teks cukup, prompt ekstraksi membaca header, pokok sengketa, posisi WP/DJP, pertimbangan, dan amar."),
        ("3B", "LLM Vision Extraction", "Jika PDF scan, halaman awal/tengah/akhir dirender menjadi gambar lalu dibaca model vision."),
        ("4", "Structured JSON", "LLM mengembalikan JSON field ekstraksi: metadata putusan, objek sengketa, pihak, argumen, pertimbangan, outcome."),
        ("5", "Validation & Storage", "Field dinormalisasi, nomor putusan dicek duplikat, lalu disimpan ke SQLite."),
        ("6", "Retrieval", "Sistem mencari putusan pembanding lokal dan peraturan PPN terkait dari database."),
        ("7", "Scoring & Risk", "Engine menghitung similarity, evidence score, distribusi outcome, faktor pendukung, faktor risiko, dan celah bukti."),
        ("8", "LLM Synthesis", "Jika diaktifkan, LLM menyusun draft rekomendasi mendalam berdasarkan input, pembanding, peraturan, dan skor."),
        ("9", "Final Output", "Hasil tampil di UI, tersimpan sebagai report, lalu dapat diunduh sebagai Word, PDF, atau Markdown."),
    ]
    for block_idx, block_steps in enumerate((steps[:6], steps[6:])):
        if block_idx == 1:
            document.add_page_break()
            add_heading(document, "Diagram 1 - Flow Utama LLM (lanjutan)", level=1)
        table = document.add_table(rows=len(block_steps), cols=3)
        set_table_fixed(table, [0.55, 1.7, 4.25])
        set_table_borders(table)
        for idx, (no, stage, detail) in enumerate(block_steps):
            set_row_cant_split(table.rows[idx])
            cells = table.rows[idx].cells
            set_cell_text(cells[0], no, bold=True, size=10, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(cells[1], stage, bold=True, size=10, color=DARK_BLUE)
            set_cell_text(cells[2], detail, size=9.5)
            shade = LIGHT_BLUE if idx % 2 == 0 else "FFFFFF"
            for cell in cells:
                set_cell_shading(cell, shade)
        document.add_paragraph()


def add_sequence_table(document):
    add_heading(document, "Diagram 2 - Sequence Proses LLM", level=1)
    rows = [
        ("1", "User -> Streamlit", "Upload dokumen atau pilih dokumen terindeks", "PDF / parameter kasus"),
        ("2", "Streamlit -> Extraction Engine", "Panggil upsert_document()", "Dokumen masuk pipeline"),
        ("3", "Extraction Engine", "Cek apakah PDF punya text layer cukup", "Cabang text atau vision"),
        ("4A", "Extraction -> OpenAI", "Prompt ekstraksi berbasis teks", "JSON label dokumen"),
        ("4B", "Extraction -> PDF Render -> OpenAI", "Render halaman scan lalu prompt vision", "JSON label dokumen"),
        ("5", "Extraction -> SQLite", "Normalisasi field dan simpan hasil", "documents, document_extractions, llm_labels, chunks"),
        ("6", "User -> Analysis", "Jalankan analisis kasus WP", "Intake final"),
        ("7", "Retrieval -> SQLite", "Ambil putusan pembanding dan peraturan PPN", "Konteks analisis"),
        ("8", "Analysis Engine", "Hitung skor, confidence, evidence, dan risiko", "Review risiko"),
        ("9", "Analysis -> OpenAI", "Opsional: sintesis draft rekomendasi mendalam", "Draft rekomendasi WP"),
        ("10", "Report Exporter", "Simpan report dan buat dokumen unduhan", "DOCX / PDF / Markdown"),
    ]
    table = document.add_table(rows=len(rows) + 1, cols=4)
    set_table_fixed(table, [0.5, 1.55, 2.65, 1.8])
    set_table_borders(table)
    headers = ["No.", "Aktor", "Aksi", "Output"]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_text(cell, header, bold=True, size=9.5, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, LIGHT_GRAY)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for row_idx, row in enumerate(rows, start=1):
        set_row_cant_split(table.rows[row_idx])
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            set_cell_text(cell, value, bold=col_idx == 0, size=9.2, align=WD_ALIGN_PARAGRAPH.CENTER if col_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_shading(cell, "FFFFFF" if row_idx % 2 else "F9FBFD")
    document.add_paragraph()


def add_component_table(document):
    add_heading(document, "Komponen LLM dan Output", level=1)
    rows = [
        ("Ekstraksi teks", "Membaca putusan yang punya text layer.", "Potongan header, pokok sengketa, posisi WP/DJP, pertimbangan, amar.", "JSON ekstraksi terstruktur."),
        ("Ekstraksi vision", "Membaca PDF scan.", "Gambar halaman awal, tengah, akhir.", "JSON ekstraksi terstruktur."),
        ("Labeling", "Klasifikasi pajak, isu, outcome, pihak, nilai sengketa.", "Teks/gambar dokumen.", "Metadata + ringkasan analitis."),
        ("Rekomendasi", "Menyusun rekomendasi untuk WP.", "Input kasus, pembanding, peraturan PPN, scoring.", "Draft rekomendasi mendalam."),
        ("Chat peraturan", "Menjawab pertanyaan aturan.", "Query user + hasil pencarian aturan.", "Jawaban aturan + lokasi/sumber."),
    ]
    table = document.add_table(rows=len(rows) + 1, cols=4)
    set_table_fixed(table, [1.35, 1.75, 2.05, 1.35])
    set_table_borders(table)
    for idx, header in enumerate(["Tahap", "Fungsi LLM", "Input utama", "Output"]):
        cell = table.cell(0, idx)
        set_cell_text(cell, header, bold=True, size=9.4, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(cell, LIGHT_GRAY)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for row_idx, row in enumerate(rows, start=1):
        set_row_cant_split(table.rows[row_idx])
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            set_cell_text(cell, value, bold=col_idx == 0, size=9.0)
            set_cell_shading(cell, "FFFFFF" if row_idx % 2 else "F9FBFD")
    document.add_paragraph()


def add_data_store_table(document):
    add_heading(document, "Data Store", level=1)
    rows = [
        ("documents", "Metadata dokumen, text layer, status ekstraksi, hasil utama."),
        ("document_extractions", "Field lengkap: metadata, objek sengketa, pihak, pokok sengketa, argumen, pertimbangan, outcome."),
        ("llm_labels", "Raw hasil labeling LLM dan confidence."),
        ("chunks", "Potongan teks dokumen untuk pencarian pembanding."),
        ("tax_regulations / regulation_chunks", "Peraturan PPN dari Ortax dan chunk pencarian."),
        ("analysis_reports", "Hasil analisis akhir yang dapat dibuka ulang dan diekspor."),
    ]
    table = document.add_table(rows=len(rows) + 1, cols=2)
    set_table_fixed(table, [2.0, 4.5])
    set_table_borders(table)
    for idx, header in enumerate(["Tabel", "Peran"]):
        cell = table.cell(0, idx)
        set_cell_text(cell, header, bold=True, size=9.5, color=DARK_BLUE)
        set_cell_shading(cell, LIGHT_GRAY)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for row_idx, row in enumerate(rows, start=1):
        set_row_cant_split(table.rows[row_idx])
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            set_cell_text(cell, value, bold=col_idx == 0, size=9.2)
            set_cell_shading(cell, "FFFFFF" if row_idx % 2 else "F9FBFD")
    document.add_paragraph()


def add_mermaid_appendix(document):
    add_heading(document, "Appendix - Mermaid Source", level=1)
    add_body(document, "Kode berikut disertakan agar diagram dapat dirender ulang di tools Markdown/Mermaid bila diperlukan.")
    code = '''flowchart TD
    A["User Input"] --> B["Document Intake"]
    B --> C{"Text cukup?"}
    C -- "Ya" --> D["LLM Text Extraction"]
    C -- "Tidak / scan" --> E["Render PDF Pages"]
    E --> F["LLM Vision Extraction"]
    D --> G["Structured JSON"]
    F --> G
    G --> H["Normalize + Duplicate Check"]
    H --> I["SQLite: documents / extractions / labels / chunks"]
    I --> J["Retrieval: Putusan + Peraturan PPN"]
    J --> K["Scoring & Risk Engine"]
    K --> L["LLM Draft Recommendation"]
    L --> M["Final UI + DOCX/PDF/Markdown"]'''
    paragraph = document.add_paragraph()
    run = paragraph.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Title", 24, BLUE, 0, 10),
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Tax Dispute Prototype - LLM Architecture")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def build():
    document = Document()
    configure_document(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("LLM Architecture - Tax Dispute Prototype")

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("Alur dari input dokumen/perkara sampai analisis akhir, rekomendasi WP, dan export dokumen.")
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string(MUTED)

    add_callout(
        document,
        "Tujuan dokumen",
        "Memberikan peta teknis yang mudah dibaca mengenai titik penggunaan LLM, data yang masuk/keluar, serta kontrol penyimpanan dan validasi.",
        fill=LIGHT_GREEN,
    )

    add_vertical_flow(document)
    add_sequence_table(document)
    add_component_table(document)
    add_data_store_table(document)

    add_heading(document, "Guardrail Saat Ini", level=1)
    for item in [
        "Dokumen duplikat ditolak berdasarkan nomor putusan yang sudah berhasil diekstrak.",
        "Jika PDF scan tidak punya text layer, sistem memakai jalur LLM vision.",
        "Jika LLM gagal membaca PDF scan, UI menampilkan error eksplisit.",
        "Draft rekomendasi bersifat indikatif dan perlu review ahli pajak/kuasa hukum.",
        "Nomor putusan, pasal, nama pihak, dan fakta baru tidak boleh dikarang oleh prompt LLM.",
    ]:
        add_bullet(document, item)

    add_heading(document, "Output Akhir", level=1)
    for item in [
        "Skor indikatif dan confidence.",
        "Review risiko, faktor pendukung, faktor risiko, dan celah bukti.",
        "Putusan pembanding dan peraturan PPN terkait.",
        "Draft rekomendasi WP yang bisa diedit.",
        "Dokumen unduhan dalam format DOCX, PDF, dan Markdown.",
    ]:
        add_numbered(document, item)

    document.add_section(WD_SECTION.NEW_PAGE)
    add_mermaid_appendix(document)

    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
